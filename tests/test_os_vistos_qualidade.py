import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


APP_DIR = Path(__file__).resolve().parents[1] / "compras_app"
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

from config import TEMPLATE_OS
from gerar_os import _inserir_vistos_qualidade_gestao, _isolar_layout_na_ultima_pagina
from os_template import mapear_tabelas_os


class VistosQualidadeGestaoTest(unittest.TestCase):
    def test_insere_vistos_antes_da_pagina_de_layout_sem_duplicar(self):
        doc = Document(TEMPLATE_OS)
        refs = mapear_tabelas_os(doc)
        tabela_layout = doc.tables[refs["layout"]]
        _isolar_layout_na_ultima_pagina(doc, tabela_layout)

        refs = mapear_tabelas_os(doc)
        _inserir_vistos_qualidade_gestao(doc, refs)
        _inserir_vistos_qualidade_gestao(doc, mapear_tabelas_os(doc))

        tabelas_vistos = [
            tabela
            for tabela in doc.tables
            if "VISTOS DE QUALIDADE E GESTÃO"
            in " ".join(cell.text for row in tabela.rows for cell in row.cells)
        ]
        self.assertEqual(1, len(tabelas_vistos))
        tabela = tabelas_vistos[0]
        self.assertEqual("Assinatura da QUALIDADE", tabela.cell(1, 0).text)
        self.assertEqual("Assinatura de GESTÃO", tabela.cell(1, 1).text)
        self.assertEqual("______________________________", tabela.cell(2, 0).text.strip())
        self.assertEqual("______________________________", tabela.cell(2, 1).text.strip())

        body = doc._body._element
        indice_vistos = body.index(tabela._tbl)
        indice_layout = body.index(tabela_layout._tbl)
        self.assertLess(indice_vistos, indice_layout)
        quebra_layout = tabela_layout._tbl.getprevious()
        self.assertEqual(qn("w:p"), quebra_layout.tag)
        self.assertTrue(quebra_layout.xpath('.//w:pageBreakBefore | .//w:br[@w:type="page"]'))
        self.assertLess(indice_vistos, body.index(quebra_layout))


if __name__ == "__main__":
    unittest.main()
