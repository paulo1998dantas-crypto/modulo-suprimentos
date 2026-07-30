from functools import wraps

from flask import Flask, render_template, request, send_file, redirect, url_for, after_this_request, session, jsonify
from flask.wrappers import Request as FlaskRequest
import csv
import hashlib
import re
import json
import io
import os
import urllib.error
import urllib.request
import logging
import tempfile
import sys
import shutil
import zipfile
import subprocess
import threading
import unicodedata
import uuid
from datetime import date, timedelta, datetime
import tempfile
import zipfile
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from werkzeug.utils import secure_filename
from werkzeug.exceptions import HTTPException
from werkzeug.datastructures import FileStorage
from openpyxl import load_workbook
from docx import Document

from config import (
    TEMPLATES_DIR,
    STATIC_DIR,
    DATA_DIR,
    PRODUTOS_FILE,
    FORNECEDORES_FILE,
    OS_PRODUTOS_FILE,
    OS_FORNECEDORES_FILE,
    OS_COMPONENTES_FILE,
    OS_PROCESSOS_FILE,
    OS_PROCESSO_RELACOES_FILE,
    OS_ITEM_POPUP_REGRAS_FILE,
    COUNTER_FILE,
    OS_COUNTER_FILE,
    HISTORICO_FILE,
    OC_IMPORT_FILE,
    OS_IMPORT_FILE,
    get_save_paths,
    set_save_paths,
    get_bom_dir,
    set_bom_dir,
    get_skus_file,
    set_skus_file,
    get_processos_dir,
    set_processos_dir,
    pasta_os,
)
from calculos import calcular_total_item
from composicao import (
    expandir_composicao_referenciada,
    normalizar_codigo,
    normalizar_componentes,
    parse_quantidade,
    resolver_composicao_final,
)
from gerar_oc import gerar_word, construir_nome_oc
from gerar_os import gerar_os_docx
from gerar_op import build_production_order_docx
from os_template import encontrar_linha_cabecalho, mapear_tabelas_os
from processos_os import PROCESSOS_ORDEM, PROCESSOS_OS, PROCESSOS_POR_KEY, identificar_nome_processo, normalizar_nome_processo
from os_setores import (
    SETOR_EXPEDICAO,
    SETOR_FATURAMENTO_DIRETO,
    SETOR_PREPARACAO,
    TIPO_REQUISICAO_FATURAMENTO_DIRETO,
    agrupar_linhas_setor,
    agrupar_linhas_por_fornecedor,
    construir_itens_os_expedicao,
    construir_itens_os_preparacao,
    construir_itens_os_setor,
    enriquecer_composicao,
    filtrar_linhas_faturamento_direto,
    filtrar_linhas_preparacao,
    filtrar_linhas_setor,
    linhas_layout_preparacao,
    propagar_setor_preparacao,
)
from processos_transformacao import (
    PROCESSO_POR_ITEM,
    RELACOES_PROCESSO_TRANSFORMACAO,
    construir_processo_por_item,
    resolver_processo_transformacao,
    resolver_processos_transformacao,
)
import supabase_catalog
import supabase_data


def _load_local_env():
    """Developer-only local settings; never overrides Render environment values."""
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env.local")
    if not os.path.isfile(path):
        return
    with open(path, "r", encoding="utf-8") as handle:
        for raw in handle:
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


_load_local_env()


def _positive_env_int(name, default):
    try:
        value = int(os.environ.get(name, default))
    except (TypeError, ValueError):
        return default
    return value if value > 0 else default


class SuprimentosRequest(FlaskRequest):
    max_form_memory_size = _positive_env_int("SUPRIMENTOS_MAX_FORM_MEMORY_BYTES", 32 * 1024 * 1024)
    max_form_parts = _positive_env_int("SUPRIMENTOS_MAX_FORM_PARTS", 20_000)


app = Flask(__name__, template_folder=TEMPLATES_DIR, static_folder=STATIC_DIR)
app.secret_key = os.environ.get("SUPRIMENTOS_SESSION_SECRET", "").strip() or "emissor_documentos"
app.request_class = SuprimentosRequest
app.config["MAX_CONTENT_LENGTH"] = _positive_env_int("SUPRIMENTOS_MAX_REQUEST_BYTES", 64 * 1024 * 1024)
app.config["MAX_FORM_MEMORY_SIZE"] = SuprimentosRequest.max_form_memory_size
app.config["MAX_FORM_PARTS"] = SuprimentosRequest.max_form_parts


def _env_bool(name, default=False):
    value = os.environ.get(name)
    if value is None:
        return default
    return value.strip().lower() not in {"0", "false", "nao", "não", "no", "n"}


def erp_feature_enabled():
    return _env_bool("ERP_FEATURE_FLAG", default=False)


def erp_feature_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if erp_feature_enabled():
            return view(*args, **kwargs)
        if request.path.startswith("/api/"):
            return jsonify({"ok": False, "error": "Integração ERP desativada pela feature flag."}), 404
        return "Integração ERP desativada pela feature flag.", 404
    return wrapped


def login_enabled():
    return _env_bool("SUPRIMENTOS_REQUIRE_LOGIN", default=supabase_data.enabled())


def shared_rbac_enabled():
    return _env_bool("ERP_SHARED_RBAC_ENABLED", default=False)


def current_user():
    return session.get("suprimentos_user")


def current_username():
    user = current_user() or {}
    return str(user.get("username") or user.get("id") or "local").strip() or "local"


def current_user_id():
    user = current_user() or {}
    value = user.get("id")
    return "" if value is None else str(value).strip()


def can(permission):
    """Return the effective permission while preserving the pre-RBAC rollout."""
    if not shared_rbac_enabled():
        return True
    user = current_user() or {}
    permissions = set(user.get("permissions") or [])
    return "*" in permissions or str(permission or "").strip() in permissions


def can_any(*permissions):
    return any(can(permission) for permission in permissions)


def _authorization_denied(permission):
    message = f"Seu perfil nao possui a permissao necessaria: {permission}."
    if request.path.startswith("/api/"):
        return jsonify({"ok": False, "error": message}), 403
    return message, 403


def permission_required(permission):
    def decorator(view):
        @wraps(view)
        def wrapped(*args, **kwargs):
            try:
                required_permission = (
                    permission(*args, **kwargs)
                    if callable(permission)
                    else permission
                )
            except Exception:
                app.logger.exception(
                    "Falha ao resolver a permissao da rota %s",
                    request.path,
                )
                if request.path.startswith("/api/"):
                    return jsonify({
                        "ok": False,
                        "error": "Nao foi possivel validar a autorizacao.",
                    }), 503
                return "Nao foi possivel validar a autorizacao.", 503
            required_permissions = (
                list(required_permission)
                if isinstance(required_permission, (list, tuple, set))
                else [required_permission]
            )
            for item in required_permissions:
                if not can(item):
                    return _authorization_denied(item)
            return view(*args, **kwargs)
        return wrapped
    return decorator


app.jinja_env.globals["can"] = can
app.jinja_env.globals["can_any"] = can_any


def _user_scoped_file(path):
    """Keep transient import data isolated per application login."""
    stem, ext = os.path.splitext(path)
    username = secure_filename(current_username()) or "local"
    return f"{stem}.{username}{ext}"


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if login_enabled() and not current_user():
            return redirect(url_for("login", next=request.path))
        return view(*args, **kwargs)

    return wrapped


@app.before_request
def require_login_global():
    if not login_enabled():
        return None
    public_endpoints = {"login", "logout", "healthz", "static"}
    if request.endpoint in public_endpoints:
        return None
    user = current_user()
    if user:
        if shared_rbac_enabled():
            try:
                fresh_user = supabase_data.revalidate_session_user(user)
            except Exception:
                app.logger.exception("Falha ao revalidar a autorizacao compartilhada")
                if request.path.startswith("/api/"):
                    return jsonify({"ok": False, "error": "Autorizacao temporariamente indisponivel."}), 503
                return "Autorizacao temporariamente indisponivel.", 503
            if not fresh_user:
                session.clear()
                if request.path.startswith("/api/"):
                    return jsonify({
                        "ok": False,
                        "error": "Sessao expirada apos alteracao de usuario ou acesso.",
                    }), 401
                return redirect(url_for("login", next=request.path))
            session["suprimentos_user"] = fresh_user
        return None
    return redirect(url_for("login", next=request.full_path if request.query_string else request.path))

RELEASE_BUILD_NAME = "ModuloSuprimentos"
RELEASE_ENVIO_DIR_NAME = "ModuloSuprimentos_envio"
RELEASE_ZIP_NAME = "ModuloSuprimentos_envio.zip"
RELEASE_BUILD_SCRIPT = "gerar_modulo_suprimentos_envio.bat"
ITEM_CAMPOS_BASE = [
    "descricao",
    "unidade",
    "grupo",
    "categoria",
    "processo_conjunto",
]
MODELO_ITENS_HEADERS = [
    "CODIGO",
    "DESCRICAO",
    "UNIDADE",
    "GRUPO",
    "CATEGORIA",
    "PROCESSO CONJUNTO",
]
MODELO_REGRAS_POPUP_HEADERS = [
    "ID_REGRA",
    "ITEM_GATILHO",
    "DESCRICAO_GATILHO",
    "ITENS_OPCOES",
    "DESCRICOES_OPCOES",
    "QUANTIDADE",
    "QUANTIDADE_EDITAVEL",
]
MODELO_RELACOES_PROCESSO_HEADERS = [
    "ITEM_CODIGO",
    "DESCRICAO_ITEM",
    "PROCESSOS",
]
HEADER_ALIASES = {
    "codigo": {
        "codigo",
        "cod",
        "cod_item",
        "codigo_item",
        "codigo_produto",
        "novo_cod",
        "novo_codigo",
        "cod_novo",
        "novo_cod_",
    },
    "descricao": {
        "descricao",
        "descricao_item",
        "descricao_produto",
        "descricao_primaria",
        "descricao_principal",
        "item_descricao",
        "produto",
        "material",
    },
    "descricao_secundaria": {
        "descricao_secundaria",
        "descricao_complementar",
        "complemento_descricao",
    },
    "sufixo": {"sufixo", "sufixo_descricao"},
    "unidade": {
        "unidade",
        "un",
        "um",
        "und",
        "unidade_medida",
        "un_medida",
        "un_medi_interna",
        "unidade_interna",
        "un_medi_comercial",
        "unidade_comercial",
    },
    "grupo": {"grupo", "grupo_produto"},
    "categoria": {"categoria", "categorias", "classificacao", "classificacao_item"},
    "processo_conjunto": {
        "processo_conjunto",
        "processo",
        "processo_vinculado",
        "arquivo_processo",
    },
    "cliente": {"cliente"},
}

CAMPOS_PRODUTO_DESCARTADOS = {
    "unidade_comercial",
    "unidade_interna",
    "tipo",
    "fornecedor",
    "ncm",
    "origem",
    "valor",
    "ipi",
    "icms",
    "cofins",
    "observacao",
}

def _is_in_dir(path, base):
    try:
        return os.path.commonpath([os.path.abspath(path), os.path.abspath(base)]) == os.path.abspath(base)
    except Exception:
        return False


def _sanitize_output_name(texto):
    texto = "" if texto is None else str(texto).strip()
    if not texto:
        return ""
    texto = re.sub(r'[<>:"/\\\\|?*]+', " ", texto)
    return " ".join(texto.split())


def _limpar_valor_busca(texto):
    texto = "" if texto is None else str(texto).strip()
    if not texto:
        return ""
    match = re.match(r"^(.*?)\s*\(([^)]+)\)\s*$", texto)
    if match:
        texto = match.group(1).strip() or match.group(2).strip()
    return " ".join(texto.split())


def _win_long_path(path):
    if os.name != "nt":
        return path
    if not path:
        return path

    abs_path = os.path.abspath(path)
    if abs_path.startswith("\\\\?\\"):
        return abs_path
    if abs_path.startswith("\\\\"):
        return "\\\\?\\UNC\\" + abs_path.lstrip("\\")
    return "\\\\?\\" + abs_path


def _open_for_read(path):
    return open(_win_long_path(path), "rb")


def _setup_logging():
    if os.environ.get("SUPRIMENTOS_FILE_LOG", "1").strip().lower() in {
        "0", "false", "no", "nao", "não", "off",
    }:
        app.logger.setLevel(logging.INFO)
        return
    log_dir = DATA_DIR
    try:
        os.makedirs(log_dir, exist_ok=True)
        log_path = os.path.join(log_dir, "emissor_documentos.log")
        with open(log_path, "a", encoding="utf-8"):
            pass
    except Exception:
        log_dir = tempfile.gettempdir()
        log_path = os.path.join(log_dir, "emissor_documentos.log")

    from logging.handlers import RotatingFileHandler
    handler = RotatingFileHandler(log_path, maxBytes=2_000_000, backupCount=3, encoding="utf-8")
    formatter = logging.Formatter("%(asctime)s %(levelname)s %(message)s")
    handler.setFormatter(formatter)
    app.logger.addHandler(handler)
    app.logger.setLevel(logging.INFO)
    app.logger.info("Log iniciado em %s", log_path)


_setup_logging()


@app.errorhandler(413)
def handle_request_too_large(_error):
    app.logger.warning(
        "Requisicao excedeu o limite em %s: content_length=%s",
        request.path,
        request.content_length,
    )
    return (
        "A O.S. ultrapassou o limite de envio do servidor. "
        "O rascunho continua salvo neste navegador. Atualize a pagina e tente novamente.",
        413,
    )


@app.errorhandler(Exception)
def handle_exception(e):
    if isinstance(e, HTTPException):
        return e
    if isinstance(e, PermissionError):
        return str(e), 500
    app.logger.exception("Erro nao tratado")
    if request.path in {"/gerar_os", "/gerar_oc"}:
        erro_id = datetime.now().strftime("%Y%m%d%H%M%S")
        return (
            "Erro ao emitir documento. "
            f"Informe este codigo ao suporte: {erro_id}. "
            f"{type(e).__name__}: {str(e)}"
        ), 500
    return "Internal Server Error", 500


def _ensure_json_file(path, default):
    if os.path.exists(path):
        return
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(default, f, ensure_ascii=False, indent=2)


def _ensure_text_file(path, content):
    if os.path.exists(path):
        return
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(str(content))


def ensure_data_storage():
    os.makedirs(DATA_DIR, exist_ok=True)
    # Para EXE: sempre sobrescreve a base local com a base embutida no build
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        packaged_data = os.path.join(sys._MEIPASS, "compras_app", "data")
        if os.path.isdir(packaged_data):
            for name in os.listdir(packaged_data):
                src = os.path.join(packaged_data, name)
                dst = os.path.join(DATA_DIR, name)
                if os.path.isfile(src):
                    try:
                        shutil.copy2(src, dst)
                    except Exception:
                        pass
    _ensure_json_file(PRODUTOS_FILE, {})
    _ensure_json_file(FORNECEDORES_FILE, {})
    _ensure_json_file(OS_PRODUTOS_FILE, {})
    _ensure_json_file(OS_FORNECEDORES_FILE, {})
    _ensure_json_file(OS_COMPONENTES_FILE, {})
    _ensure_json_file(OS_PROCESSOS_FILE, {})
    _ensure_json_file(OS_PROCESSO_RELACOES_FILE, {})
    _ensure_json_file(OS_ITEM_POPUP_REGRAS_FILE, [])
    _ensure_json_file(HISTORICO_FILE, [])
    _ensure_json_file(OC_IMPORT_FILE, {})
    _ensure_json_file(OS_IMPORT_FILE, {})
    _ensure_text_file(COUNTER_FILE, "1")
    _ensure_text_file(OS_COUNTER_FILE, "1")


ensure_data_storage()


def _find_release_source_dir():
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return os.path.dirname(sys.executable)
    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    envio_dir = os.path.join(project_root, RELEASE_ENVIO_DIR_NAME)
    if os.path.isdir(envio_dir) and os.path.isdir(os.path.join(envio_dir, "_internal")):
        return envio_dir
    dist_dir = os.path.join(project_root, "dist")
    if os.path.isdir(dist_dir):
        preferred_names = [RELEASE_BUILD_NAME, "Emissor documentos"]
        for preferred_name in preferred_names:
            preferred = os.path.join(dist_dir, preferred_name)
            if os.path.isdir(preferred):
                return preferred
        for name in os.listdir(dist_dir):
            cand = os.path.join(dist_dir, name)
            if os.path.isdir(cand) and os.path.isdir(os.path.join(cand, "_internal")):
                return cand
    return None


def _run_release_build():
    if getattr(sys, "frozen", False):
        return True, None
    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    script_candidates = [
        RELEASE_BUILD_SCRIPT,
        "gerar_exe.bat",
    ]
    bat_path = None
    for script_name in script_candidates:
        candidate = os.path.join(project_root, script_name)
        if os.path.isfile(candidate):
            bat_path = candidate
            break
    if not bat_path:
        return False, f"Arquivo {RELEASE_BUILD_SCRIPT} nao encontrado no projeto."
    try:
        result = subprocess.run(
            f"\"{bat_path}\"",
            cwd=project_root,
            shell=True,
            capture_output=True,
            text=True,
        )
    except Exception as exc:
        return False, f"Falha ao executar gerar_exe.bat: {exc}"
    if result.returncode != 0:
        msg = (result.stdout or "") + "\n" + (result.stderr or "")
        return False, f"Erro ao gerar EXE. Saida:\n{msg.strip()}"
    return True, None


def _build_release_zip():
    source_dir = _find_release_source_dir()
    if not source_dir:
        return None
    temp_root = tempfile.mkdtemp(prefix="emissor_pack_")
    staging = os.path.join(temp_root, RELEASE_ENVIO_DIR_NAME)
    shutil.copytree(source_dir, staging, dirs_exist_ok=True)

    zip_path = os.path.join(temp_root, RELEASE_ZIP_NAME)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, _, files in os.walk(staging):
            for fname in files:
                full = os.path.join(root, fname)
                rel = os.path.relpath(full, staging)
                zf.write(full, rel)
    return zip_path, temp_root


def _reset_base_data():
    arquivos = [
        PRODUTOS_FILE,
        FORNECEDORES_FILE,
        OS_PRODUTOS_FILE,
        OS_FORNECEDORES_FILE,
        OS_COMPONENTES_FILE,
        OS_PROCESSOS_FILE,
        OS_PROCESSO_RELACOES_FILE,
        OS_ITEM_POPUP_REGRAS_FILE,
        HISTORICO_FILE,
        OC_IMPORT_FILE,
        OS_IMPORT_FILE,
        COUNTER_FILE,
        OS_COUNTER_FILE,
    ]
    for path in arquivos:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass
    ensure_data_storage()


def _carregar_historico_local():
    if not os.path.exists(HISTORICO_FILE):
        return []
    try:
        with open(HISTORICO_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return [dict(item, id=item.get("id") or f"local-{idx}") for idx, item in enumerate(data)]
        if isinstance(data, dict):
            merged = []
            for _, items in data.items():
                if isinstance(items, list):
                    merged.extend(items)
            return [dict(item, id=item.get("id") or f"local-{idx}") for idx, item in enumerate(merged)]
    except Exception:
        return []
    return []


def carregar_historico():
    if supabase_data.enabled():
        try:
            return supabase_data.carregar_documentos()
        except Exception:
            app.logger.exception("Falha ao carregar historico de documentos do Supabase")
    return _carregar_historico_local()


def salvar_historico(entries):
    try:
        with open(HISTORICO_FILE, "w", encoding="utf-8") as f:
            json.dump(entries, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def _line_id(prefix):
    safe_prefix = re.sub(r"[^a-z0-9_-]+", "-", str(prefix or "linha").lower()).strip("-") or "linha"
    return f"{safe_prefix}-{uuid.uuid4().hex}"


def _normalizar_line_id(value):
    texto = str(value or "").strip()
    return texto if texto else ""


def _linha_signature(linha, campos):
    if not isinstance(linha, dict):
        return ()
    return tuple(str(linha.get(campo, "") or "").strip().lower() for campo in campos or ())


def _ids_existentes_por_signature(linhas, campos):
    por_signature = {}
    for linha in linhas or []:
        if not isinstance(linha, dict):
            continue
        line_id = _normalizar_line_id(
            linha.get("line_id") or linha.get("linha_id") or linha.get("id_linha") or linha.get("lineId")
        )
        if not line_id:
            continue
        por_signature.setdefault(_linha_signature(linha, campos), []).append(line_id)
    return por_signature


def _linhas_existentes_por_line_id(linhas):
    por_line_id = {}
    for linha in linhas or []:
        if not isinstance(linha, dict):
            continue
        line_id = _normalizar_line_id(
            linha.get("line_id") or linha.get("linha_id") or linha.get("id_linha") or linha.get("lineId")
        )
        if line_id:
            por_line_id[line_id] = linha
    return por_line_id


def _atribuir_line_ids(linhas, prefix, existentes=None, campos_chave=None):
    campos_chave = tuple(campos_chave or ())
    ids_por_signature = _ids_existentes_por_signature(existentes, campos_chave)
    existentes_por_line_id = _linhas_existentes_por_line_id(existentes)
    resultado = []
    usados = set()
    for linha in linhas or []:
        if not isinstance(linha, dict):
            continue
        item = dict(linha)
        line_id = _normalizar_line_id(
            item.get("line_id") or item.get("linha_id") or item.get("id_linha") or item.get("lineId")
        )
        if not line_id:
            candidatos = ids_por_signature.get(_linha_signature(item, campos_chave), [])
            while candidatos and candidatos[0] in usados:
                candidatos.pop(0)
            if candidatos:
                line_id = candidatos.pop(0)
        while not line_id or line_id in usados:
            line_id = _line_id(prefix)
        item["line_id"] = line_id
        existente = existentes_por_line_id.get(line_id)
        if existente:
            for key in ("line_status", "line_status_atualizado_por", "line_status_atualizado_em"):
                if not item.get(key) and existente.get(key):
                    item[key] = existente.get(key)
        usados.add(line_id)
        resultado.append(item)
    return resultado


def _atribuir_line_ids_processos(processos, existentes=None):
    resultado = {}
    existentes = existentes if isinstance(existentes, dict) else {}
    for grupo, linhas in (processos or {}).items():
        resultado[grupo] = _atribuir_line_ids(
            linhas or [],
            "os-proc",
            existentes.get(grupo),
            ("atividade", "responsavel"),
        )
    return resultado


def registrar_historico(
    tipo,
    numero,
    dados,
    itens=None,
    processos=None,
    componentes=None,
    composicao=None,
    documento_id=None,
    status="emitido",
    submit_token="",
):
    existente = obter_historico_documento(documento_id) if documento_id else None
    usuario = current_username()
    itens = _atribuir_line_ids(
        itens or [],
        f"{tipo}-item",
        (existente or {}).get("itens"),
        ("codigo", "descricao", "qtd", "unidade", "serie"),
    )
    composicao = _atribuir_line_ids(
        composicao or [],
        f"{tipo}-comp",
        (existente or {}).get("composicao"),
        ("item", "codigo", "descricao", "qtd", "unidade", "level", "setor"),
    )
    processos = _atribuir_line_ids_processos(processos or {}, (existente or {}).get("processos"))
    entry = {
        "tipo": tipo,
        "numero": str(numero),
        "data_criacao": (existente or {}).get("data_criacao") or datetime.now().strftime("%Y-%m-%d"),
        "status": status or (existente or {}).get("status") or "emitido",
        "submit_token": submit_token or (existente or {}).get("submit_token") or "",
        "criado_por": (existente or {}).get("criado_por") or usuario,
        "atualizado_por": usuario,
        "erp_purchase_order_id": (existente or {}).get("erp_purchase_order_id") or None,
        "erp_work_order_id": (existente or {}).get("erp_work_order_id") or None,
        "dados": dados or {},
        "itens": itens,
        "processos": processos,
        "componentes": componentes or {},
        "composicao": composicao,
    }
    if documento_id:
        entry["id"] = documento_id
    if supabase_data.enabled():
        try:
            if documento_id:
                supabase_data.atualizar_documento(documento_id, entry)
            else:
                salvo = supabase_data.salvar_documento(entry)
                if isinstance(salvo, dict) and salvo.get("id") is not None:
                    entry["id"] = salvo["id"]
        except Exception:
            app.logger.exception("Falha ao salvar historico de documento no Supabase")
            raise
        return entry
    entries = _carregar_historico_local()
    atualizado = False
    for idx, registro_local in enumerate(entries):
        mesmo_id = documento_id and str(registro_local.get("id")) == str(documento_id)
        mesmo_token = submit_token and registro_local.get("submit_token") == submit_token
        if mesmo_id or mesmo_token:
            if not entry.get("id"):
                entry["id"] = registro_local.get("id")
            if not entry.get("criado_por"):
                entry["criado_por"] = registro_local.get("criado_por", "")
            entries[idx] = entry
            atualizado = True
            break
    if not atualizado:
        entry.setdefault("id", f"local-{submit_token or len(entries)}")
        entries.append(entry)
    salvar_historico(entries)
    return entry


def obter_historico_documento(documento_id):
    if supabase_data.enabled():
        try:
            return supabase_data.obter_documento(documento_id)
        except Exception:
            app.logger.exception("Falha ao consultar documento no Supabase")
    for entry in _carregar_historico_local():
        if str(entry.get("id")) == str(documento_id):
            return entry
    return None


def obter_historico_por_submit_token(submit_token):
    token = str(submit_token or "").strip()
    if not token:
        return None
    if supabase_data.enabled():
        return supabase_data.obter_documento_por_submit_token(token)
    for entry in _carregar_historico_local():
        if str(entry.get("submit_token") or "").strip() == token:
            return entry
    return None


def excluir_historico_documento(documento_id):
    if supabase_data.enabled():
        supabase_data.excluir_documento(documento_id)
        return
    entries = [
        entry
        for entry in _carregar_historico_local()
        if str(entry.get("id")) != str(documento_id)
    ]
    salvar_historico(entries)


def atualizar_status_historico_documento(documento_id, status):
    status = str(status or "").strip().lower()
    if status not in {"rascunho", "emitido", "cancelado", "concluido"}:
        raise ValueError("Status de documento invalido.")
    documento = obter_historico_documento(documento_id)
    if not documento:
        return None
    documento["status"] = status
    documento["atualizado_por"] = current_username()
    if supabase_data.enabled():
        supabase_data.atualizar_documento(documento_id, documento)
        return documento
    entries = _carregar_historico_local()
    for idx, entry in enumerate(entries):
        if str(entry.get("id")) == str(documento_id):
            entries[idx] = documento
            break
    salvar_historico(entries)
    return documento


def salvar_historico_documento_atualizado(documento_id, documento):
    if not documento:
        return None
    documento["atualizado_por"] = current_username()
    if supabase_data.enabled():
        supabase_data.atualizar_documento(documento_id, documento)
        return documento
    entries = _carregar_historico_local()
    atualizado = False
    for idx, entry in enumerate(entries):
        if str(entry.get("id")) == str(documento_id):
            entries[idx] = documento
            atualizado = True
            break
    if not atualizado:
        return None
    salvar_historico(entries)
    return documento


def vincular_documento_erp(documento, campo, entity_id):
    """Persist an immutable ERP id inside the existing JSON document.

    Keeping this link in ``dados`` is backward compatible with the current
    Supabase table and avoids matching new records by visible O.C./O.S. number.
    """
    if not documento or not entity_id:
        return documento
    documento = dict(documento)
    dados = dict(documento.get("dados") or {})
    dados[campo] = str(entity_id)
    documento["dados"] = dados
    if campo in {"erp_purchase_order_id", "erp_work_order_id"}:
        documento[campo] = str(entity_id)
    documento_id = documento.get("id")
    if documento_id:
        salvar_historico_documento_atualizado(documento_id, documento)
    return documento


def _agrupar_por_data(entries, tipo, campo_soma=None):
    resumo = {}
    for e in entries:
        if e.get("tipo") != tipo:
            continue
        if e.get("status") == "cancelado":
            continue
        data = (e.get("data_criacao") or "")[:10]
        if not data:
            continue
        if data not in resumo:
            resumo[data] = 0
        if campo_soma:
            valor = e.get("dados", {}).get(campo_soma, 0)
            try:
                resumo[data] += float(valor)
            except Exception:
                pass
        else:
            resumo[data] += 1
    return [{"data": k, "valor": resumo[k]} for k in sorted(resumo.keys())]


def _dashboard_totais(entries):
    total_oc = 0.0
    qtd_oc = 0
    qtd_os = 0
    for entry in entries:
        if entry.get("status") == "cancelado":
            continue
        if entry.get("tipo") == "oc":
            qtd_oc += 1
            total_oc += _parse_numero_form((entry.get("dados") or {}).get("total_pedido"), 0.0)
        elif entry.get("tipo") == "os":
            qtd_os += 1
    return {"qtd_oc": qtd_oc, "qtd_os": qtd_os, "total_oc": total_oc}


def _dashboard_recentes(entries, limit=20):
    recentes = sorted(
        entries,
        key=lambda item: (str(item.get("data_criacao") or ""), str(item.get("numero") or "")),
        reverse=True,
    )
    return recentes[:limit]


def _dashboard_documentos(entries):
    documentos = []
    for entry in entries or []:
        tipo = str(entry.get("tipo") or "").strip().lower()
        if tipo not in {"oc", "os"}:
            continue
        dados = entry.get("dados") or {}
        documentos.append({
            "id": str(entry.get("id") or ""),
            "tipo": tipo,
            "numero": str(entry.get("numero") or ""),
            "data": str(entry.get("data_criacao") or "")[:10],
            "status": str(entry.get("status") or "emitido").strip().lower() or "emitido",
            "nome": str(
                dados.get("cliente")
                if tipo == "os"
                else dados.get("fornecedor") or dados.get("razao_social")
                or ""
            ),
            "chassis": str(dados.get("chassis") or ""),
            "mmv": str(dados.get("mmv") or ""),
            "total": _parse_numero_form(dados.get("total_pedido"), 0.0) if tipo == "oc" else 0.0,
            "itens": len(entry.get("itens") or []),
            "ordem": str(entry.get("updated_at") or entry.get("created_at") or ""),
        })
    return documentos


def _get_free_port():
    import socket
    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sock.bind(("127.0.0.1", 0))
    port = sock.getsockname()[1]
    sock.close()
    return port


def _abrir_navegador(port):
    try:
        import webbrowser
        webbrowser.open_new(f"http://127.0.0.1:{port}")
    except Exception:
        pass




def carregar_produtos():
    if supabase_catalog.enabled():
        try:
            return supabase_catalog.carregar_produtos()
        except Exception as exc:
            app.logger.exception("Falha ao carregar SKUs do Supabase")
            return {}

    if not os.path.exists(PRODUTOS_FILE):
        return {}

    with open(PRODUTOS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def carregar_fornecedores():
    if supabase_data.enabled():
        try:
            return supabase_data.carregar_pessoas("fornecedor")
        except Exception:
            app.logger.exception("Falha ao carregar fornecedores do Supabase")
            return {}

    if not os.path.exists(FORNECEDORES_FILE):
        return {}

    with open(FORNECEDORES_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def carregar_os_produtos():
    if supabase_catalog.enabled():
        return carregar_produtos()

    if not os.path.exists(OS_PRODUTOS_FILE):
        return {}

    with open(OS_PRODUTOS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def carregar_os_fornecedores():
    if supabase_data.enabled():
        try:
            return supabase_data.carregar_pessoas("cliente")
        except Exception:
            app.logger.exception("Falha ao carregar clientes do Supabase")
            return {}

    if not os.path.exists(OS_FORNECEDORES_FILE):
        return {}

    with open(OS_FORNECEDORES_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def _resolver_nome_cliente_os(valor, clientes=None):
    valor = _limpar_valor_busca(valor)
    if not valor:
        return ""

    clientes = clientes if isinstance(clientes, dict) else carregar_os_fornecedores()
    info = clientes.get(valor)
    if not isinstance(info, dict):
        valor_normalizado = valor.casefold()
        for chave, candidato in clientes.items():
            if str(chave).strip().casefold() == valor_normalizado:
                info = candidato
                break
            if not isinstance(candidato, dict):
                continue
            nomes = (
                candidato.get("cliente"),
                candidato.get("nome_fantasia"),
                candidato.get("razao_social"),
            )
            if any(str(nome or "").strip().casefold() == valor_normalizado for nome in nomes):
                info = candidato
                break

    if not isinstance(info, dict):
        return valor
    return _limpar_valor_busca(
        info.get("cliente")
        or info.get("nome_fantasia")
        or info.get("razao_social")
        or valor
    )


def salvar_fornecedores(fornecedores):
    if supabase_data.enabled():
        supabase_data.salvar_pessoas_legacy(fornecedores or {}, "fornecedor")
        return
    salvar_json(FORNECEDORES_FILE, fornecedores or {})


def salvar_os_fornecedores(fornecedores):
    if supabase_data.enabled():
        supabase_data.salvar_pessoas_legacy(fornecedores or {}, "cliente")
        return
    salvar_json(OS_FORNECEDORES_FILE, fornecedores or {})


def carregar_os_componentes():
    if supabase_data.enabled():
        try:
            return supabase_data.carregar_bom_componentes()
        except Exception:
            app.logger.exception("Falha ao carregar B.O.M. do Supabase")
            return {}

    if not os.path.exists(OS_COMPONENTES_FILE):
        return {}

    with open(OS_COMPONENTES_FILE, "r", encoding="utf-8") as f:
        return normalizar_componentes(json.load(f))


def carregar_os_processos():
    if supabase_data.enabled():
        try:
            data = supabase_data.carregar_processos()
        except Exception:
            app.logger.exception("Falha ao carregar processos do Supabase")
            data = {}
        for conjunto in list(data):
            for nome in PROCESSOS_ORDEM:
                data[conjunto].setdefault(nome, [])
        return data

    if not os.path.exists(OS_PROCESSOS_FILE):
        return {}

    with open(OS_PROCESSOS_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)

    # suporta formato antigo: {processo: [linhas]}
    if data and isinstance(next(iter(data.values())), list):
        data = {"PADRAO": data}

    def normalizar_processo(nome):
        if not nome:
            return nome
        n = str(nome).strip().upper()
        mapa = {
            "CORTE": "CORTE",
            "AR CONDICIONADO": "AR CONDICIONADO",
            "PREPARACAO DE PECAS": "PREPARAÇÃO DE PEÇAS",
            "PREPARAÇÃO DE PEÇAS": "PREPARAÇÃO DE PEÇAS",
            "PREPARA??O DE PE?AS": "PREPARAÇÃO DE PEÇAS",
            "PREPARAÃ‡ÃƒO DE PEÃ‡AS": "PREPARAÇÃO DE PEÇAS",
            "ISOLAMENTO": "ISOLAMENTO",
            "DESMONTAGEM E ISOLAMENTO": "ISOLAMENTO",
            "REVESTIMENTO": "REVESTIMENTO",
            "BANCOS": "BANCOS",
            "ELETRICA": "ELÉTRICA 2",
            "ELÉTRICA": "ELÉTRICA 2",
            "ELÉTRICA 2": "ELÉTRICA 2",
            "EL?TRICA 2": "ELÉTRICA 2",
            "ELÃ‰TRICA 2": "ELÉTRICA 2",
            "LIMPEZA/LIBERACAO": "LIMPEZA/LIBERAÇÃO",
            "LIMPEZA/LIBERAÇÃO": "LIMPEZA/LIBERAÇÃO",
            "LIMPEZA/LIBERA??O": "LIMPEZA/LIBERAÇÃO",
            "LIMPEZA/LIBERAÃ‡ÃƒO": "LIMPEZA/LIBERAÇÃO",
        }
        return mapa.get(n, nome)

    normalizado = {}
    for conjunto, processos in data.items():
        normalizado.setdefault(conjunto, {})
        for nome_proc, linhas in processos.items():
            chave = normalizar_processo(nome_proc)
            normalizado[conjunto].setdefault(chave, [])
            normalizado[conjunto][chave].extend(linhas)

    return normalizado


def carregar_relacoes_processo_item():
    if supabase_data.enabled():
        try:
            return supabase_data.carregar_relacoes()
        except Exception:
            app.logger.exception("Falha ao carregar relacoes processo x item do Supabase")
            return {}

    if not os.path.exists(OS_PROCESSO_RELACOES_FILE):
        return {}
    try:
        with open(OS_PROCESSO_RELACOES_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return {}
    if isinstance(data, dict):
        relacoes = {}
        for chave, valor in data.items():
            codigo = normalizar_codigo(chave)
            processos = valor if isinstance(valor, list) else [valor]
            processos = [
                str(processo or "").strip()
                for processo in processos
                if str(processo or "").strip()
            ]
            if codigo and processos:
                relacoes[codigo] = list(dict.fromkeys(processos))
        return relacoes
    if isinstance(data, list):
        relacoes = {}
        for relacao in data:
            if not isinstance(relacao, dict):
                continue
            codigo = normalizar_codigo(
                relacao.get("codigo")
                or relacao.get("item_codigo")
                or relacao.get("item")
                or ""
            )
            processos = relacao.get("processos") or relacao.get("opcoes") or [
                relacao.get("processo_conjunto") or relacao.get("processo") or ""
            ]
            if not isinstance(processos, list):
                processos = [processos]
            processos = [
                str(processo or "").strip()
                for processo in processos
                if str(processo or "").strip()
            ]
            if codigo and processos:
                relacoes[codigo] = list(dict.fromkeys(processos))
        return relacoes
    return {}


def salvar_relacoes_processo_item(relacoes):
    if supabase_data.enabled():
        supabase_data.salvar_relacoes(relacoes or {})
        return
    salvar_json(OS_PROCESSO_RELACOES_FILE, relacoes or {})


def carregar_regras_popup_item():
    if supabase_data.enabled():
        try:
            return supabase_data.carregar_regras()
        except Exception:
            app.logger.exception("Falha ao carregar parametros de item relacionado do Supabase")
            return []

    if not os.path.exists(OS_ITEM_POPUP_REGRAS_FILE):
        return []
    try:
        with open(OS_ITEM_POPUP_REGRAS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return []
    if not isinstance(data, list):
        return []
    regras = []
    for idx, regra in enumerate(data):
        if not isinstance(regra, dict):
            continue
        gatilho = normalizar_codigo(regra.get("gatilho") or regra.get("item_gatilho") or "")
        opcoes = [
            normalizar_codigo(codigo)
            for codigo in (regra.get("opcoes") or [])
            if normalizar_codigo(codigo)
        ]
        if not gatilho or not opcoes:
            continue
        try:
            quantidade = float(regra.get("quantidade", 1) or 1)
        except Exception:
            quantidade = 1
        if quantidade <= 0:
            quantidade = 1
        regras.append(
            {
                "id": str(regra.get("id") or f"regra-{idx + 1}"),
                "gatilho": gatilho,
                "opcoes": list(dict.fromkeys(opcoes)),
                "quantidade": quantidade,
                "quantidade_editavel": bool(regra.get("quantidade_editavel", False)),
            }
        )
    return regras


def salvar_regras_popup_item(regras):
    if supabase_data.enabled():
        supabase_data.salvar_regras(regras or [])
        return
    salvar_json(OS_ITEM_POPUP_REGRAS_FILE, regras or [])


POPUP_ITEM_NAO_APLICAVEL = "__NAO_APLICAVEL__"


def _gatilhos_popup_na_bom(codigo_raiz, regras_por_gatilho, componentes):
    codigo_raiz = normalizar_codigo(codigo_raiz)
    encontrados = []
    gatilhos_vistos = set()
    pais_visitados = set()

    def visitar(codigo_pai, caminho):
        if codigo_pai in pais_visitados:
            return
        pais_visitados.add(codigo_pai)
        for componente in (componentes or {}).get(codigo_pai, []) or []:
            codigo_filho = normalizar_codigo((componente or {}).get("codigo", ""))
            if not codigo_filho or codigo_filho in caminho:
                continue
            caminho_filho = [*caminho, codigo_filho]
            if regras_por_gatilho.get(codigo_filho) and codigo_filho not in gatilhos_vistos:
                gatilhos_vistos.add(codigo_filho)
                encontrados.append({"codigo": codigo_filho, "caminho": caminho_filho})
            visitar(codigo_filho, caminho_filho)

    if codigo_raiz:
        visitar(codigo_raiz, [codigo_raiz])
    return encontrados


def _resolver_selecoes_popup_item(codigo_raiz, selecoes, regras_por_gatilho, componentes=None):
    codigo_raiz = normalizar_codigo(codigo_raiz)
    selecoes = [selecao for selecao in (selecoes or []) if isinstance(selecao, dict)]
    por_chave = {
        str(selecao.get("chave", "") or ""): selecao
        for selecao in selecoes
        if str(selecao.get("chave", "") or "")
    }
    resolvidas = []

    def resolver(gatilho, contexto, ancestrais, caminho, incluir_bom=True):
        for regra in regras_por_gatilho.get(gatilho, []):
            regra_id = str(regra.get("id", "") or "")
            chave = f"{contexto}|{regra_id}"
            selecao = por_chave.get(chave)
            if selecao is None and not ancestrais:
                selecao = next(
                    (
                        candidata
                        for candidata in selecoes
                        if not str(candidata.get("chave", "") or "")
                        and str(candidata.get("regra_id", "") or "") == regra_id
                    ),
                    None,
                )
            selecionado = normalizar_codigo((selecao or {}).get("codigo", ""))
            opcoes = {
                normalizar_codigo(codigo)
                for codigo in (regra.get("opcoes") or [])
                if normalizar_codigo(codigo)
            }
            quantidade = _parse_numero_form((selecao or {}).get("qtd", 0), 0)
            if selecao and selecionado == POPUP_ITEM_NAO_APLICAVEL:
                normalizada = dict(selecao)
                normalizada.update(
                    {
                        "regra_id": regra_id,
                        "gatilho": gatilho,
                        "codigo": POPUP_ITEM_NAO_APLICAVEL,
                        "qtd": 0,
                        "chave": chave,
                        "ancestrais": list(ancestrais),
                    }
                )
                resolvidas.append(normalizada)
                continue
            if not selecao or selecionado not in opcoes or quantidade <= 0:
                return f"Selecione o item relacionado obrigatorio para {gatilho}."
            if selecionado in caminho:
                trilha = " -> ".join([*caminho, selecionado])
                return f"Ciclo nos parametros de item relacionado: {trilha}."

            normalizada = dict(selecao)
            normalizada.update(
                {
                    "regra_id": regra_id,
                    "gatilho": gatilho,
                    "codigo": selecionado,
                    "qtd": quantidade,
                    "chave": chave,
                    "ancestrais": list(ancestrais),
                }
            )
            resolvidas.append(normalizada)
            erro = resolver(
                selecionado,
                f"{chave}>{selecionado}",
                [*ancestrais, chave],
                [*caminho, selecionado],
            )
            if erro:
                return erro
        if incluir_bom:
            for gatilho_bom in _gatilhos_popup_na_bom(gatilho, regras_por_gatilho, componentes):
                caminho_bom = gatilho_bom["caminho"]
                erro = resolver(
                    gatilho_bom["codigo"],
                    f"{contexto}>bom:{'>'.join(caminho_bom)}",
                    ancestrais,
                    [*caminho, *caminho_bom[1:]],
                    incluir_bom=False,
                )
                if erro:
                    return erro
        return ""

    erro = resolver(codigo_raiz, f"root:{codigo_raiz}", [], [codigo_raiz])
    return resolvidas, erro


def carregar_os_processos():
    if supabase_data.enabled():
        try:
            data = supabase_data.carregar_processos()
        except Exception:
            app.logger.exception("Falha ao carregar processos do Supabase")
            data = {}
        for conjunto in list(data):
            for nome in PROCESSOS_ORDEM:
                data[conjunto].setdefault(nome, [])
        return data

    if not os.path.exists(OS_PROCESSOS_FILE):
        return {}

    with open(OS_PROCESSOS_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)

    if data and isinstance(next(iter(data.values())), list):
        data = {"PADRAO": data}

    normalizado = {}
    for conjunto, processos in data.items():
        normalizado.setdefault(conjunto, {})
        for nome_proc, linhas in processos.items():
            chave = normalizar_nome_processo(nome_proc)
            normalizado[conjunto].setdefault(chave, [])
            normalizado[conjunto][chave].extend(linhas)
        for nome in PROCESSOS_ORDEM:
            normalizado[conjunto].setdefault(nome, [])

    return normalizado


def salvar_os_processos(processos):
    if supabase_data.enabled():
        supabase_data.salvar_processos(processos or {})
        return
    salvar_json(OS_PROCESSOS_FILE, processos or {})


def salvar_json(path, data):
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as exc:
        raise PermissionError(
            f"Falha ao salvar dados em '{path}'. Verifique permissao/OneDrive."
        ) from exc


OS_MODE_CONFIGS = {
    "completa": {
        "titulo": "O.S Completa",
        "doc_mode": "completa",
        "itens_source": "originais",
        "incluir_cliente_nome": False,
    },
    "expedicao": {
        "titulo": "REQUISI\u00c7\u00c3O EXPEDI\u00c7\u00c3O",
        "doc_mode": "expedicao",
        "itens_source": "expedicao",
        "incluir_cliente_nome": False,
    },
    "preparacao": {
        "titulo": "REQUISI\u00c7\u00c3O PREPARA\u00c7\u00c3O",
        "doc_mode": "preparacao",
        "itens_source": "preparacao",
        "incluir_cliente_nome": False,
    },
    "producao": {
        "titulo": "O.S Producao",
        "doc_mode": "producao",
        "itens_source": "originais",
        "incluir_cliente_nome": False,
    },
    "mascara": {
        "titulo": "Mascara",
        "doc_mode": "mascara",
        "itens_source": "originais",
    },
    "resumida": {
        "titulo": "O.S Resumida",
        "doc_mode": "resumida",
        "itens_source": "originais",
    },
}


def _formatar_qtd_saida(valor):
    try:
        numero = float(valor)
    except Exception:
        return valor
    if numero.is_integer():
        return int(numero)
    texto = f"{numero:.4f}".rstrip("0").rstrip(".")
    return texto.replace(".", ",")


def _parse_numero_form(valor, default=0.0):
    texto = str(valor or "").strip()
    if not texto:
        return default
    texto = texto.replace("R$", "").replace("%", "").strip()
    if "," in texto:
        texto = texto.replace(".", "").replace(",", ".")
    try:
        return float(texto)
    except Exception:
        return default


def _eh_faturamento_direto(descricao):
    texto = _corrigir_mojibake(descricao)
    texto = unicodedata.normalize("NFKD", str(texto or "").strip().upper())
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    return texto.startswith("FATURAMENTO DIRETO")


def _categoria_ar_condicionado(categoria):
    texto = _corrigir_mojibake(categoria)
    texto = unicodedata.normalize("NFKD", str(texto or "").strip().upper())
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    return texto.startswith("10") and "AR CONDICIONADO" in texto


def _fornecedor_faturamento_direto(descricao, categoria, fornecedor_informado="", fornecedor_cadastro=""):
    fornecedor_informado = str(fornecedor_informado or "").strip()
    fornecedor_cadastro = str(fornecedor_cadastro or "").strip()
    if _categoria_ar_condicionado(categoria):
        texto = _corrigir_mojibake(descricao)
        texto_norm = unicodedata.normalize("NFKD", str(texto or "").strip().upper())
        texto_norm = "".join(ch for ch in texto_norm if not unicodedata.combining(ch))
        partes = [parte for parte in re.split(r"[^A-Z0-9]+", texto_norm) if parte]
        if partes:
            ultimo = partes[-1]
            if ultimo in {"GE", "CLIM"}:
                return ultimo
            if ultimo in {"EURO", "GRUPOEURO"}:
                return "GE"
            if ultimo in {"CLIMAUTO", "CLIMATIZAR"}:
                return "CLIM"
    return fornecedor_informado or fornecedor_cadastro


def _resolve_output_path(path):
    dir_path = os.path.dirname(path)
    try:
        arquivos = os.listdir(dir_path)
    except Exception:
        arquivos = []
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    for idx in range(1, 100):
        candidato = f"{base} - R{idx:02d}{ext}"
        if not os.path.exists(candidato):
            return candidato
    return path


def _save_workbook_safe(wb, path):
    try:
        wb.save(path)
        return path
    except Exception:
        try:
            long_path = _win_long_path(path)
            wb.save(long_path)
            return path
        except Exception:
            fallback_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            os.makedirs(fallback_dir, exist_ok=True)
            fallback_path = _resolve_output_path(os.path.join(fallback_dir, os.path.basename(path)))
            wb.save(fallback_path)
            return fallback_path


def _agrupar_linhas_requisicao(linhas):
    agrupado = {}
    ordem = []
    for linha in linhas or []:
        codigo = normalizar_codigo(linha.get("codigo", ""))
        unidade = str(linha.get("unidade", "") or "").strip()
        setor = str(linha.get("setor", "") or "").strip()
        fornecedor = str(linha.get("fornecedor", "") or "").strip()
        tipo_requisicao = str(linha.get("tipo_requisicao", "") or "").strip()
        chave = (codigo, unidade, setor, fornecedor, tipo_requisicao)
        if chave not in agrupado:
            agrupado[chave] = {
                "codigo": codigo,
                "descricao": linha.get("descricao", "") or "",
                "unidade": unidade,
                "grupo": linha.get("grupo", "") or "",
                "categoria": linha.get("categoria", "") or "",
                "fornecedor": fornecedor,
                "setor": setor,
                "tipo_requisicao": tipo_requisicao,
                "qtd": 0.0,
            }
            ordem.append(chave)
        agrupado[chave]["qtd"] += parse_quantidade(linha.get("qtd", 0))
    return [agrupado[chave] for chave in ordem]


def _formatar_titulo_planilha(ws, titulo, total_colunas):
    ultima_coluna = max(total_colunas, 1)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ultima_coluna)
    cell = ws.cell(row=1, column=1)
    cell.value = titulo.upper()
    cell.font = Font(bold=True, size=14)
    cell.alignment = Alignment(horizontal="center", vertical="center")


def _formatar_cabecalho_planilha(ws, row_idx, total_colunas):
    for col_idx in range(1, total_colunas + 1):
        cell = ws.cell(row=row_idx, column=col_idx)
        if cell.value:
            cell.value = str(cell.value).upper()
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")


def _criar_planilha_requisicao_materiais(numero_os, dados, linhas, titulo_arquivo):
    pasta_destino = pasta_os(numero_os, dados)
    chassi = (dados.get("chassis", "") or "").strip()
    partes_nome = [titulo_arquivo]
    if chassi:
        partes_nome.append(chassi)
    nome_arquivo = " - ".join(partes_nome).strip(" -") + ".xlsx"
    caminho = _resolve_output_path(
        os.path.join(pasta_destino, nome_arquivo)
    )

    wb = Workbook()
    ws_detalhe = wb.active
    ws_detalhe.title = "Requisicao"
    _formatar_titulo_planilha(ws_detalhe, "REQUISI\u00c7\u00c3O DE MATERIAIS", 14)
    ws_detalhe.append(
        [
            "numero_os",
            "cliente",
            "chassis",
            "mmv",
            "item_os",
            "codigo",
            "descricao",
            "grupo",
            "categoria",
            "fornecedor",
            "tipo_requisicao",
            "setor",
            "unidade",
            "qtd",
        ]
    )
    _formatar_cabecalho_planilha(ws_detalhe, 2, 14)

    if not linhas:
        ws_detalhe.append(
            [
                numero_os,
                dados.get("cliente", ""),
                dados.get("chassis", ""),
                dados.get("mmv", ""),
                "",
                "",
                "Sem itens classificados para requisicao de materiais.",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ]
        )
    else:
        for linha in linhas:
            ws_detalhe.append(
                [
                    numero_os,
                    dados.get("cliente", ""),
                    dados.get("chassis", ""),
                    dados.get("mmv", ""),
                    linha.get("item", ""),
                    linha.get("codigo", ""),
                    linha.get("descricao", ""),
                    linha.get("grupo", ""),
                    linha.get("categoria", ""),
                    linha.get("fornecedor", ""),
                    linha.get("tipo_requisicao", ""),
                    linha.get("setor", ""),
                    linha.get("unidade", ""),
                    _formatar_qtd_saida(linha.get("qtd", "")),
                ]
            )

    ws_resumo = wb.create_sheet("Somatorio")
    _formatar_titulo_planilha(ws_resumo, "REQUISI\u00c7\u00c3O DE MATERIAIS", 9)
    ws_resumo.append(["codigo", "descricao", "grupo", "categoria", "fornecedor", "tipo_requisicao", "setor", "unidade", "qtd_total"])
    _formatar_cabecalho_planilha(ws_resumo, 2, 9)
    for linha in _agrupar_linhas_requisicao(linhas):
        ws_resumo.append(
            [
                linha.get("codigo", ""),
                linha.get("descricao", ""),
                linha.get("grupo", ""),
                linha.get("categoria", ""),
                linha.get("fornecedor", ""),
                linha.get("tipo_requisicao", ""),
                linha.get("setor", ""),
                linha.get("unidade", ""),
                _formatar_qtd_saida(linha.get("qtd", "")),
            ]
        )

    for ws in wb.worksheets:
        for col_idx, column_cells in enumerate(ws.columns, start=1):
            valores = [len(str(cell.value or "")) for cell in column_cells]
            letra = get_column_letter(col_idx)
            ws.column_dimensions[letra].width = min(max(max(valores, default=0) + 2, 12), 40)

    return _save_workbook_safe(wb, caminho)


def _criar_zip_temporario(paths, download_name):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
    tmp.close()
    with zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            if not path or not os.path.exists(path):
                continue
            zf.write(path, arcname=os.path.basename(path))
    return tmp.name, download_name


def carregar_importacao(path):
    if not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        try:
            return json.load(f)
        except json.JSONDecodeError:
            return {}


def limpar_importacao(path):
    if os.path.exists(path):
        os.remove(path)


def _limpar_placeholder(texto):
    if texto is None:
        return ""
    texto = str(texto).strip()
    if "{" in texto and "}" in texto:
        return ""
    return texto


def _parse_money(texto):
    if texto is None:
        return ""
    texto = str(texto)
    texto = texto.replace("R$", "").replace(" ", "").replace(".", "").replace(",", ".")
    try:
        return float(texto)
    except ValueError:
        return ""


def _normalizar_qtd(texto):
    texto = _limpar_placeholder(texto)
    if texto == "":
        return ""
    texto = str(texto).strip().replace(" ", "")
    if "," in texto and "." in texto:
        if texto.rfind(",") > texto.rfind("."):
            texto = texto.replace(".", "").replace(",", ".")
        else:
            texto = texto.replace(",", "")
    elif "," in texto:
        texto = texto.replace(",", ".")
    try:
        numero = float(texto)
    except ValueError:
        return ""
    if abs(numero - int(numero)) < 1e-9:
        return str(int(numero))
    return f"{numero:.6f}".rstrip("0").rstrip(".")


def _parse_date_ddmmyyyy(texto):
    texto = _limpar_placeholder(texto)
    if not texto:
        return ""
    try:
        dt = datetime.strptime(texto, "%d/%m/%Y")
        return dt.strftime("%Y-%m-%d")
    except ValueError:
        return ""


def _parse_datetime_ddmmyyyy(texto):
    texto = _limpar_placeholder(texto)
    if not texto:
        return ""
    texto = texto.replace(" - ", " ")
    try:
        dt = datetime.strptime(texto, "%d/%m/%Y %H:%M")
        return dt.strftime("%Y-%m-%dT%H:%M")
    except ValueError:
        return ""


def _best_descricao(cells, code_idx, desc_idx, ignore_idx=None):
    desc = _limpar_placeholder(cells[desc_idx]) if len(cells) > desc_idx else ""
    if desc:
        return desc
    parts = []
    for i, c in enumerate(cells):
        if i == code_idx:
            continue
        if ignore_idx and i in ignore_idx:
            continue
        txt = _limpar_placeholder(c)
        if not txt:
            continue
        if re.fullmatch(r"[\\d\\.,]+", txt) or "R$" in txt:
            continue
        parts.append(txt)
    return " ".join(parts).strip()


def _extract_pdf_text(file_storage):
    try:
        from PyPDF2 import PdfReader
    except Exception:
        return ""
    file_storage.stream.seek(0)
    reader = PdfReader(file_storage.stream)
    textos = []
    for page in reader.pages:
        textos.append(page.extract_text() or "")
    return "\n".join(textos)


def _buscar_valor_linha(linhas, label):
    label_upper = label.upper()
    for idx, linha in enumerate(linhas):
        if label_upper in linha.upper():
            partes = linha.split(":", 1)
            if len(partes) > 1 and partes[1].strip():
                return partes[1].strip()
            if idx + 1 < len(linhas):
                return linhas[idx + 1].strip()
    return ""


def parse_oc_docx(file_storage):
    file_storage.stream.seek(0)
    doc = Document(file_storage.stream)
    data = {}

    def cell_text_safe(table, row_idx, col_idx):
        if not table or row_idx < 0 or col_idx < 0:
            return ""
        if row_idx >= len(table.rows):
            return ""
        row = table.rows[row_idx]
        if col_idx >= len(row.cells):
            return ""
        return row.cells[col_idx].text

    if len(doc.tables) < 5:
        return data

    t0 = doc.tables[0]
    data["data_emissao"] = _parse_date_ddmmyyyy(t0.cell(1, 1).text if len(t0.rows) > 1 else "")
    oc_txt = t0.cell(1, 2).text if len(t0.rows) > 1 else ""
    oc_num = re.findall(r"\d+", oc_txt)
    data["oc_numero"] = oc_num[0] if oc_num else ""

    t2 = doc.tables[2]
    data["fornecedor"] = _limpar_placeholder(t2.cell(0, 1).text)
    data["razao_social"] = _limpar_placeholder(t2.cell(0, 4).text)
    data["cnpj"] = _limpar_placeholder(t2.cell(1, 1).text)
    data["endereco"] = _limpar_placeholder(t2.cell(1, 4).text)
    data["bairro"] = _limpar_placeholder(t2.cell(2, 1).text)
    data["cep"] = _limpar_placeholder(t2.cell(2, 4).text)
    data["cidade"] = _limpar_placeholder(t2.cell(3, 1).text)
    data["uf"] = _limpar_placeholder(t2.cell(3, 2).text)
    data["telefone"] = _limpar_placeholder(t2.cell(3, 4).text)
    data["email"] = _limpar_placeholder(t2.cell(4, 1).text)
    data["previsao"] = _parse_date_ddmmyyyy(t2.cell(4, 4).text)

    t3 = doc.tables[3]
    itens = []
    for row in t3.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):
            continue
        itens.append({
            "codigo": _limpar_placeholder(cells[0]) if len(cells) > 0 else "",
            "descricao": _best_descricao(cells, 0, 1, ignore_idx={2, 3, 4, 5, 6}),
            "unidade": _limpar_placeholder(cells[2]) if len(cells) > 2 else "",
            "qtd": _normalizar_qtd(cells[3]) if len(cells) > 3 else "",
            "valor": _parse_money(cells[4]) if len(cells) > 4 else "",
            "desconto": _parse_money(cells[5]) if len(cells) > 5 else "",
        })
    data["itens"] = itens

    t4 = doc.tables[4]
    if len(t4.rows) > 1:
        data["tipo_frete"] = _limpar_placeholder(cell_text_safe(t4, 1, 0))
        data["frete"] = _parse_money(cell_text_safe(t4, 1, 1))
        data["total_itens"] = _parse_money(cell_text_safe(t4, 1, 2))
        total_pedido = _parse_money(cell_text_safe(t4, 1, 3))
        if total_pedido == "":
            total_pedido = _parse_money(cell_text_safe(t4, 1, 4))
        data["total_pedido"] = total_pedido
        data["forma_pagamento"] = _limpar_placeholder(cell_text_safe(t4, 1, 6))
        data["prazo"] = _limpar_placeholder(cell_text_safe(t4, 1, 7))
        data["vencimento"] = _limpar_placeholder(cell_text_safe(t4, 1, 8))

    obs = []
    for row in t4.rows[2:]:
        for cell in row.cells:
            texto = cell.text.strip()
            if "Observa" in texto:
                if ":" in texto:
                    obs_txt = texto.split(":", 1)[1].strip()
                    if obs_txt:
                        obs.append(obs_txt)
    data["obs"] = "\n".join(dict.fromkeys(obs))
    return data


def parse_os_docx(file_storage):
    file_storage.stream.seek(0)
    doc = Document(file_storage.stream)
    data = {}

    if len(doc.tables) < 5:
        return data

    t0 = doc.tables[0]
    os_txt = t0.cell(0, 2).text if len(t0.rows) > 0 else ""
    os_num = re.findall(r"\d+", os_txt)
    data["os_numero"] = os_num[0] if os_num else ""

    t1 = doc.tables[1]
    data["chassis"] = _limpar_placeholder(t1.cell(1, 1).text)
    data["municipio"] = _limpar_placeholder(t1.cell(1, 3).text)
    data["cliente"] = _limpar_placeholder(t1.cell(2, 1).text)
    data["mmv"] = _limpar_placeholder(t1.cell(2, 3).text)
    data["previsao_inicio"] = _parse_datetime_ddmmyyyy(t1.cell(3, 1).text)
    data["previsao_termino"] = _parse_datetime_ddmmyyyy(t1.cell(3, 3).text)

    t2 = doc.tables[2]
    itens = []
    for row in t2.rows[2:]:
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):
            continue
        itens.append({
            "codigo": _limpar_placeholder(cells[0]) if len(cells) > 0 else "",
            "descricao": _best_descricao(cells, 0, 1, ignore_idx={2, 3, 4}),
            "qtd": _normalizar_qtd(cells[2]) if len(cells) > 2 else "",
            "serie": _limpar_placeholder(cells[3]) if len(cells) > 3 else "",
            "unidade": _limpar_placeholder(cells[4]) if len(cells) > 4 else "",
        })
    data["itens"] = itens

    t3 = doc.tables[3]
    composicao = []
    for row in t3.rows[2:]:
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):
            continue
        composicao.append({
            "codigo": _limpar_placeholder(cells[0]) if len(cells) > 0 else "",
            "descricao": _limpar_placeholder(cells[1]) if len(cells) > 1 else "",
            "qtd": _normalizar_qtd(cells[2]) if len(cells) > 2 else "",
            "unidade": _limpar_placeholder(cells[3]) if len(cells) > 3 else "",
        })
    data["composicao"] = composicao

    t4 = doc.tables[4]
    if len(t4.rows) > 1:
        data["obs_materiais"] = _limpar_placeholder(t4.cell(1, 0).text)

    processos_map = {
        "CORTE": 5,
        "AR CONDICIONADO": 6,
        "PREPARAÃ‡ÃƒO DE PEÃ‡AS": 7,
        "ISOLAMENTO": 8,
        "REVESTIMENTO": 9,
        "BANCOS": 10,
        "ELÃ‰TRICA 2": 11,
        "LIMPEZA/LIBERAÃ‡ÃƒO": 12,
    }
    processos = {}
    for nome, idx in processos_map.items():
        if len(doc.tables) <= idx:
            continue
        t = doc.tables[idx]
        linhas = []
        for row in t.rows[2:]:
            atividade = row.cells[1].text.strip() if len(row.cells) > 1 else ""
            if not atividade:
                continue
            responsavel_idx = 2 if len(row.cells) > 6 else 3
            data_idx = 3 if len(row.cells) > 6 else None
            inicio_idx = 4 if len(row.cells) > 6 else None
            fim_idx = 5 if len(row.cells) > 6 else None
            feito_idx = 6 if len(row.cells) > 6 else 4 if len(row.cells) > 4 else None
            linha = {
                "atividade": atividade,
                "responsavel": row.cells[responsavel_idx].text.strip() if len(row.cells) > responsavel_idx else "",
            }
            if data_idx is not None:
                linha["data"] = row.cells[data_idx].text.strip() if len(row.cells) > data_idx else ""
                linha["inicio"] = row.cells[inicio_idx].text.strip() if len(row.cells) > inicio_idx else ""
                linha["fim"] = row.cells[fim_idx].text.strip() if len(row.cells) > fim_idx else ""
                linha["feito"] = row.cells[feito_idx].text.strip() if len(row.cells) > feito_idx else ""
            linhas.append(linha)
        processos[nome] = linhas
    data["processos"] = processos

    obs_final = ""
    for p in doc.paragraphs:
        if p.text.strip().upper().startswith("OBS FINAL"):
            obs_final = p.text.split(":", 1)[-1].strip()
            break
    data["obs"] = obs_final
    return data


def parse_os_docx_atualizado(file_storage):
    file_storage.stream.seek(0)
    doc = Document(file_storage.stream)
    refs = mapear_tabelas_os(doc)
    data = {}

    def cell_text_safe(table, row_idx, col_idx):
        if table is None or row_idx < 0 or col_idx < 0:
            return ""
        if row_idx >= len(table.rows):
            return ""
        row = table.rows[row_idx]
        if col_idx >= len(row.cells):
            return ""
        return (row.cells[col_idx].text or "").strip()

    if refs.get("cabecalho") is not None:
        t0 = doc.tables[refs["cabecalho"]]
        os_txt = cell_text_safe(t0, 0, 2)
        os_num = re.findall(r"\d+", os_txt)
        data["os_numero"] = os_num[0] if os_num else ""

    if refs.get("dados") is not None:
        t1 = doc.tables[refs["dados"]]
        data["chassis"] = _limpar_placeholder(cell_text_safe(t1, 1, 1))
        data["municipio"] = _limpar_placeholder(cell_text_safe(t1, 1, 3))
        data["cliente"] = _limpar_placeholder(cell_text_safe(t1, 2, 1))
        data["mmv"] = _limpar_placeholder(cell_text_safe(t1, 2, 3))
        data["previsao_inicio"] = _parse_datetime_ddmmyyyy(cell_text_safe(t1, 3, 1))
        data["previsao_termino"] = _parse_datetime_ddmmyyyy(cell_text_safe(t1, 3, 3))

    itens = []
    if refs.get("itens") is not None:
        tabela_itens = doc.tables[refs["itens"]]
        header_idx = encontrar_linha_cabecalho(tabela_itens, "CODIGO", "QTD")
        inicio = (header_idx or 0) + 1
        for row in tabela_itens.rows[inicio:]:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue
            itens.append({
                "codigo": _limpar_placeholder(cells[0]) if len(cells) > 0 else "",
                "descricao": _best_descricao(cells, 0, 1, ignore_idx={2, 3, 4}),
                "qtd": _normalizar_qtd(cells[2]) if len(cells) > 2 else "",
                "serie": _limpar_placeholder(cells[3]) if len(cells) > 3 else "",
                "unidade": _limpar_placeholder(cells[4]) if len(cells) > 4 else "",
            })
    data["itens"] = itens

    composicao = []
    if refs.get("composicao") is not None:
        tabela_comp = doc.tables[refs["composicao"]]
        header_idx = encontrar_linha_cabecalho(tabela_comp, "CODIGO", "QTD")
        inicio = (header_idx or 0) + 1
        for row in tabela_comp.rows[inicio:]:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue
            composicao.append({
                "item": "",
                "codigo": _limpar_placeholder(cells[0]) if len(cells) > 0 else "",
                "descricao": _limpar_placeholder(cells[1]) if len(cells) > 1 else "",
                "qtd": _normalizar_qtd(cells[2]) if len(cells) > 2 else "",
                "unidade": _limpar_placeholder(cells[3]) if len(cells) > 3 else "",
                "level": 0,
            })
    data["composicao"] = composicao

    if refs.get("observacoes") is not None:
        tabela_obs = doc.tables[refs["observacoes"]]
        if len(tabela_obs.rows) > 1:
            data["obs_materiais"] = _limpar_placeholder(cell_text_safe(tabela_obs, 1, 0))

    processos = {nome: [] for nome in PROCESSOS_ORDEM}
    for nome, idx in refs.get("processos", {}).items():
        tabela_proc = doc.tables[idx]
        header_idx = encontrar_linha_cabecalho(tabela_proc, "ATIVIDADE", "RESPONS")
        inicio = (header_idx or 0) + 1
        linhas = []
        for row in tabela_proc.rows[inicio:]:
            cells = [c.text.strip() for c in row.cells]
            atividade = cells[1] if len(cells) > 1 else ""
            if not atividade:
                continue
            responsavel_idx = 2 if len(cells) > 6 else 3
            data_idx = 3 if len(cells) > 6 else None
            inicio_idx = 4 if len(cells) > 6 else None
            fim_idx = 5 if len(cells) > 6 else None
            feito_idx = 6 if len(cells) > 6 else 4 if len(cells) > 4 else None
            linha = {
                "atividade": atividade,
                "responsavel": cells[responsavel_idx] if len(cells) > responsavel_idx else "",
            }
            if data_idx is not None:
                linha["data"] = cells[data_idx] if len(cells) > data_idx else ""
                linha["inicio"] = cells[inicio_idx] if len(cells) > inicio_idx else ""
                linha["fim"] = cells[fim_idx] if len(cells) > fim_idx else ""
                linha["feito"] = cells[feito_idx] if len(cells) > feito_idx else ""
            linhas.append(linha)
        processos[nome] = linhas
    data["processos"] = processos

    obs_final = ""
    for p in doc.paragraphs:
        if p.text.strip().upper().startswith("OBS FINAL"):
            obs_final = p.text.split(":", 1)[-1].strip()
            break
    data["obs"] = obs_final
    return data


def parse_oc_pdf(file_storage):
    texto = _extract_pdf_text(file_storage)
    linhas = [l.strip() for l in texto.splitlines() if l.strip()]
    data = {
        "fornecedor": _buscar_valor_linha(linhas, "NOME FANTASIA"),
        "razao_social": _buscar_valor_linha(linhas, "RAZÃO SOCIAL") or _buscar_valor_linha(linhas, "RAZAO SOCIAL"),
        "cnpj": _buscar_valor_linha(linhas, "CNPJ/CPF"),
        "endereco": _buscar_valor_linha(linhas, "ENDEREÇO") or _buscar_valor_linha(linhas, "ENDERECO"),
        "bairro": _buscar_valor_linha(linhas, "BAIRRO"),
        "cep": _buscar_valor_linha(linhas, "CEP"),
        "cidade": _buscar_valor_linha(linhas, "CIDADE"),
        "uf": _buscar_valor_linha(linhas, "UF"),
        "telefone": _buscar_valor_linha(linhas, "TELEFONE"),
        "email": _buscar_valor_linha(linhas, "E-MAIL"),
        "previsao": _parse_date_ddmmyyyy(_buscar_valor_linha(linhas, "PREVISÃO DE CHEGADA") or _buscar_valor_linha(linhas, "PREVISAO DE CHEGADA")),
        "tipo_frete": _buscar_valor_linha(linhas, "TIPO DO FRETE"),
        "frete": _parse_money(_buscar_valor_linha(linhas, "VALOR DO FRETE")),
        "total_itens": _parse_money(_buscar_valor_linha(linhas, "VALOR DO PEDIDO")),
        "total_pedido": _parse_money(_buscar_valor_linha(linhas, "VALOR TOTAL DO PEDIDO")),
        "forma_pagamento": _buscar_valor_linha(linhas, "FORMA DE PAGAMENTO"),
        "prazo": _buscar_valor_linha(linhas, "PRAZO"),
        "vencimento": _buscar_valor_linha(linhas, "VENCIMENTO"),
        "obs": "",
        "itens": [],
    }

    for linha in linhas:
        if "OBSERVA" in linha.upper() and ":" in linha:
            data["obs"] = linha.split(":", 1)[1].strip()
            break

    itens = []
    for linha in linhas:
        if "R$" not in linha:
            continue
        money = re.findall(r"R\\$\\s*[\\d\\.,]+", linha)
        if len(money) < 2:
            continue
        partes = linha.split()
        if len(partes) < 4:
            continue
        codigo = partes[0]
        pre_money = linha.split("R$")[0].strip()
        numeros = re.findall(r"[\\d\\.,]+", pre_money)
        qtd = _normalizar_qtd(numeros[-1]) if numeros else ""
        desc = pre_money.replace(codigo, "", 1).replace(qtd, "", 1).strip()
        itens.append({
            "codigo": codigo,
            "descricao": desc,
            "unidade": "",
            "qtd": qtd,
            "valor": _parse_money(money[-3]) if len(money) >= 3 else _parse_money(money[-2]),
            "desconto": _parse_money(money[-2]) if len(money) >= 3 else "",
        })
    data["itens"] = itens
    return data


def parse_os_pdf(file_storage):
    texto = _extract_pdf_text(file_storage)
    linhas = [l.strip() for l in texto.splitlines() if l.strip()]
    data = {
        "os_numero": "",
        "chassis": _buscar_valor_linha(linhas, "CHASSIS"),
        "municipio": _buscar_valor_linha(linhas, "MUNICÍPIO") or _buscar_valor_linha(linhas, "MUNICIPIO"),
        "cliente": _buscar_valor_linha(linhas, "CLIENTE"),
        "mmv": _buscar_valor_linha(linhas, "MMV"),
        "previsao_inicio": _parse_datetime_ddmmyyyy(_buscar_valor_linha(linhas, "PREVISÃO INICIO") or _buscar_valor_linha(linhas, "PREVISAO INICIO")),
        "previsao_termino": _parse_datetime_ddmmyyyy(_buscar_valor_linha(linhas, "PREVISÃO TÉRMINO") or _buscar_valor_linha(linhas, "PREVISAO TERMINO")),
        "obs_materiais": _buscar_valor_linha(linhas, "OBSERVAÇÕES") or _buscar_valor_linha(linhas, "OBSERVACOES"),
        "obs": "",
        "itens": [],
        "processos": {},
    }

    for linha in linhas:
        if "OBS FINAL" in linha.upper() and ":" in linha:
            data["obs"] = linha.split(":", 1)[1].strip()
            break

    itens = []
    capturar = False
    for linha in linhas:
        if "PRODUTOS:" in linha.upper():
            capturar = True
            continue
        if "COMPOSIÇÃO" in linha.upper() or "COMPOSICAO" in linha.upper():
            capturar = False
        if not capturar:
            continue
        partes = linha.split()
        if len(partes) < 3:
            continue
        codigo = partes[0]
        qtd = _normalizar_qtd(partes[-2]) if len(partes) >= 2 else ""
        un = partes[-1] if len(partes) >= 1 else ""
        descricao = " ".join(partes[1:-2]).strip()
        if codigo and descricao:
            itens.append({
                "codigo": codigo,
                "descricao": descricao,
                "qtd": qtd,
                "serie": "",
                "unidade": un,
            })
    data["itens"] = itens

    composicao = []
    capturar = False
    for linha in linhas:
        up = linha.upper()
        if "COMPOSIÇÃO" in up or "COMPOSICAO" in up:
            capturar = True
            continue
        if "OBSERVA" in up or "PROCESSOS DE PRODU" in up or "PROCESSOS DE PRODU" in up:
            capturar = False
        if not capturar:
            continue
        partes = linha.split()
        if len(partes) < 3:
            continue
        codigo = partes[0]
        qtd = _normalizar_qtd(partes[-2]) if len(partes) >= 2 else ""
        un = partes[-1] if len(partes) >= 1 else ""
        descricao = " ".join(partes[1:-2]).strip()
        if codigo and descricao:
            composicao.append({
                "codigo": codigo,
                "descricao": descricao,
                "qtd": qtd,
                "unidade": un,
            })
    data["composicao"] = composicao

    processos = {}
    current = None
    for linha in linhas:
        up = linha.upper()
        if "PROCESSOS DE PRODUÇÃO" in up or "PROCESSOS DE PRODUCAO" in up:
            if "CORTE" in up:
                current = "CORTE"
            elif "AR CONDICIONADO" in up:
                current = "AR CONDICIONADO"
            elif "PREPARA" in up:
                current = "PREPARAÃ‡ÃƒO DE PEÃ‡AS"
            elif "ISOLAMENTO" in up:
                current = "ISOLAMENTO"
            elif "REVESTIMENTO" in up:
                current = "REVESTIMENTO"
            elif "BANCOS" in up:
                current = "BANCOS"
            elif "ELÉTRICA" in up or "ELETRICA" in up:
                current = "ELÃ‰TRICA 2"
            elif "LIMPEZA" in up:
                current = "LIMPEZA/LIBERAÃ‡ÃƒO"
            processos.setdefault(current, [])
            continue
        if current:
            if up.startswith("#") or up.startswith("ATIVIDADE") or up.startswith("RESPONS"):
                continue
            texto = re.sub(r"^\\d+\\s+", "", linha).strip()
            if texto:
                processos[current].append({"atividade": texto, "responsavel": ""})
    data["processos"] = processos
    return data



def _criar_modelo_xlsx(headers, nome_arquivo, header_row=1):
    wb = Workbook()
    ws = wb.active
    for _ in range(max(header_row - 1, 0)):
        ws.append(["" for _ in headers])
    ws.append(headers)
    _formatar_cabecalho_planilha(ws, header_row, len(headers))

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    temp.close()
    wb.save(temp.name)
    return temp.name, nome_arquivo


def _criar_modelo_os_processos_xlsx():
    wb = Workbook()
    ws = wb.active
    headers = [
        "OPERACAO", "10 - CORTE",
        "OPERACAO", "20 - AR CONDICIONADO",
        "OPERACAO", "30 - PREPARACAO DE PECAS",
        "OPERACAO", "40 - ELETRICA 1",
        "OPERACAO", "50 - ISOLAMENTO",
        "OPERACAO", "60 - REVESTIMENTO 1",
        "OPERACAO", "70 - REVESTIMENTO 2",
        "OPERACAO", "80 - ELETRICA 2",
        "OPERACAO", "90 - LIMPEZA/LIBERACAO",
        "OPERACAO", "95 - ACESSORIOS",
        "OPERACAO", "100 - BANCO",
    ]
    ws.append(headers)

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    temp.close()
    wb.save(temp.name)
    return temp.name, "modelo_os_processos.xlsx"


def _catalogo_itens_regras():
    catalogo = dict(carregar_produtos() or {})
    catalogo.update(carregar_os_produtos() or {})
    return catalogo


def _descricao_item_regras(codigo, catalogo):
    info = (catalogo or {}).get(normalizar_codigo(codigo), {}) or {}
    return str(info.get("descricao", "") or "").strip()


def _formatar_planilha_regras(ws, headers):
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}1"
    ws.sheet_view.showGridLines = False
    fill = PatternFill("solid", fgColor="1F4E78")
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 24


def _adicionar_instrucoes_planilha(wb, titulo, instrucoes):
    ws = wb.create_sheet("Instrucoes")
    ws.sheet_view.showGridLines = False
    ws["A1"] = titulo
    ws["A1"].font = Font(bold=True, size=14, color="1F4E78")
    for idx, instrucao in enumerate(instrucoes, 3):
        ws.cell(row=idx, column=1, value=f"{idx - 2}.")
        ws.cell(row=idx, column=2, value=instrucao)
        ws.cell(row=idx, column=2).alignment = Alignment(wrap_text=True, vertical="top")
    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 110


def _salvar_planilha_temporaria(wb, nome_arquivo):
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    temp.close()
    wb.save(temp.name)
    return temp.name, nome_arquivo


def _criar_planilha_regras_popup_item(regras=None, nome_arquivo="modelo_parametros_item_relacionado.xlsx"):
    wb = Workbook()
    ws = wb.active
    ws.title = "Parametros Item Relacionado"
    _formatar_planilha_regras(ws, MODELO_REGRAS_POPUP_HEADERS)
    catalogo = _catalogo_itens_regras()

    for regra in regras or []:
        gatilho = normalizar_codigo(regra.get("gatilho", ""))
        opcoes = [normalizar_codigo(codigo) for codigo in regra.get("opcoes", []) if normalizar_codigo(codigo)]
        ws.append(
            [
                str(regra.get("id", "") or ""),
                gatilho,
                _descricao_item_regras(gatilho, catalogo),
                "; ".join(opcoes),
                " | ".join(_descricao_item_regras(codigo, catalogo) for codigo in opcoes),
                regra.get("quantidade", 1),
                "SIM" if regra.get("quantidade_editavel") else "NAO",
            ]
        )

    widths = [16, 18, 54, 42, 72, 14, 24]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    for row_idx in range(2, ws.max_row + 1):
        ws.cell(row=row_idx, column=2).number_format = "@"
        ws.cell(row=row_idx, column=4).number_format = "@"
    ws.auto_filter.ref = f"A1:G{max(ws.max_row, 1)}"
    validacao = DataValidation(type="list", formula1='"SIM,NAO"', allow_blank=False)
    ws.add_data_validation(validacao)
    validacao.add("G2:G1000")

    _adicionar_instrucoes_planilha(
        wb,
        "Parametros de Item Relacionado",
        [
            "Use uma linha por regra. Separe varios codigos em ITENS_OPCOES com ponto e virgula (;).",
            "ID_REGRA e opcional para regras novas. No arquivo exportado, mantenha o ID para atualizar a regra existente.",
            "QUANTIDADE deve ser maior que zero. QUANTIDADE_EDITAVEL aceita SIM ou NAO.",
            "A importacao mescla os dados: atualiza IDs existentes e adiciona novas regras sem apagar as demais.",
            "As colunas de descricao servem apenas para conferencia e nao sao usadas na importacao.",
        ],
    )
    return _salvar_planilha_temporaria(wb, nome_arquivo)


def _criar_planilha_relacoes_processo_item(relacoes=None, nome_arquivo="modelo_relacao_processo_item.xlsx"):
    wb = Workbook()
    ws = wb.active
    ws.title = "Relacao Processo x Item"
    _formatar_planilha_regras(ws, MODELO_RELACOES_PROCESSO_HEADERS)
    catalogo = _catalogo_itens_regras()

    for codigo, processos in (relacoes or {}).items():
        codigo = normalizar_codigo(codigo)
        ws.append(
            [
                codigo,
                _descricao_item_regras(codigo, catalogo),
                "; ".join(str(processo or "").strip() for processo in processos if str(processo or "").strip()),
            ]
        )

    widths = [18, 58, 96]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    for row_idx in range(2, ws.max_row + 1):
        ws.cell(row=row_idx, column=1).number_format = "@"
    ws.auto_filter.ref = f"A1:C{max(ws.max_row, 1)}"

    _adicionar_instrucoes_planilha(
        wb,
        "Relacao Processo x Item",
        [
            "Use uma linha por item. Separe varios processos na coluna PROCESSOS com ponto e virgula (;).",
            "O nome do processo deve ser identico ao nome do arquivo/conjunto importado na base de processos.",
            "A importacao mescla os dados: substitui os processos dos itens informados e preserva os demais itens.",
            "A coluna DESCRICAO_ITEM serve apenas para conferencia e nao e usada na importacao.",
        ],
    )
    return _salvar_planilha_temporaria(wb, nome_arquivo)


def normalizar_header(texto):
    if texto is None:
        return ""
    texto = str(texto).strip().lower()
    return (
        texto.replace(" ", "_")
        .replace("-", "_")
        .replace("/", "_")
        .replace("ç", "c")
        .replace("ã", "a")
        .replace("á", "a")
        .replace("à", "a")
        .replace("â", "a")
        .replace("é", "e")
        .replace("ê", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ô", "o")
        .replace("õ", "o")
        .replace("ú", "u")
    )


def ler_linhas_arquivo(file_storage):
    filename = secure_filename(file_storage.filename or "")
    ext = os.path.splitext(filename)[1].lower()

    if ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
        wb = load_workbook(file_storage, data_only=True)
        ws = wb.active
        linhas = []
        for row in ws.iter_rows(values_only=True):
            linha = ["" if cell is None else str(cell).strip() for cell in row]
            if any(linha):
                linhas.append(linha)
        return linhas

    return []


BAIXA_HEADERS = {
    "documento_id": {
        "id",
        "documento_id",
        "id_documento",
        "historico_id",
        "id_historico",
        "doc_id",
    },
    "line_id": {
        "id_linha",
        "linha_id",
        "line_id",
        "id_item",
        "item_id",
        "id_da_linha",
        "codigo_linha",
    },
    "acao": {
        "acao",
        "ação",
        "baixa",
        "status",
        "situacao",
        "situação",
        "operacao",
        "operação",
    },
    "tipo": {"tipo", "documento", "tipo_documento"},
    "numero": {"numero", "número", "oc", "os", "o_c", "o_s"},
}


def _canonicalizar_header_baixa(texto):
    normalizado = normalizar_header(texto)
    if not normalizado:
        return ""
    for canonico, aliases in BAIXA_HEADERS.items():
        aliases_norm = {normalizar_header(alias) for alias in aliases}
        if normalizado == canonico or normalizado in aliases_norm:
            return canonico
    return normalizado


def _resolver_header_baixa(linhas):
    melhor_idx = None
    melhor_score = 0
    for idx in range(min(len(linhas), 8)):
        headers = [_canonicalizar_header_baixa(valor) for valor in linhas[idx]]
        score = sum(1 for header in headers if header in {"documento_id", "line_id", "acao", "tipo", "numero"})
        if "acao" in headers:
            score += 3
        if score > melhor_score:
            melhor_idx = idx
            melhor_score = score
    if melhor_idx is None:
        return {}, []
    header = [_canonicalizar_header_baixa(h) for h in linhas[melhor_idx]]
    mapa = {}
    for idx, nome in enumerate(header):
        if nome and nome not in mapa:
            mapa[nome] = idx
    return mapa, linhas[melhor_idx + 1:]


def _linhas_worksheet(ws):
    linhas = []
    for row in ws.iter_rows(values_only=True):
        linha = ["" if cell is None else str(cell).strip() for cell in row]
        if any(linha):
            linhas.append(linha)
    return linhas


def _ler_abas_baixa(file_storage):
    filename = secure_filename(file_storage.filename or "")
    ext = os.path.splitext(filename)[1].lower()
    try:
        file_storage.stream.seek(0)
    except Exception:
        pass

    if ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
        wb = load_workbook(file_storage, data_only=True, read_only=True)
        try:
            return [(ws.title, _linhas_worksheet(ws)) for ws in wb.worksheets]
        finally:
            wb.close()

    if ext == ".csv":
        raw = file_storage.stream.read()
        texto = raw.decode("utf-8-sig", errors="replace")
        amostra = texto[:2048]
        try:
            dialect = csv.Sniffer().sniff(amostra, delimiters=";,\t")
            reader = csv.reader(io.StringIO(texto), dialect)
        except Exception:
            reader = csv.reader(io.StringIO(texto), delimiter=";")
        linhas = [[cell.strip() for cell in row] for row in reader]
        return [(filename or "CSV", [linha for linha in linhas if any(linha)])]

    raise ValueError("Envie uma planilha XLSX ou CSV.")


def _normalizar_acao_baixa(value):
    texto = normalizar_header(value)
    if texto in {"concluir", "concluido", "concluida", "baixar", "baixa", "finalizar", "finalizado", "finalizada"}:
        return "concluido"
    if texto in {"cancelar", "cancelado", "cancelada"}:
        return "cancelado"
    if texto in {"excluir", "excluido", "excluida", "deletar", "apagar", "remover"}:
        return "excluir"
    return ""


def _normalizar_tipo_documento(value):
    texto = normalizar_header(value)
    if texto in {"oc", "o_c", "compra", "compras", "pedido", "pedido_compra", "ordem_compra"}:
        return "oc"
    if texto in {"os", "o_s", "ordem_servico", "ordens_servico", "servico", "producao"}:
        return "os"
    return ""


def _line_ids_documento(documento):
    for item in documento.get("itens", []) or []:
        line_id = _normalizar_line_id((item or {}).get("line_id"))
        if line_id:
            yield line_id
    for comp in documento.get("composicao", []) or []:
        line_id = _normalizar_line_id((comp or {}).get("line_id"))
        if line_id:
            yield line_id
    for linhas in (documento.get("processos", {}) or {}).values():
        for linha in linhas or []:
            line_id = _normalizar_line_id((linha or {}).get("line_id"))
            if line_id:
                yield line_id


def _linha_status(linha):
    return str((linha or {}).get("line_status") or "emitido").strip().lower() or "emitido"


def _linha_tem_id(linha, line_id):
    return _normalizar_line_id((linha or {}).get("line_id")) == line_id


def _aplicar_acao_linha_lista(linhas, line_id, acao, usuario):
    resultado = []
    aplicado = False
    for linha in linhas or []:
        if not isinstance(linha, dict) or not _linha_tem_id(linha, line_id):
            resultado.append(linha)
            continue
        aplicado = True
        if acao == "excluir":
            continue
        atualizada = dict(linha)
        atualizada["line_status"] = acao
        atualizada["line_status_atualizado_por"] = usuario
        atualizada["line_status_atualizado_em"] = datetime.now().isoformat(timespec="seconds")
        resultado.append(atualizada)
    return resultado, aplicado


def aplicar_baixa_linha_documento(documento, line_id, acao):
    line_id = _normalizar_line_id(line_id)
    if not line_id:
        return False, ""
    usuario = current_username()

    itens, aplicado = _aplicar_acao_linha_lista(documento.get("itens", []) or [], line_id, acao, usuario)
    if aplicado:
        documento["itens"] = itens
        return True, "itens"

    composicao, aplicado = _aplicar_acao_linha_lista(documento.get("composicao", []) or [], line_id, acao, usuario)
    if aplicado:
        documento["composicao"] = composicao
        return True, "composicao"

    processos = documento.get("processos", {}) or {}
    for grupo, linhas in list(processos.items()):
        linhas_atualizadas, aplicado = _aplicar_acao_linha_lista(linhas or [], line_id, acao, usuario)
        if aplicado:
            processos[grupo] = linhas_atualizadas
            documento["processos"] = processos
            return True, f"processos/{grupo}"

    return False, ""


def _indexar_historico_para_baixa(historico):
    por_id = {}
    por_linha = {}
    por_tipo_numero = {}
    for documento in historico or []:
        documento_id = str(documento.get("id") or "").strip()
        tipo = str(documento.get("tipo") or "").strip().lower()
        numero = str(documento.get("numero") or "").strip()
        if documento_id:
            por_id[documento_id] = documento
        if tipo and numero:
            por_tipo_numero[(tipo, numero)] = documento
        for line_id in _line_ids_documento(documento):
            por_linha[line_id] = documento
    return por_id, por_linha, por_tipo_numero


def _valor_baixa(row, mapa, campo):
    idx = mapa.get(campo)
    if idx is None or idx >= len(row):
        return ""
    return _corrigir_mojibake(row[idx])


def _resolver_documento_baixa(row, mapa, indices):
    por_id, por_linha, por_tipo_numero = indices
    documento_id = _valor_baixa(row, mapa, "documento_id")
    if documento_id and documento_id in por_id:
        return por_id[documento_id], "ID"
    line_id = _valor_baixa(row, mapa, "line_id")
    if line_id and line_id in por_linha:
        return por_linha[line_id], "ID Linha"
    tipo = _normalizar_tipo_documento(_valor_baixa(row, mapa, "tipo"))
    numero = _valor_baixa(row, mapa, "numero")
    if tipo and numero and (tipo, numero) in por_tipo_numero:
        return por_tipo_numero[(tipo, numero)], "Tipo/Numero"
    return None, ""


def importar_baixas_documentos(file_storage, tipo_filtro=""):
    tipo_filtro = str(tipo_filtro or "").strip().lower()
    if tipo_filtro not in {"", "oc", "os"}:
        raise ValueError("Tipo de baixa invalido.")

    historico = carregar_historico()
    indices = _indexar_historico_para_baixa(historico)
    por_id, por_linha, _ = indices
    pendentes_documentos = {}
    pendentes_linhas = {}
    resultado = {
        "linhas": 0,
        "atualizados": 0,
        "excluidos": 0,
        "linhas_atualizadas": 0,
        "linhas_excluidas": 0,
        "ignorados": 0,
        "erros": [],
    }

    for aba, linhas in _ler_abas_baixa(file_storage):
        if not linhas:
            continue
        mapa, rows = _resolver_header_baixa(linhas)
        if "acao" not in mapa:
            continue
        for row_idx, row in enumerate(rows, start=2):
            acao = _normalizar_acao_baixa(_valor_baixa(row, mapa, "acao"))
            if not acao:
                resultado["ignorados"] += 1
                continue
            resultado["linhas"] += 1

            line_id = _normalizar_line_id(_valor_baixa(row, mapa, "line_id"))
            if line_id:
                documento = por_linha.get(line_id)
                if not documento:
                    resultado["erros"].append(f"{aba} linha {row_idx}: ID Linha nao encontrado")
                    continue
                if tipo_filtro and documento.get("tipo") != tipo_filtro:
                    resultado["ignorados"] += 1
                    continue
                documento_id = str(documento.get("id") or "").strip()
                if not documento_id:
                    resultado["erros"].append(f"{aba} linha {row_idx}: documento da linha sem ID")
                    continue
                if documento_id in pendentes_documentos:
                    resultado["erros"].append(
                        f"{aba} linha {row_idx}: documento {documento.get('numero') or documento_id} ja possui baixa de documento inteiro"
                    )
                    continue
                key = (documento_id, line_id)
                existente = pendentes_linhas.get(key)
                if existente and existente["acao"] != acao:
                    resultado["erros"].append(f"{aba} linha {row_idx}: acao conflitante para ID Linha {line_id}")
                    continue
                pendentes_linhas[key] = {"acao": acao, "documento": documento, "aba": aba}
                continue

            documento, origem = _resolver_documento_baixa(row, mapa, indices)
            if not documento:
                resultado["erros"].append(f"{aba} linha {row_idx}: documento nao encontrado")
                continue
            if tipo_filtro and documento.get("tipo") != tipo_filtro:
                resultado["ignorados"] += 1
                continue
            documento_id = str(documento.get("id") or "").strip()
            if not documento_id:
                resultado["erros"].append(f"{aba} linha {row_idx}: documento sem ID")
                continue
            if any(key[0] == documento_id for key in pendentes_linhas):
                resultado["erros"].append(
                    f"{aba} linha {row_idx}: documento {documento.get('numero') or documento_id} ja possui baixa por linha"
                )
                continue
            existente = pendentes_documentos.get(documento_id)
            if existente and existente["acao"] != acao:
                resultado["erros"].append(
                    f"{aba} linha {row_idx}: acao conflitante para documento {documento.get('numero') or documento_id}"
                )
                continue
            pendentes_documentos[documento_id] = {"acao": acao, "documento": documento, "origem": origem}

    documentos_linhas = {}
    for (documento_id, line_id), item in pendentes_linhas.items():
        documento = documentos_linhas.get(documento_id)
        if documento is None:
            base_doc = obter_historico_documento(documento_id) or item["documento"]
            documento = dict(base_doc)
            documentos_linhas[documento_id] = documento
        aplicado, _escopo = aplicar_baixa_linha_documento(documento, line_id, item["acao"])
        if not aplicado:
            resultado["erros"].append(f"ID Linha {line_id}: nao foi possivel aplicar baixa")
            continue
        if item["acao"] == "excluir":
            resultado["linhas_excluidas"] += 1
        else:
            resultado["linhas_atualizadas"] += 1

    for documento_id, documento in documentos_linhas.items():
        salvar_historico_documento_atualizado(documento_id, documento)

    for documento_id, item in pendentes_documentos.items():
        acao = item["acao"]
        if acao == "excluir":
            excluir_historico_documento(documento_id)
            resultado["excluidos"] += 1
        else:
            atualizado = atualizar_status_historico_documento(documento_id, acao)
            if atualizado:
                resultado["atualizados"] += 1
            else:
                resultado["erros"].append(f"documento {documento_id}: nao foi possivel atualizar")

    return resultado


def _corrigir_mojibake(texto):
    texto = "" if texto is None else str(texto).strip()
    if not texto:
        return ""
    if any(token in texto for token in ("Ã", "Â", "�")):
        for encoding in ("latin1", "cp1252"):
            try:
                convertido = texto.encode(encoding).decode("utf-8")
            except Exception:
                continue
            if convertido:
                texto = convertido
                break
    return texto


def normalizar_header(texto):
    if texto is None:
        return ""
    texto = _corrigir_mojibake(texto)
    texto = unicodedata.normalize("NFKD", str(texto).strip().lower())
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    texto = re.sub(r"[^a-z0-9]+", "_", texto)
    return texto.strip("_")


def _canonicalizar_header(texto):
    normalizado = normalizar_header(texto)
    if not normalizado:
        return ""
    for canonico, aliases in HEADER_ALIASES.items():
        if normalizado == canonico or normalizado in aliases:
            return canonico
    return normalizado


def _resolver_header_importacao(linhas, campos_esperados=None):
    campos_esperados = set(campos_esperados or [])
    melhor_idx = 0
    melhor_score = -1
    for idx in range(min(len(linhas), 5)):
        headers_canonicos = [_canonicalizar_header(valor) for valor in linhas[idx]]
        score = 0
        for header in headers_canonicos:
            if not header:
                continue
            if header in campos_esperados:
                score += 3
            elif header in HEADER_ALIASES:
                score += 2
            elif any(header in aliases for aliases in HEADER_ALIASES.values()):
                score += 1
        if score > melhor_score:
            melhor_idx = idx
            melhor_score = score

    header_raw = linhas[melhor_idx]
    header = [_canonicalizar_header(h) for h in header_raw]
    mapa = {}
    for idx, nome in enumerate(header):
        if nome and nome not in mapa:
            mapa[nome] = idx
    return header_raw, header, mapa, linhas[melhor_idx + 1:]


def _valor_coluna(row, mapa, *colunas):
    for coluna in colunas:
        idx = mapa.get(coluna)
        if idx is None or idx >= len(row):
            continue
        valor = _corrigir_mojibake(row[idx])
        if valor != "":
            return valor
    return ""


def _coletar_campos_extras(row, header_raw, header, usados):
    extras = {}
    for idx, _ in enumerate(header_raw):
        nome = header[idx] if idx < len(header) else ""
        if not nome or nome in usados or nome in CAMPOS_PRODUTO_DESCARTADOS or idx >= len(row):
            continue
        valor = _corrigir_mojibake(row[idx])
        if valor == "":
            continue
        extras[nome] = valor
    return extras


def _montar_descricao_importacao(row, mapa):
    partes = [
        _valor_coluna(row, mapa, "descricao"),
        _valor_coluna(row, mapa, "descricao_secundaria"),
        _valor_coluna(row, mapa, "sufixo"),
    ]
    return " ".join(parte for parte in partes if parte).strip()


def _mesclar_dados_item(atual, novos, extras=None):
    item = dict(atual or {})
    for campo in CAMPOS_PRODUTO_DESCARTADOS:
        item.pop(campo, None)
    for campo in ITEM_CAMPOS_BASE:
        valor = novos.get(campo, "")
        if valor != "":
            item[campo] = valor
        else:
            item[campo] = item.get(campo, "")

    extras_existentes = dict(item.get("campos_extras", {}) or {})
    for chave, valor in (extras or {}).items():
        if valor != "":
            extras_existentes[chave] = valor
    if extras_existentes:
        item["campos_extras"] = extras_existentes
    elif "campos_extras" in item:
        item["campos_extras"] = item.get("campos_extras", {})
    return item


def ler_linhas_arquivo(file_storage):
    filename = secure_filename(file_storage.filename or "")
    ext = os.path.splitext(filename)[1].lower()

    if ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
        wb = load_workbook(file_storage, data_only=True)
        ws = wb.active
        linhas = []
        for row in ws.iter_rows(values_only=True):
            linha = [_corrigir_mojibake(cell) for cell in row]
            if any(linha):
                linhas.append(linha)
        return linhas

    return []


def importar_produtos(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        return 0

    header_raw, header, mapa, rows = _resolver_header_importacao(
        linhas,
        campos_esperados={"codigo", "descricao", "unidade", "grupo", "categoria"},
    )

    produtos = carregar_produtos()
    count = 0

    for row in rows:
        codigo = normalizar_codigo(_valor_coluna(row, mapa, "codigo"))
        if not codigo:
            continue
        atual = produtos.get(codigo, {})
        usados = {
            "codigo",
            "descricao",
            "descricao_secundaria",
            "sufixo",
            "unidade",
            "grupo",
            "categoria",
            "processo_conjunto",
        }
        novos = {
            "descricao": _montar_descricao_importacao(row, mapa),
            "unidade": _valor_coluna(row, mapa, "unidade"),
            "grupo": _valor_coluna(row, mapa, "grupo"),
            "categoria": _valor_coluna(row, mapa, "categoria"),
            "processo_conjunto": _valor_coluna(row, mapa, "processo_conjunto"),
        }
        extras = _coletar_campos_extras(row, header_raw, header, usados)
        produtos[codigo] = _mesclar_dados_item(atual, novos, extras=extras)
        count += 1

    salvar_json(PRODUTOS_FILE, produtos)
    return count


def atualizar_skus_arquivo(skus_file=None, somente_se_mais_novo=False):
    if supabase_catalog.enabled():
        if not somente_se_mais_novo:
            supabase_catalog.clear_cache()
        try:
            produtos = supabase_catalog.carregar_produtos(force=not somente_se_mais_novo)
            return {
                "arquivo": supabase_catalog.status().get("url", ""),
                "linhas": len(produtos),
                "atualizado": True,
                "ignorado": False,
                "erro": "",
            }
        except Exception as exc:
            app.logger.exception("Falha ao atualizar SKUs do Supabase")
            return {
                "arquivo": supabase_catalog.status().get("url", ""),
                "linhas": 0,
                "atualizado": False,
                "ignorado": False,
                "erro": f"Falha ao buscar SKUs no Supabase: {exc}",
            }

    skus_file = skus_file or get_skus_file()
    resultado = {
        "arquivo": skus_file,
        "linhas": 0,
        "atualizado": False,
        "ignorado": False,
        "erro": "",
    }
    if not skus_file or not os.path.isfile(_win_long_path(skus_file)):
        resultado["erro"] = f"Arquivo de SKUs nao encontrado: {skus_file}"
        return resultado

    if somente_se_mais_novo and os.path.exists(PRODUTOS_FILE):
        try:
            if os.path.getmtime(_win_long_path(skus_file)) <= os.path.getmtime(PRODUTOS_FILE):
                resultado["ignorado"] = True
                return resultado
        except Exception:
            pass

    try:
        with _open_for_read(skus_file) as arquivo:
            storage = FileStorage(stream=arquivo, filename=os.path.basename(skus_file))
            resultado["linhas"] = importar_produtos(storage)
        resultado["atualizado"] = resultado["linhas"] > 0
        app.logger.info("Importado cadastro de SKUs %s (%s linhas)", skus_file, resultado["linhas"])
    except PermissionError:
        app.logger.exception("Falha ao importar cadastro de SKUs (arquivo em uso) %s", skus_file)
        resultado["erro"] = f"Arquivo de SKUs em uso no Excel: {os.path.basename(skus_file)}"
    except Exception as exc:
        app.logger.exception("Falha ao importar cadastro de SKUs %s", skus_file)
        resultado["erro"] = f"Falha ao importar SKUs: {exc}"
    return resultado


def atualizar_skus_automatico():
    return atualizar_skus_arquivo(somente_se_mais_novo=True)


def importar_fornecedores(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        return 0

    _, _, mapa, rows = _resolver_header_importacao(
        linhas,
        campos_esperados={"fornecedor", "cnpj", "razao_social", "email"},
    )

    fornecedores = carregar_fornecedores()
    count = 0

    for row in rows:
        fornecedor = _valor_coluna(row, mapa, "fornecedor")
        cnpj = _valor_coluna(row, mapa, "cnpj")
        chave = cnpj if cnpj else fornecedor
        if not chave:
            continue

        atual = fornecedores.get(chave, {})
        def _pick(campo, valor):
            return valor if valor != "" else atual.get(campo, "")
        fornecedores[chave] = {
            "fornecedor": _pick("fornecedor", fornecedor),
            "razao_social": _pick("razao_social", _valor_coluna(row, mapa, "razao_social")),
            "cnpj": _pick("cnpj", cnpj),
            "email": _pick("email", _valor_coluna(row, mapa, "email")),
            "telefone": _pick("telefone", _valor_coluna(row, mapa, "telefone")),
            "endereco": _pick("endereco", _valor_coluna(row, mapa, "endereco")),
            "bairro": _pick("bairro", _valor_coluna(row, mapa, "bairro")),
            "cidade": _pick("cidade", _valor_coluna(row, mapa, "cidade")),
            "uf": _pick("uf", _valor_coluna(row, mapa, "uf")),
            "cep": _pick("cep", _valor_coluna(row, mapa, "cep")),
        }
        count += 1

    salvar_fornecedores(fornecedores)
    return count


def importar_os_produtos(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        return 0

    header_raw, header, mapa, rows = _resolver_header_importacao(
        linhas,
        campos_esperados={"codigo", "descricao", "unidade", "grupo", "categoria"},
    )

    produtos = carregar_os_produtos()
    count = 0

    for row in rows:
        codigo = normalizar_codigo(_valor_coluna(row, mapa, "codigo"))
        if not codigo:
            continue
        atual = produtos.get(codigo, {})
        usados = {
            "codigo",
            "descricao",
            "descricao_secundaria",
            "sufixo",
            "unidade",
            "grupo",
            "categoria",
        }
        novos = {
            "descricao": _montar_descricao_importacao(row, mapa),
            "unidade": _valor_coluna(row, mapa, "unidade"),
            "grupo": _valor_coluna(row, mapa, "grupo"),
            "categoria": _valor_coluna(row, mapa, "categoria"),
        }
        extras = _coletar_campos_extras(row, header_raw, header, usados)
        produtos[codigo] = _mesclar_dados_item(atual, novos, extras=extras)
        count += 1

    salvar_json(OS_PRODUTOS_FILE, produtos)
    return count


def importar_os_fornecedores(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        return 0

    _, _, mapa, rows = _resolver_header_importacao(
        linhas,
        campos_esperados={"cliente", "fornecedor"},
    )

    fornecedores = carregar_os_fornecedores()
    count = 0

    for row in rows:
        fornecedor = _valor_coluna(row, mapa, "fornecedor", "cliente")
        chave = fornecedor
        if not chave:
            continue

        atual = fornecedores.get(chave, {})
        fornecedores[chave] = {
            "cliente": fornecedor if fornecedor != "" else atual.get("cliente", ""),
        }
        count += 1

    salvar_os_fornecedores(fornecedores)
    return count


PESSOA_IMPORT_FIELDS = {
    "data_registro": ("data_de_registro", "data_registro"),
    "pessoa_fisica": ("pessoafisica", "pessoa_fisica"),
    "nome_fantasia": ("nomefantasia", "nome_fantasia", "fornecedor", "cliente"),
    "razao_social": ("razaosocial", "razao_social"),
    "cnpj_cpf": ("cnpj_cpf", "cnpj", "cpf"),
    "codigo_identificador_unico": ("codigo_identificador_unico",),
    "rg": ("rg",),
    "ie": ("ie",),
    "logradouro": ("logradouro", "endereco"),
    "logradouro_numero": ("logradouronumero", "logradouro_numero", "numero"),
    "complemento": ("complemento",),
    "bairro": ("bairro",),
    "cidade": ("cidade",),
    "codigo_municipio": ("codigomunicipio", "codigo_municipio"),
    "pais": ("pais",),
    "codigo_pais": ("codigopais", "codigo_pais"),
    "cep": ("cep",),
    "uf": ("uf",),
    "codigo_uf": ("codigouf", "codigo_uf"),
    "telefone": ("telefone",),
    "whatsapp": ("whatsapp",),
    "celular": ("celular",),
    "email": ("email",),
    "site": ("site",),
    "cliente": ("cliente",),
    "fornecedor": ("fornecedor",),
    "colaborador": ("colaborador",),
    "transportadora": ("transportadora",),
    "pessoa_grupo": ("pessoagrupo", "pessoa_grupo"),
    "identificador": ("identificador",),
    "vendedor_padrao": ("vendedorpadrao", "vendedor_padrao"),
    "categoria": ("categoria",),
    "tabela_preco": ("tabelapreco", "tabela_preco"),
    "observacoes": ("observacoes", "observacao"),
    "limite_credito": ("limite_de_credito", "limite_credito"),
    "periodicidade_venda_compra_dias": (
        "periodicidade_venda_compra_dias",
        "periodicidade_venda_compra",
    ),
    "validation": ("validation",),
    "valor_minimo_compra": ("valorminimocompra", "valor_minimo_compra"),
    "data_nascimento_fundacao": (
        "datanascimentofundacao",
        "data_nascimento_fundacao",
    ),
}


def importar_pessoas(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        return 0

    header_raw, header, mapa, rows = _resolver_header_importacao(
        linhas,
        campos_esperados={"nomefantasia", "razaosocial", "cnpj_cpf", "identificador"},
    )
    pessoas = []
    for row in rows:
        pessoa = {}
        for campo, aliases in PESSOA_IMPORT_FIELDS.items():
            pessoa[campo] = _valor_coluna(row, mapa, *aliases)
        if not any(pessoa.get(campo) for campo in ("nome_fantasia", "razao_social", "cnpj_cpf", "identificador")):
            continue
        pessoa["payload"] = _coletar_campos_extras(row, header_raw, header, set())
        pessoas.append(pessoa)

    if not supabase_data.enabled():
        fornecedores = carregar_fornecedores()
        clientes = carregar_os_fornecedores()
        for pessoa in pessoas:
            nome = pessoa.get("nome_fantasia") or pessoa.get("razao_social") or pessoa.get("cnpj_cpf")
            if not nome:
                continue
            legacy = {
                "fornecedor": nome,
                "cliente": nome,
                "razao_social": pessoa.get("razao_social", ""),
                "cnpj": pessoa.get("cnpj_cpf", ""),
                "email": pessoa.get("email", ""),
                "telefone": pessoa.get("telefone", ""),
                "endereco": pessoa.get("logradouro", ""),
                "bairro": pessoa.get("bairro", ""),
                "cidade": pessoa.get("cidade", ""),
                "uf": pessoa.get("uf", ""),
                "cep": pessoa.get("cep", ""),
            }
            if str(pessoa.get("fornecedor", "")).strip().upper() in {"SIM", "S", "1", "TRUE"}:
                fornecedores[pessoa.get("cnpj_cpf") or nome] = legacy
            if str(pessoa.get("cliente", "")).strip().upper() in {"SIM", "S", "1", "TRUE"}:
                clientes[nome] = {"cliente": nome}
        salvar_fornecedores(fornecedores)
        salvar_os_fornecedores(clientes)
        return len(pessoas)

    return supabase_data.salvar_pessoas(pessoas)


def _normalizar_codigo_planilha(valor):
    texto = _corrigir_mojibake(valor)
    if re.fullmatch(r"\d+\.0+", texto):
        texto = texto.split(".", 1)[0]
    return normalizar_codigo(texto)


def _separar_lista_planilha(valor):
    texto = _corrigir_mojibake(valor)
    valores = []
    for parte in re.split(r"[;\r\n]+", texto):
        parte = str(parte or "").strip()
        if parte and parte not in valores:
            valores.append(parte)
    return valores


def _quantidade_regra_planilha(valor):
    texto = _corrigir_mojibake(valor).replace(",", ".")
    if not texto:
        return 1.0
    try:
        quantidade = float(texto)
    except Exception as exc:
        raise ValueError("quantidade invalida") from exc
    if quantidade <= 0:
        raise ValueError("quantidade deve ser maior que zero")
    return quantidade


def _booleano_planilha(valor):
    texto = normalizar_header(_corrigir_mojibake(valor))
    if texto in {"sim", "s", "1", "true", "verdadeiro", "x"}:
        return True
    if texto in {"", "nao", "n", "0", "false", "falso"}:
        return False
    raise ValueError("quantidade_editavel deve ser SIM ou NAO")


def _assinatura_regra_popup(regra):
    return (
        normalizar_codigo(regra.get("gatilho", "")),
        tuple(normalizar_codigo(codigo) for codigo in regra.get("opcoes", []) if normalizar_codigo(codigo)),
        float(regra.get("quantidade", 1) or 1),
        bool(regra.get("quantidade_editavel", False)),
    )


def importar_regras_popup_item_planilha(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        raise ValueError("A planilha esta vazia ou nao esta em formato XLSX.")

    _, _, mapa, rows = _resolver_header_importacao(
        linhas,
        campos_esperados={
            "id_regra",
            "item_gatilho",
            "itens_opcoes",
            "quantidade",
            "quantidade_editavel",
        },
    )
    if "item_gatilho" not in mapa or "itens_opcoes" not in mapa:
        raise ValueError("Cabecalho invalido. Use ITEM_GATILHO e ITENS_OPCOES.")

    regras = carregar_regras_popup_item()
    resultado = [dict(regra) for regra in regras]
    indice_por_id = {
        str(regra.get("id", "") or "").strip(): idx
        for idx, regra in enumerate(resultado)
        if str(regra.get("id", "") or "").strip()
    }
    ids_usados = set(indice_por_id)
    assinaturas = {_assinatura_regra_popup(regra) for regra in resultado}
    proximo_numero = max(
        [int(re.sub(r"\D", "", regra_id) or 0) for regra_id in ids_usados] or [0]
    ) + 1
    novas = 0
    atualizadas = 0
    ignoradas = 0
    erros = []

    for numero_linha, row in enumerate(rows, 2):
        gatilho = _normalizar_codigo_planilha(_valor_coluna(row, mapa, "item_gatilho", "gatilho"))
        opcoes = [
            _normalizar_codigo_planilha(valor)
            for valor in _separar_lista_planilha(_valor_coluna(row, mapa, "itens_opcoes", "opcoes", "item_opcao"))
        ]
        opcoes = list(dict.fromkeys(codigo for codigo in opcoes if codigo))
        regra_id = _valor_coluna(row, mapa, "id_regra", "id").strip()
        if not gatilho and not opcoes and not regra_id:
            continue
        if not gatilho or not opcoes:
            ignoradas += 1
            erros.append(f"linha {numero_linha}: informe gatilho e opcoes")
            continue
        try:
            quantidade = _quantidade_regra_planilha(_valor_coluna(row, mapa, "quantidade"))
            quantidade_editavel = _booleano_planilha(
                _valor_coluna(row, mapa, "quantidade_editavel", "alteravel")
            )
        except ValueError as exc:
            ignoradas += 1
            erros.append(f"linha {numero_linha}: {exc}")
            continue

        registro = {
            "id": regra_id,
            "gatilho": gatilho,
            "opcoes": opcoes,
            "quantidade": quantidade,
            "quantidade_editavel": quantidade_editavel,
        }
        if regra_id and regra_id in indice_por_id:
            resultado[indice_por_id[regra_id]] = registro
            atualizadas += 1
            assinaturas.add(_assinatura_regra_popup(registro))
            continue

        assinatura = _assinatura_regra_popup(registro)
        if not regra_id and assinatura in assinaturas:
            ignoradas += 1
            continue
        if not regra_id:
            while f"regra-{proximo_numero}" in ids_usados:
                proximo_numero += 1
            regra_id = f"regra-{proximo_numero}"
            proximo_numero += 1
            registro["id"] = regra_id

        ids_usados.add(regra_id)
        indice_por_id[regra_id] = len(resultado)
        resultado.append(registro)
        assinaturas.add(assinatura)
        novas += 1

    if novas == 0 and atualizadas == 0 and ignoradas == 0:
        raise ValueError("Nenhuma regra foi encontrada na planilha.")
    if novas or atualizadas:
        salvar_regras_popup_item(resultado)
    return {
        "novas": novas,
        "atualizadas": atualizadas,
        "ignoradas": ignoradas,
        "erros": erros,
    }


def importar_relacoes_processo_item_planilha(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        raise ValueError("A planilha esta vazia ou nao esta em formato XLSX.")

    _, _, mapa, rows = _resolver_header_importacao(
        linhas,
        campos_esperados={"item_codigo", "processos"},
    )
    if "item_codigo" not in mapa or not any(campo in mapa for campo in ("processos", "processo", "processo_conjunto")):
        raise ValueError("Cabecalho invalido. Use ITEM_CODIGO e PROCESSOS.")

    importadas = {}
    ignoradas = 0
    erros = []
    for numero_linha, row in enumerate(rows, 2):
        codigo = _normalizar_codigo_planilha(_valor_coluna(row, mapa, "item_codigo", "codigo", "item"))
        processos = _separar_lista_planilha(
            _valor_coluna(row, mapa, "processos", "processo", "processo_conjunto")
        )
        if not codigo and not processos:
            continue
        if not codigo or not processos:
            ignoradas += 1
            erros.append(f"linha {numero_linha}: informe item e processos")
            continue
        destino = importadas.setdefault(codigo, [])
        for processo in processos:
            if processo not in destino:
                destino.append(processo)

    if not importadas and ignoradas == 0:
        raise ValueError("Nenhuma relacao foi encontrada na planilha.")

    relacoes = carregar_relacoes_processo_item()
    novas = 0
    atualizadas = 0
    sem_alteracao = 0
    for codigo, processos in importadas.items():
        atuais = relacoes.get(codigo)
        if atuais is None:
            novas += 1
        elif atuais == processos:
            sem_alteracao += 1
        else:
            atualizadas += 1
        relacoes[codigo] = processos

    if novas or atualizadas:
        salvar_relacoes_processo_item(relacoes)
    return {
        "novas": novas,
        "atualizadas": atualizadas,
        "ignoradas": ignoradas + sem_alteracao,
        "erros": erros,
    }


def importar_os_componentes(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        return 0

    header = [normalizar_header(h) for h in linhas[0]]
    mapa = {name: idx for idx, name in enumerate(header)}

    componentes = carregar_os_componentes()
    produtos = carregar_os_produtos()
    count = 0
    vistos = set()
    item_codigo_atual = ""

    def normalizar_descricao_item(valor):
        return " ".join(str(valor or "").strip().upper().split())

    descricao_para_codigos = {}
    for codigo_produto, info in (produtos or {}).items():
        descricao_produto = normalizar_descricao_item(info.get("descricao", ""))
        if not descricao_produto:
            continue
        descricao_para_codigos.setdefault(descricao_produto, set()).add(codigo_produto)

    item_bloco_descricao = ""
    idx_item_bloco = mapa.get("item_bloco")
    if idx_item_bloco is not None and (idx_item_bloco + 1) < len(linhas[0]):
        item_bloco_descricao = str(linhas[0][idx_item_bloco + 1]).strip()

    item_codigo_alias = ""
    if item_bloco_descricao:
        codigos_match = sorted(
            descricao_para_codigos.get(normalizar_descricao_item(item_bloco_descricao), set())
        )
        if len(codigos_match) == 1:
            item_codigo_alias = codigos_match[0]

    for row in linhas[1:]:
        def pegar(col):
            idx = mapa.get(col)
            if idx is None or idx >= len(row):
                return ""
            return str(row[idx]).strip()

        item_codigo_linha = normalizar_codigo(
            pegar("item_codigo")
            or pegar("codigo_item")
            or pegar("codigo_pai")
            or pegar("item")
            or pegar("codigo")
        )
        if item_codigo_linha:
            item_codigo_atual = item_codigo_linha

        item_codigo = item_codigo_atual
        if not item_codigo:
            continue

        comp = {
            "codigo": normalizar_codigo(pegar("componente_codigo") or pegar("codigo_componente")),
            "descricao": pegar("descricao"),
            "unidade": pegar("unidade"),
            "quantidade": pegar("quantidade") or pegar("qtd"),
        }

        if not (comp["codigo"] or comp["descricao"] or comp["unidade"] or comp["quantidade"]):
            continue

        destinos = [item_codigo]
        if item_codigo_alias and item_codigo_alias not in destinos:
            destinos.append(item_codigo_alias)

        for item_destino in destinos:
            if item_destino not in vistos:
                componentes[item_destino] = []
                vistos.add(item_destino)
            componentes.setdefault(item_destino, []).append(dict(comp))
        count += 1

    salvar_json(OS_COMPONENTES_FILE, componentes)
    return count


EXCEL_EXTS_IMPORTACAO = {".xls", ".xlsx", ".xlsm", ".xltx", ".xltm"}


def listar_arquivos_excel_base(pasta):
    arquivos = []
    for root, _, nomes in os.walk(pasta):
        for nome in nomes:
            if nome.startswith("~$"):
                continue
            if os.path.splitext(nome)[1].lower() in EXCEL_EXTS_IMPORTACAO:
                arquivos.append(os.path.join(root, nome))
    arquivos.sort()
    return arquivos


def importar_bom_diretorio(bom_dir, somente_se_mais_novo=False):
    arquivos = listar_arquivos_excel_base(bom_dir)
    if not arquivos:
        return {
            "arquivos": 0,
            "linhas": 0,
            "falhas": 0,
            "em_uso": [],
            "com_erro": [],
            "atualizado": False,
        }

    if somente_se_mais_novo and os.path.exists(OS_COMPONENTES_FILE):
        base_mtime = os.path.getmtime(OS_COMPONENTES_FILE)
        mtimes = []
        for caminho in arquivos:
            try:
                mtimes.append(os.path.getmtime(caminho))
            except FileNotFoundError:
                app.logger.warning("B.O.M. ignorada porque nao esta mais acessivel: %s", caminho)
        if not mtimes:
            return {
                "arquivos": 0,
                "linhas": 0,
                "falhas": 0,
                "em_uso": [],
                "com_erro": [],
                "atualizado": False,
            }
        mais_novo = max(mtimes)
        if mais_novo <= base_mtime:
            return {
                "arquivos": len(arquivos),
                "linhas": 0,
                "falhas": 0,
                "em_uso": [],
                "com_erro": [],
                "atualizado": False,
            }

    arquivos_processados = 0
    linhas_importadas = 0
    falhas = 0
    arquivos_em_uso = []
    arquivos_com_erro = []
    for caminho in arquivos:
        try:
            with _open_for_read(caminho) as arquivo:
                storage = FileStorage(stream=arquivo, filename=os.path.basename(caminho))
                linhas_importadas += importar_os_componentes(storage)
            arquivos_processados += 1
            app.logger.info("Importado B.O.M. %s", caminho)
        except PermissionError:
            app.logger.exception("Falha ao importar B.O.M. (arquivo em uso) %s", caminho)
            arquivos_em_uso.append(os.path.basename(caminho))
            falhas += 1
        except Exception:
            app.logger.exception("Falha ao importar B.O.M. %s", caminho)
            arquivos_com_erro.append(os.path.basename(caminho))
            falhas += 1

    return {
        "arquivos": arquivos_processados,
        "linhas": linhas_importadas,
        "falhas": falhas,
        "em_uso": arquivos_em_uso,
        "com_erro": arquivos_com_erro,
        "atualizado": arquivos_processados > 0,
    }


def importar_os_processos(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        return 0

    header_raw = ["" if h is None else str(h).strip() for h in linhas[0]]
    header_norm = [normalizar_header(h) for h in header_raw]

    processos = {"PADRAO": {}}
    count = 0

    # formato "longo": colunas processo/atividade/responsavel
    if "processo" in header_norm:
        mapa = {name: idx for idx, name in enumerate(header_norm)}
        staged_por_conjunto = {}

        for row in linhas[1:]:
            def pegar(col):
                idx = mapa.get(col)
                if idx is None or idx >= len(row):
                    return ""
                return str(row[idx]).strip()

            conjunto = pegar("conjunto") or pegar("grupo") or "PADRAO"
            processo = pegar("processo")
            atividade = pegar("atividade")
            responsavel = pegar("responsavel")
            if not processo or not atividade:
                continue

            processo_norm = processo.strip().upper()
            mapa_proc = {
                "CORTE": "CORTE",
                "AR CONDICIONADO": "AR CONDICIONADO",
                "PREPARACAO DE PECAS": "PREPARAÇÃO DE PEÇAS",
                "PREPARAÇÃO DE PEÇAS": "PREPARAÇÃO DE PEÇAS",
                "ISOLAMENTO": "ISOLAMENTO",
                "DESMONTAGEM E ISOLAMENTO": "ISOLAMENTO",
                "REVESTIMENTO": "REVESTIMENTO",
                "BANCOS": "BANCOS",
                "ELETRICA": "ELÉTRICA 2",
                "ELÉTRICA": "ELÉTRICA 2",
                "ELÉTRICA 2": "ELÉTRICA 2",
                "LIMPEZA/LIBERACAO": "LIMPEZA/LIBERAÇÃO",
                "LIMPEZA/LIBERAÇÃO": "LIMPEZA/LIBERAÇÃO",
            }
            processo_key = normalizar_nome_processo(mapa_proc.get(processo_norm, processo))

            processos.setdefault(conjunto, {})
            processos[conjunto].setdefault(processo_key, []).append({
                "atividade": atividade,
                "responsavel": responsavel,
            })
            count += 1

        salvar_os_processos(processos)
        return count

    # formato "largo": cada coluna = processo, cada linha = atividade
    conjunto = "PADRAO"
    processos.setdefault(conjunto, {})

    for col_idx, processo in enumerate(header_raw):
        processo = processo.strip()
        if not processo:
            continue
        processo_norm = processo.upper()
        mapa_proc = {
            "CORTE": "CORTE",
            "AR CONDICIONADO": "AR CONDICIONADO",
            "PREPARACAO DE PECAS": "PREPARAÇÃO DE PEÇAS",
            "PREPARAÇÃO DE PEÇAS": "PREPARAÇÃO DE PEÇAS",
            "ISOLAMENTO": "ISOLAMENTO",
            "DESMONTAGEM E ISOLAMENTO": "ISOLAMENTO",
            "REVESTIMENTO": "REVESTIMENTO",
            "BANCOS": "BANCOS",
            "ELETRICA": "ELÉTRICA 2",
            "ELÉTRICA": "ELÉTRICA 2",
            "ELÉTRICA 2": "ELÉTRICA 2",
            "LIMPEZA/LIBERACAO": "LIMPEZA/LIBERAÇÃO",
            "LIMPEZA/LIBERAÇÃO": "LIMPEZA/LIBERAÇÃO",
        }
        processo_key = normalizar_nome_processo(mapa_proc.get(processo_norm, processo))
        processos[conjunto].setdefault(processo_key, [])

        for row in linhas[1:]:
            if col_idx >= len(row):
                continue
            atividade = str(row[col_idx]).strip()
            if not atividade:
                continue
            processos[conjunto][processo_key].append({
                "atividade": atividade,
                "responsavel": "",
            })
            count += 1

    salvar_os_processos(processos)
    return count


def importar_os_processos_atualizado(file_storage):
    linhas = ler_linhas_arquivo(file_storage)
    if not linhas:
        return 0

    def nome_conjunto_arquivo(filename):
        nome = os.path.splitext(os.path.basename(filename or ""))[0].strip()
        if not nome:
            return "PADRAO"
        partes = [parte.strip() for parte in nome.split(" - ") if parte.strip()]
        if len(partes) >= 3 and normalizar_header(partes[1]).startswith("processo_transformacao"):
            nome = " - ".join(partes[2:])
        if normalizar_header(nome) == "template":
            return "PADRAO"
        return nome or "PADRAO"

    def remover_sufixo_banco(nome):
        partes = [parte.strip() for parte in str(nome or "").split(" - ") if parte.strip()]
        if len(partes) >= 2 and normalizar_nome_processo(partes[-1]) == "BANCO":
            return " - ".join(partes[:-1]).strip() or nome
        return nome

    header_raw = ["" if h is None else str(h).strip() for h in linhas[0]]
    header_norm = [normalizar_header(h) for h in header_raw]

    processos = carregar_os_processos()
    conjunto_arquivo = nome_conjunto_arquivo(getattr(file_storage, "filename", ""))
    count = 0
    conjuntos_tocados = set()

    def preparar_conjunto(conjunto, reset_all=False, preservar_processos=None):
        if conjunto in conjuntos_tocados:
            return
        base = processos.get(conjunto, {})
        processos[conjunto] = {
            nome: list(base.get(nome, [])) if not reset_all else []
            for nome in PROCESSOS_ORDEM
        }
        for nome in preservar_processos or []:
            if nome in processos[conjunto]:
                processos[conjunto][nome] = list(base.get(nome, []))
        conjuntos_tocados.add(conjunto)

    if "processo" in header_norm:
        mapa = {name: idx for idx, name in enumerate(header_norm)}
        staged_por_conjunto = {}

        for row in linhas[1:]:
            def pegar(col):
                idx = mapa.get(col)
                if idx is None or idx >= len(row):
                    return ""
                return str(row[idx]).strip()

            conjunto = pegar("conjunto") or pegar("grupo") or conjunto_arquivo or "PADRAO"
            processo = normalizar_nome_processo(pegar("processo"))
            atividade = pegar("atividade")
            responsavel = pegar("responsavel")
            if processo not in PROCESSOS_ORDEM or not atividade:
                continue

            staged_por_conjunto.setdefault(conjunto, {}).setdefault(processo, []).append({
                "atividade": atividade,
                "responsavel": responsavel,
            })
            count += 1

        if count == 0:
            return 0

        for conjunto, staged in staged_por_conjunto.items():
            preservar = set(PROCESSOS_ORDEM) - set(staged)
            preparar_conjunto(conjunto, reset_all=True, preservar_processos=preservar)
            for processo, linhas_processo in staged.items():
                processos[conjunto][processo] = linhas_processo

        salvar_os_processos(processos)
        return count

    conjunto = conjunto_arquivo or "PADRAO"
    colunas_processo = []
    for col_idx, processo in enumerate(header_raw):
        processo_nome = normalizar_nome_processo(processo)
        if processo_nome not in PROCESSOS_ORDEM:
            continue
        colunas_processo.append((col_idx, processo_nome))

    if not colunas_processo:
        return 0

    staged = {}
    for col_idx, processo_nome in colunas_processo:
        linhas_processo = []
        for row in linhas[1:]:
            if col_idx >= len(row):
                continue
            atividade = str(row[col_idx]).strip()
            if not atividade:
                continue
            linhas_processo.append({
                "atividade": atividade,
                "responsavel": "",
            })
        staged[processo_nome] = linhas_processo
        count += len(linhas_processo)

    # Nao limpamos a base quando o arquivo e apenas um template vazio.
    if count == 0:
        return 0

    nomes_importados = set(staged)
    somente_banco = nomes_importados == {"BANCO"}
    if somente_banco:
        conjunto = remover_sufixo_banco(conjunto)
    preservar = set(PROCESSOS_ORDEM) - nomes_importados
    preparar_conjunto(conjunto, reset_all=True, preservar_processos=preservar)

    for processo_nome, linhas_processo in staged.items():
        processos[conjunto][processo_nome] = linhas_processo

    salvar_os_processos(processos)
    return count


def mesclar_processos_modelo(processos, conjuntos):
    processos_final = {nome: [] for nome in PROCESSOS_ORDEM}
    vistos = {nome: set() for nome in PROCESSOS_ORDEM}
    for conjunto in conjuntos or []:
        modelo = processos.get(conjunto) or {}
        for nome in PROCESSOS_ORDEM:
            for linha in modelo.get(nome) or []:
                atividade = str((linha or {}).get("atividade", "")).strip()
                responsavel = str((linha or {}).get("responsavel", "")).strip()
                if not atividade:
                    continue
                chave = (atividade, responsavel)
                if chave in vistos[nome]:
                    continue
                processos_final[nome].append({
                    "atividade": atividade,
                    "responsavel": responsavel,
                })
                vistos[nome].add(chave)
    return processos_final


_document_number_lock = threading.Lock()


def _proximo_numero_documento(tipo, counter_file):
    if supabase_data.enabled():
        return supabase_data.proximo_numero_documento(tipo)
    with _document_number_lock:
        numeros = []
        for documento in _carregar_historico_local():
            if documento.get("tipo") != tipo:
                continue
            numero = str(documento.get("numero") or "").strip()
            if numero.isdigit():
                numeros.append(int(numero))
        numero = max(numeros, default=0) + 1
        with open(counter_file, "w", encoding="utf-8") as f:
            f.write(str(numero))
        return numero


def proximo_numero_oc():
    return _proximo_numero_documento("oc", COUNTER_FILE)


def proximo_numero_os():
    return _proximo_numero_documento("os", OS_COUNTER_FILE)


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        try:
            user = supabase_data.verify_user(username, password)
        except Exception:
            app.logger.exception("Falha ao validar login no Supabase")
            if shared_rbac_enabled():
                return render_template(
                    "login.html",
                    erro="Autorizacao compartilhada temporariamente indisponivel.",
                    next_url=request.args.get("next") or url_for("index"),
                ), 503
            user = None
        if user:
            session.clear()
            session["suprimentos_user"] = user
            return redirect(request.args.get("next") or url_for("index"))
        return render_template(
            "login.html",
            erro="Usuario ou senha invalidos.",
            next_url=request.args.get("next") or url_for("index"),
        )
    if current_user():
        return redirect(url_for("index"))
    return render_template(
        "login.html",
        erro="",
        next_url=request.args.get("next") or url_for("index"),
    )


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


def _enriquecer_compras_integradas(historico):
    """Attach the live ERP purchase-order read model without rewriting history."""
    if not erp_feature_enabled():
        return historico
    try:
        compras = _erp_stock_request("dashboard").get("orders", [])
        por_id = {
            str(ordem.get("id") or ""): ordem
            for ordem in compras
            if ordem.get("id")
        }
        por_chave = {
            str(ordem.get("idempotency_key") or ""): ordem
            for ordem in compras
            if ordem.get("idempotency_key")
        }
        for documento in historico:
            if documento.get("tipo") != "oc":
                continue
            erp_id = str(
                documento.get("erp_purchase_order_id")
                or (documento.get("dados") or {}).get("erp_purchase_order_id")
                or ""
            )
            chave = f"suprimentos-oc:{documento.get('id')}"
            if erp_id in por_id:
                documento["erp_purchase_order"] = por_id[erp_id]
            elif chave in por_chave:
                documento["erp_purchase_order"] = por_chave[chave]
    except Exception:
        app.logger.exception("Falha ao consultar estados integrados de O.C.; historico legado preservado.")
    return historico


def _enriquecer_historico_integrado(historico, tab):
    """Attach live read models without rewriting legacy document JSON."""
    if not erp_feature_enabled() or tab not in {"dashboard", "gestao-oc", "gestao-os"}:
        return historico

    _enriquecer_compras_integradas(historico)

    try:
        ordens_mes = _erp_mes_request("work-orders").get("orders", [])
        por_id = {
            str(ordem.get("id") or ""): ordem
            for ordem in ordens_mes
            if ordem.get("id")
        }
        por_numero = {}
        for ordem in ordens_mes:
            numero = str(ordem.get("numero_os") or "").strip()
            if numero and numero not in por_numero:
                por_numero[numero] = ordem
        for documento in historico:
            if documento.get("tipo") != "os":
                continue
            erp_id = str(
                documento.get("erp_work_order_id")
                or (documento.get("dados") or {}).get("erp_work_order_id")
                or ""
            )
            numero = str(documento.get("numero") or "").strip()
            if erp_id in por_id:
                documento["mes_work_order"] = por_id[erp_id]
            elif numero in por_numero:
                documento["mes_work_order"] = por_numero[numero]
    except Exception:
        app.logger.exception("Falha ao consultar estados do MES; historico legado preservado.")
    return historico


@app.route("/")
def index():
    if request.method == "HEAD":
        return "", 200

    resultado_skus_auto = atualizar_skus_automatico()
    produtos = carregar_produtos()
    fornecedores = carregar_fornecedores()
    os_produtos = carregar_os_produtos()
    os_fornecedores = carregar_os_fornecedores()
    os_componentes = carregar_os_componentes()
    os_processos = carregar_os_processos()
    relacoes_processo_item = carregar_relacoes_processo_item()
    regras_popup_item = carregar_regras_popup_item()
    processo_por_item = construir_processo_por_item(
        os_produtos,
        os_processos.keys(),
        relacoes_processo_item,
    )
    oc_prefill = carregar_importacao(_user_scoped_file(OC_IMPORT_FILE))
    os_prefill = carregar_importacao(_user_scoped_file(OS_IMPORT_FILE))
    requested_tab = str(request.args.get("tab") or "").strip().lower()
    tab = requested_tab or "oc"
    tab_permissions = {
        "oc": "suprimentos.purchase.create",
        "dashboard": "suprimentos.dashboard.view",
        "gestao-oc": "suprimentos.purchase.view",
        "gestao-os": "suprimentos.work_order.view",
        "os": "suprimentos.work_order.manage",
        "cadastro": "suprimentos.master_data.manage",
    }
    if shared_rbac_enabled() and not can(tab_permissions.get(tab, "suprimentos.dashboard.view")):
        if requested_tab:
            return _authorization_denied(tab_permissions.get(tab, "suprimentos.dashboard.view"))
        for candidate in ("dashboard", "gestao-oc", "gestao-os", "oc", "os", "cadastro"):
            if can(tab_permissions[candidate]):
                tab = candidate
                break
    historico = _enriquecer_historico_integrado(carregar_historico(), tab)
    oc_totais = _agrupar_por_data(historico, "oc", "total_pedido")
    os_quantidades = _agrupar_por_data(historico, "os", None)
    dashboard = {
        "oc_totais": oc_totais,
        "os_quantidades": os_quantidades,
        "totais": _dashboard_totais(historico),
        "recentes": _dashboard_recentes(historico),
        "documentos": _dashboard_documentos(historico),
        "ordens_compra": [entry for entry in historico if entry.get("tipo") == "oc"],
        "ordens_servico": [entry for entry in historico if entry.get("tipo") == "os"],
        "persistencia": "supabase" if supabase_data.enabled() else "local",
    }
    pedidos_dir, os_dir = get_save_paths()
    bom_dir = get_bom_dir()
    skus_file = get_skus_file()
    processos_dir = get_processos_dir()
    save_paths = {
        "pedidos_dir": pedidos_dir,
        "os_dir": os_dir,
        "bom_dir": bom_dir,
        "skus_file": skus_file,
        "processos_dir": processos_dir,
    }
    catalogo_status = supabase_catalog.status()
    dados_status = supabase_data.status()
    bom_status = request.args.get("bom_status")
    skus_status = request.args.get("skus_status")
    if not skus_status and resultado_skus_auto.get("atualizado"):
        if catalogo_status.get("enabled"):
            skus_status = f"Catalogo Supabase conectado ({resultado_skus_auto['linhas']} SKUs)."
        else:
            skus_status = f"SKUs atualizados automaticamente ({resultado_skus_auto['linhas']} linhas)."
    os_processos_status = request.args.get("os_processos_status")
    pessoas_status = request.args.get("pessoas_status")
    regras_popup_status = request.args.get("regras_popup_status")
    relacoes_processo_status = request.args.get("relacoes_processo_status")
    documento_status = request.args.get("documento_status")

    return render_template(
        "index.html",
        produtos=produtos,
        fornecedores=fornecedores,
        os_produtos=os_produtos,
        os_fornecedores=os_fornecedores,
        os_componentes=os_componentes,
        os_processos=os_processos,
        oc_prefill=oc_prefill,
        os_prefill=os_prefill,
        tab=tab,
        dashboard=dashboard,
        save_paths=save_paths,
        catalogo_status=catalogo_status,
        dados_status=dados_status,
        current_user=current_user(),
        bom_dir=bom_dir,
        skus_file=skus_file,
        processos_dir=processos_dir,
        bom_status=bom_status,
        skus_status=skus_status,
        os_processos_status=os_processos_status,
        pessoas_status=pessoas_status,
        regras_popup_status=regras_popup_status,
        relacoes_processo_status=relacoes_processo_status,
        documento_status=documento_status,
        processos_os=PROCESSOS_OS,
        processo_transformacao_por_item=processo_por_item,
        relacoes_processo_transformacao=RELACOES_PROCESSO_TRANSFORMACAO,
        relacoes_processo_item=relacoes_processo_item,
        regras_popup_item=regras_popup_item,
    )


@app.route("/api/historico/os/<documento_id>")
@permission_required("suprimentos.work_order.view")
def api_historico_os(documento_id):
    documento = obter_historico_documento(documento_id)
    if not documento or documento.get("tipo") != "os":
        return jsonify({"ok": False, "erro": "O.S nao encontrada."}), 404
    return jsonify({"ok": True, "documento": documento})


@app.route("/api/historico/os/<documento_id>/excluir", methods=["POST"])
@permission_required("suprimentos.work_order.manage")
def api_excluir_historico_os(documento_id):
    documento = obter_historico_documento(documento_id)
    if not documento or documento.get("tipo") != "os":
        return jsonify({"ok": False, "erro": "O.S nao encontrada."}), 404
    try:
        excluir_historico_documento(documento_id)
    except Exception:
        app.logger.exception("Falha ao excluir historico da O.S %s", documento_id)
        return jsonify({"ok": False, "erro": "Nao foi possivel excluir a O.S."}), 500
    return jsonify({"ok": True})


@app.route("/api/historico/oc/<documento_id>")
@permission_required("suprimentos.purchase.view")
def api_historico_oc(documento_id):
    documento = obter_historico_documento(documento_id)
    if not documento or documento.get("tipo") != "oc":
        return jsonify({"ok": False, "erro": "O.C nao encontrada."}), 404
    _enriquecer_compras_integradas([documento])
    return jsonify({"ok": True, "documento": documento})


@app.route("/api/historico/oc/<documento_id>/excluir", methods=["POST"])
@permission_required("suprimentos.purchase.edit")
def api_excluir_historico_oc(documento_id):
    documento = obter_historico_documento(documento_id)
    if not documento or documento.get("tipo") != "oc":
        return jsonify({"ok": False, "erro": "O.C nao encontrada."}), 404
    if erp_feature_enabled() and str(documento.get("status") or "").lower() != "rascunho":
        return jsonify({"ok": False, "erro": "O.C. emitida nao pode ser excluida. Use Cancelar para manter a auditoria e a integracao."}), 409
    try:
        excluir_historico_documento(documento_id)
    except Exception:
        app.logger.exception("Falha ao excluir historico da O.C %s", documento_id)
        return jsonify({"ok": False, "erro": "Nao foi possivel excluir a O.C."}), 500
    return jsonify({"ok": True})


def _historico_status_required_permission(tipo, documento_id):
    document_type = str(tipo or "").strip().lower()
    payload = request.get_json(silent=True) or request.form
    new_status = str(payload.get("status") or "").strip().lower()
    target_permission = {
        ("oc", "cancelado"): "suprimentos.purchase.cancel",
        ("oc", "concluido"): "suprimentos.purchase.technical_close",
        ("os", "concluido"): "suprimentos.work_order.technical_close",
    }.get((document_type, new_status))
    if not target_permission and document_type == "oc":
        target_permission = "suprimentos.purchase.edit"
    if not target_permission and document_type == "os":
        target_permission = "suprimentos.work_order.manage"
    if not target_permission:
        return "suprimentos.system.admin"

    permissions = [target_permission]
    document = obter_historico_documento(documento_id)
    current_status = str((document or {}).get("status") or "").strip().lower()
    if current_status != new_status:
        if document_type == "oc" and current_status == "cancelado":
            permissions.append("suprimentos.purchase.cancel")
        if document_type == "oc" and current_status == "concluido":
            permissions.append("suprimentos.purchase.technical_close")
        if document_type == "os" and current_status == "concluido":
            permissions.append("suprimentos.work_order.technical_close")
    return tuple(dict.fromkeys(permissions))


@app.route("/api/historico/<tipo>/<documento_id>/status", methods=["POST"])
@permission_required(_historico_status_required_permission)
def api_status_historico(tipo, documento_id):
    tipo = str(tipo or "").strip().lower()
    if tipo not in {"oc", "os"}:
        return jsonify({"ok": False, "erro": "Tipo de documento invalido."}), 400
    documento = obter_historico_documento(documento_id)
    if not documento or documento.get("tipo") != tipo:
        return jsonify({"ok": False, "erro": "Documento nao encontrado."}), 404
    payload = request.get_json(silent=True) or request.form
    try:
        novo_status = str(payload.get("status") or "").strip().lower()
        if tipo == "oc" and novo_status == "cancelado" and erp_feature_enabled():
            _cancel_emitted_legacy_oc_in_erp(documento, "Cancelada pelo comprador no Suprimentos.")
        if tipo == "oc" and novo_status == "concluido" and erp_feature_enabled():
            _close_emitted_legacy_oc_in_erp(documento, "Concluida tecnicamente pelo comprador no Suprimentos.")
        if tipo == "os" and erp_feature_enabled():
            if novo_status == "concluido":
                _close_linked_legacy_os_in_mes(
                    documento,
                    "Concluída tecnicamente pelo PCP em Suprimentos.",
                )
            elif str(documento.get("status") or "").lower() == "concluido":
                _reopen_linked_legacy_os_in_mes(
                    documento,
                    "Conclusão técnica reaberta em Suprimentos.",
                )
        atualizado = atualizar_status_historico_documento(documento_id, payload.get("status"))
    except ValueError as exc:
        return jsonify({"ok": False, "erro": str(exc)}), 400
    except Exception:
        app.logger.exception("Falha ao atualizar status do documento %s", documento_id)
        return jsonify({"ok": False, "erro": "Nao foi possivel atualizar o status."}), 500
    return jsonify({"ok": True, "documento": atualizado})


OS_LUMINARIAS_CODIGOS = {"10260092", "10260095"}


def _revisao_nome_arquivo_os(nome):
    match = re.search(r"(?:^|[ _-])R(\d+)(?:\D|$)", str(nome or ""), re.IGNORECASE)
    return int(match.group(1)) if match else 0


def _fontes_os_reconciliacao(arquivos):
    fontes = []
    total_descompactado = 0
    ordem = 0
    for arquivo in arquivos or []:
        nome_arquivo = secure_filename(getattr(arquivo, "filename", "") or "")
        if not nome_arquivo:
            continue
        extensao = os.path.splitext(nome_arquivo)[1].lower()
        conteudo = arquivo.read() or b""
        if extensao == ".json":
            try:
                manifesto = json.loads(conteudo.decode("utf-8"))
            except (UnicodeDecodeError, json.JSONDecodeError) as exc:
                raise ValueError(f"Manifesto JSON invalido: {nome_arquivo}.") from exc
            entradas = manifesto.get("ordens", []) if isinstance(manifesto, dict) else manifesto
            if not isinstance(entradas, list) or len(entradas) > 500:
                raise ValueError("O manifesto deve conter uma lista de ate 500 O.S.")
            for entrada in entradas:
                if not isinstance(entrada, dict) or not isinstance(entrada.get("dados"), dict):
                    raise ValueError("O manifesto possui uma entrada de O.S. invalida.")
                nome = secure_filename(entrada.get("nome") or f"os-{ordem + 1}.docx")
                try:
                    data_arquivo = datetime.fromisoformat(str(entrada.get("data_arquivo") or ""))
                except ValueError:
                    data_arquivo = datetime.now()
                fontes.append({
                    "nome": nome,
                    "conteudo": b"",
                    "dados": entrada["dados"],
                    "hash": str(entrada.get("hash") or "").strip(),
                    "data_arquivo": data_arquivo,
                    "revisao": int(entrada.get("revisao") or _revisao_nome_arquivo_os(nome)),
                    "ordem": ordem,
                })
                ordem += 1
            continue
        if extensao == ".docx":
            fontes.append({
                "nome": nome_arquivo,
                "conteudo": conteudo,
                "data_arquivo": datetime.now(),
                "revisao": _revisao_nome_arquivo_os(nome_arquivo),
                "ordem": ordem,
            })
            ordem += 1
            continue
        if extensao != ".zip":
            raise ValueError(f"Arquivo nao suportado: {nome_arquivo}.")
        try:
            with zipfile.ZipFile(io.BytesIO(conteudo)) as pacote:
                membros = [
                    info for info in pacote.infolist()
                    if not info.is_dir() and os.path.splitext(info.filename)[1].lower() == ".docx"
                ]
                if len(membros) > 500:
                    raise ValueError("O ZIP excede o limite de 500 arquivos DOCX.")
                for info in membros:
                    total_descompactado += int(info.file_size or 0)
                    if total_descompactado > 128 * 1024 * 1024:
                        raise ValueError("O ZIP excede o limite descompactado de 128 MB.")
                    nome = secure_filename(os.path.basename(info.filename))
                    fontes.append({
                        "nome": nome,
                        "conteudo": pacote.read(info),
                        "data_arquivo": datetime(*info.date_time),
                        "revisao": _revisao_nome_arquivo_os(nome),
                        "ordem": ordem,
                    })
                    ordem += 1
        except zipfile.BadZipFile as exc:
            raise ValueError(f"ZIP invalido: {nome_arquivo}.") from exc
    if not fontes:
        raise ValueError("Nenhum arquivo DOCX foi encontrado para reconciliar.")
    return fontes


def _parsear_fontes_os_reconciliacao(fontes):
    por_numero = {}
    descartadas = 0
    for fonte in fontes:
        dados = fonte.get("dados")
        if not isinstance(dados, dict):
            armazenamento = FileStorage(
                stream=io.BytesIO(fonte["conteudo"]),
                filename=fonte["nome"],
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            try:
                dados = parse_os_docx_atualizado(armazenamento)
            except Exception as exc:
                raise ValueError(f"Nao foi possivel ler {fonte['nome']}.") from exc
        numero = str(dados.get("os_numero") or "").strip()
        if not numero:
            raise ValueError(f"O arquivo {fonte['nome']} nao possui numero de O.S.")
        if not (dados.get("itens") or []):
            raise ValueError(f"A O.S. {numero} de {fonte['nome']} nao possui itens.")
        candidata = dict(fonte)
        candidata["dados"] = dados
        candidata["hash"] = fonte.get("hash") or hashlib.sha256(
            fonte.get("conteudo") or json.dumps(dados, ensure_ascii=False, sort_keys=True).encode("utf-8")
        ).hexdigest()
        atual = por_numero.get(numero)
        if atual is None:
            por_numero[numero] = candidata
            continue
        descartadas += 1
        chave_atual = (atual["data_arquivo"], atual["revisao"], atual["ordem"])
        chave_nova = (candidata["data_arquivo"], candidata["revisao"], candidata["ordem"])
        if candidata["hash"] != atual["hash"] and chave_nova > chave_atual:
            por_numero[numero] = candidata
    return por_numero, descartadas


def _referencia_os_recente(documentos, numero, chassis=""):
    numero = str(numero or "").strip()
    chassis_norm = str(chassis or "").strip().casefold()
    candidatas = [
        documento for documento in (documentos or [])
        if documento.get("tipo") == "os" and str(documento.get("numero") or "").strip() == numero
    ]
    exatas = [
        documento for documento in candidatas
        if str((documento.get("dados") or {}).get("chassis") or "").strip().casefold() == chassis_norm
    ]
    if exatas:
        candidatas = exatas
    if not candidatas:
        return None
    return max(
        candidatas,
        key=lambda documento: (
            str(documento.get("updated_at") or ""),
            str(documento.get("created_at") or ""),
            str(documento.get("data_criacao") or ""),
            str(documento.get("id") or ""),
        ),
    )


def _fornecedor_item_referencia(referencia, codigo):
    codigo = normalizar_codigo(codigo)
    for item in (referencia or {}).get("itens", []) or []:
        if normalizar_codigo((item or {}).get("codigo", "")) != codigo:
            continue
        fornecedor = str((item or {}).get("fornecedor") or "").strip()
        if fornecedor:
            return fornecedor
    return ""


def _linhas_inferencia_os(dados_origem, referencia):
    linhas = []
    for linha in (referencia or {}).get("composicao", []) or []:
        if isinstance(linha, dict):
            copia = dict(linha)
            copia["_origem_inferencia"] = "referencia"
            linhas.append(copia)
    for linha in dados_origem.get("composicao", []) or []:
        if isinstance(linha, dict):
            copia = dict(linha)
            copia["_origem_inferencia"] = "arquivo"
            linhas.append(copia)
    return linhas


def _selecionar_codigo_relacionado(linhas, item_raiz, opcoes):
    item_raiz = normalizar_codigo(item_raiz)
    opcoes = {normalizar_codigo(codigo) for codigo in opcoes if normalizar_codigo(codigo)}
    grupos = [
        [
            linha for linha in linhas
            if linha.get("_origem_inferencia") == "arquivo"
            and normalizar_codigo(linha.get("item", "")) == item_raiz
            and normalizar_codigo(linha.get("codigo", "")) in opcoes
        ],
        [
            linha for linha in linhas
            if linha.get("_origem_inferencia") == "arquivo"
            and normalizar_codigo(linha.get("codigo", "")) in opcoes
        ],
        [
            linha for linha in linhas
            if linha.get("_origem_inferencia") == "referencia"
            and normalizar_codigo(linha.get("item", "")) == item_raiz
            and normalizar_codigo(linha.get("codigo", "")) in opcoes
        ],
        [
            linha for linha in linhas
            if linha.get("_origem_inferencia") == "referencia"
            and normalizar_codigo(linha.get("codigo", "")) in opcoes
        ],
    ]
    for candidatas in grupos:
        codigos = list(dict.fromkeys(normalizar_codigo(linha.get("codigo", "")) for linha in candidatas))
        if len(codigos) == 1:
            codigo = codigos[0]
            linha = next(linha for linha in candidatas if normalizar_codigo(linha.get("codigo", "")) == codigo)
            return codigo, _parse_numero_form(linha.get("qtd", 1), 1.0) or 1.0
    return "", 0


def _luminarias_historicas_por_item(documentos):
    candidatas = {}
    for documento in documentos or []:
        if str((documento or {}).get("tipo") or "").strip().lower() != "os":
            continue
        for linha in (documento or {}).get("composicao", []) or []:
            if not isinstance(linha, dict) or _linha_status(linha) == "cancelado":
                continue
            item = normalizar_codigo(linha.get("item", ""))
            codigo = normalizar_codigo(linha.get("codigo", ""))
            if not item.startswith("4034") or codigo not in OS_LUMINARIAS_CODIGOS:
                continue
            qtd = _parse_numero_form(linha.get("qtd", 1), 1.0) or 1.0
            candidatas.setdefault(item, set()).add((codigo, qtd))

    return {
        item: {"codigo": codigo, "qtd": qtd}
        for item, relacoes in candidatas.items()
        if len(relacoes) == 1
        for codigo, qtd in relacoes
    }


def _gatilhos_popup_alcancaveis(codigo_raiz, componentes, regras_por_gatilho):
    encontrados = []
    visitados = set()

    def visitar(codigo):
        codigo = normalizar_codigo(codigo)
        if not codigo or codigo in visitados:
            return
        visitados.add(codigo)
        if codigo in regras_por_gatilho:
            encontrados.append(codigo)
        for componente in (componentes or {}).get(codigo, []) or []:
            visitar((componente or {}).get("codigo", ""))

    visitar(codigo_raiz)
    return encontrados


def _montar_os_reconciliada(fonte, referencia, contexto):
    dados_origem = fonte["dados"]
    os_produtos = contexto["os_produtos"]
    produtos_catalogo = contexto["produtos_catalogo"]
    componentes = contexto["componentes"]
    regras_por_gatilho = contexto["regras_por_gatilho"]
    itens = []
    extras = []
    linhas_inferencia = _linhas_inferencia_os(dados_origem, referencia)

    for item_origem in dados_origem.get("itens", []) or []:
        codigo = normalizar_codigo(item_origem.get("codigo", ""))
        if not codigo:
            continue
        item_info = os_produtos.get(codigo, {}) or {}
        qtd = _parse_numero_form(item_origem.get("qtd", 1), 1.0)
        if qtd <= 0:
            qtd = 1.0
        descricao = item_info.get("descricao") or item_origem.get("descricao", "")
        categoria = item_info.get("categoria", "") or ""
        fornecedor = item_info.get("fornecedor", "") or ""
        if _eh_faturamento_direto(descricao):
            fornecedor_referencia = _fornecedor_item_referencia(referencia, codigo)
            fornecedor = _fornecedor_faturamento_direto(
                descricao,
                categoria,
                fornecedor_referencia or ("" if _categoria_ar_condicionado(categoria) else "SJ"),
                fornecedor,
            )
            if not fornecedor and not _categoria_ar_condicionado(categoria):
                raise ValueError(f"O.S. {dados_origem.get('os_numero')}: fornecedor ausente no item {codigo}.")
        itens.append({
            "codigo": codigo,
            "descricao": descricao,
            "qtd": qtd,
            "serie": item_origem.get("serie", "") or "",
            "unidade": item_origem.get("unidade", "") or item_info.get("unidade", "") or "",
            "grupo": item_info.get("grupo", "") or "",
            "categoria": categoria,
            "fornecedor": fornecedor,
            "valor": 0,
            "total": calcular_total_item(qtd, 0, 0),
        })

        if codigo.startswith("4034"):
            luminaria_codigo, luminaria_qtd = _selecionar_codigo_relacionado(
                linhas_inferencia,
                codigo,
                OS_LUMINARIAS_CODIGOS,
            )
            if not luminaria_codigo:
                luminaria_historica = (contexto.get("luminarias_por_item") or {}).get(codigo, {})
                luminaria_codigo = normalizar_codigo(luminaria_historica.get("codigo", ""))
                luminaria_qtd = _parse_numero_form(luminaria_historica.get("qtd", 1), 1.0) or 1.0
            if not luminaria_codigo:
                raise ValueError(f"O.S. {dados_origem.get('os_numero')}: luminaria nao identificada para {codigo}.")
            luminaria_info = produtos_catalogo.get(luminaria_codigo, {}) or os_produtos.get(luminaria_codigo, {}) or {}
            extras.append({
                "item": codigo,
                "codigo": luminaria_codigo,
                "descricao": luminaria_info.get("descricao", "") or luminaria_info.get("nome", "") or "",
                "qtd": luminaria_qtd,
                "unidade": luminaria_info.get("unidade", "") or "",
                "level": 0,
            })

        for gatilho in _gatilhos_popup_alcancaveis(codigo, componentes, regras_por_gatilho):
            for regra in regras_por_gatilho.get(gatilho, []):
                relacionado, relacionado_qtd = _selecionar_codigo_relacionado(
                    linhas_inferencia,
                    codigo,
                    regra.get("opcoes") or [],
                )
                if not relacionado:
                    continue
                relacionado_info = produtos_catalogo.get(relacionado, {}) or os_produtos.get(relacionado, {}) or {}
                extras.append({
                    "item": codigo,
                    "codigo": relacionado,
                    "descricao": relacionado_info.get("descricao", "") or relacionado_info.get("nome", "") or "",
                    "qtd": relacionado_qtd or regra.get("quantidade", 1) or 1,
                    "unidade": relacionado_info.get("unidade", "") or "",
                    "level": 0,
                })

    if not itens:
        raise ValueError(f"O.S. {dados_origem.get('os_numero')}: nenhum item valido para salvar.")

    extras_unicos = []
    chaves_extras = set()
    for extra in extras:
        chave = (normalizar_codigo(extra.get("item", "")), normalizar_codigo(extra.get("codigo", "")))
        if chave in chaves_extras:
            continue
        chaves_extras.add(chave)
        extras_unicos.append(extra)

    processos = contexto["processos"]
    processo_por_item = contexto["processo_por_item"]
    codigos_processo = [item["codigo"] for item in itens] + [extra["codigo"] for extra in extras_unicos]
    conjuntos = resolver_processos_transformacao(codigos_processo, processo_por_item)
    processos_final = mesclar_processos_modelo(processos, conjuntos)

    composicao_final = resolver_composicao_final(itens, componentes)
    extras_composicao = expandir_composicao_referenciada(extras_unicos, componentes)
    chaves_composicao = {
        (normalizar_codigo(linha.get("item", "")), normalizar_codigo(linha.get("codigo", "")))
        for linha in composicao_final
    }
    for linha in extras_composicao:
        chave = (normalizar_codigo(linha.get("item", "")), normalizar_codigo(linha.get("codigo", "")))
        if chave not in chaves_composicao:
            composicao_final.append(linha)
            chaves_composicao.add(chave)
    composicao_final = propagar_setor_preparacao(
        enriquecer_composicao(composicao_final, os_produtos),
        os_produtos,
        componentes,
    )

    total_itens = sum(item["total"] for item in itens)
    dados = {
        "cliente": _resolver_nome_cliente_os(dados_origem.get("cliente", "")),
        "previsao_inicio": dados_origem.get("previsao_inicio", "") or "",
        "previsao_termino": dados_origem.get("previsao_termino", "") or "",
        "chassis": dados_origem.get("chassis", "") or "",
        "municipio": dados_origem.get("municipio", "") or "",
        "mmv": dados_origem.get("mmv", "") or "",
        "descricao_servico": dados_origem.get("descricao_servico", "") or "",
        "total_itens": total_itens,
        "total_pedido": total_itens,
        "obs_materiais": dados_origem.get("obs_materiais", "") or "",
        "obs": dados_origem.get("obs", "") or "",
        "processo_conjunto": " + ".join(conjuntos),
        "modo_os": "pacote_os",
    }
    usuario = current_username()
    numero = str(dados_origem.get("os_numero") or "").strip()
    return {
        "tipo": "os",
        "numero": numero,
        "data_criacao": str((referencia or {}).get("data_criacao") or fonte["data_arquivo"].date().isoformat()),
        "status": "emitido",
        "submit_token": f"marco-zero-os-{numero}-{uuid.uuid4().hex}",
        "criado_por": usuario,
        "atualizado_por": usuario,
        "dados": dados,
        "itens": _atribuir_line_ids(itens, "os-item", campos_chave=("codigo", "descricao", "qtd", "unidade", "serie")),
        "processos": _atribuir_line_ids_processos(processos_final),
        "componentes": componentes,
        "composicao": _atribuir_line_ids(
            composicao_final,
            "os-comp",
            campos_chave=("item", "codigo", "descricao", "qtd", "unidade", "level", "setor"),
        ),
    }


def _contexto_reconciliacao_os(documentos_referencia=None):
    os_produtos = carregar_os_produtos()
    componentes = carregar_os_componentes()
    processos = carregar_os_processos()
    relacoes = carregar_relacoes_processo_item()
    processo_por_item = construir_processo_por_item(os_produtos, processos.keys(), relacoes)
    regras_por_gatilho = {}
    for regra in carregar_regras_popup_item():
        regras_por_gatilho.setdefault(regra.get("gatilho", ""), []).append(regra)
    return {
        "os_produtos": os_produtos,
        "produtos_catalogo": carregar_produtos(),
        "componentes": componentes,
        "processos": processos,
        "processo_por_item": processo_por_item,
        "regras_por_gatilho": regras_por_gatilho,
        "luminarias_por_item": _luminarias_historicas_por_item(documentos_referencia),
    }


def _alvo_os_importada(documentos, numero, chassis):
    numero = str(numero or "").strip()
    chassis_norm = str(chassis or "").strip().casefold()
    candidatas = [
        documento for documento in (documentos or [])
        if documento.get("tipo") == "os" and str(documento.get("numero") or "").strip() == numero
    ]
    if chassis_norm:
        candidatas = [
            documento for documento in candidatas
            if str((documento.get("dados") or {}).get("chassis") or "").strip().casefold() == chassis_norm
        ]
    if len(candidatas) != 1:
        raise ValueError(
            f"O.S. {numero}: esperado um registro com chassi {chassis or '-'}, encontrados {len(candidatas)}."
        )
    if candidatas[0].get("id") is None:
        raise ValueError(f"O.S. {numero}: registro sem ID no Supabase.")
    return candidatas[0]


def _assinatura_os_recalculada(documento):
    def normalizar(valor):
        if isinstance(valor, dict):
            return {str(chave): normalizar(conteudo) for chave, conteudo in valor.items()}
        if isinstance(valor, list):
            return [normalizar(conteudo) for conteudo in valor]
        if isinstance(valor, float):
            return int(valor) if valor.is_integer() else float(f"{valor:.12g}")
        return valor

    conteudo = {
        "numero": str(documento.get("numero") or ""),
        "dados": documento.get("dados") or {},
        "itens": documento.get("itens") or [],
        "processos": documento.get("processos") or {},
        "composicao": documento.get("composicao") or [],
    }
    serializado = json.dumps(
        normalizar(conteudo),
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=str,
    )
    return hashlib.sha256(serializado.encode("utf-8")).hexdigest()


def recalcular_os_importadas(arquivos):
    if not supabase_data.enabled():
        raise ValueError("O recalculo seletivo de O.S. exige a base Supabase.")
    fontes = _fontes_os_reconciliacao(arquivos)
    fontes_por_numero, duplicadas_fonte = _parsear_fontes_os_reconciliacao(fontes)
    historico_atual = supabase_data.carregar_documentos(force=True)
    os_atuais = [documento for documento in historico_atual if documento.get("tipo") == "os"]
    referencias = [*historico_atual, *_carregar_historico_local()]
    contexto = _contexto_reconciliacao_os(historico_atual)

    atualizacoes = []
    for numero, fonte in sorted(fontes_por_numero.items(), key=lambda item: item[0]):
        dados_origem = fonte["dados"]
        alvo = _alvo_os_importada(os_atuais, numero, dados_origem.get("chassis", ""))
        referencia = _referencia_os_recente(referencias, numero, dados_origem.get("chassis", "")) or alvo
        corrigida = _montar_os_reconciliada(fonte, referencia, contexto)
        corrigida.update({
            "data_criacao": alvo.get("data_criacao") or corrigida.get("data_criacao"),
            "status": alvo.get("status") or "emitido",
            "submit_token": alvo.get("submit_token") or corrigida.get("submit_token"),
            "criado_por": alvo.get("criado_por") or corrigida.get("criado_por"),
            "atualizado_por": current_username(),
        })
        atualizacoes.append({
            "id": str(alvo["id"]),
            "original": alvo,
            "corrigida": corrigida,
        })

    ids_atualizados = []
    try:
        for atualizacao in atualizacoes:
            confirmado = supabase_data.atualizar_documento(atualizacao["id"], atualizacao["corrigida"])
            if not confirmado:
                raise RuntimeError(f"O Supabase nao confirmou a O.S. {atualizacao['corrigida']['numero']}.")
            ids_atualizados.append(atualizacao["id"])

        final = supabase_data.carregar_documentos(force=True)
        final_por_id = {str(documento.get("id")): documento for documento in final}
        divergentes = [
            atualizacao["corrigida"]["numero"]
            for atualizacao in atualizacoes
            if atualizacao["id"] not in final_por_id
            or _assinatura_os_recalculada(final_por_id[atualizacao["id"]])
            != _assinatura_os_recalculada(atualizacao["corrigida"])
        ]
        if divergentes:
            raise RuntimeError("A verificacao final divergiu nas O.S.: " + ", ".join(divergentes[:10]))
    except Exception:
        for atualizacao in reversed(atualizacoes):
            if atualizacao["id"] not in ids_atualizados:
                continue
            try:
                supabase_data.atualizar_documento(atualizacao["id"], atualizacao["original"])
            except Exception:
                app.logger.exception("Falha ao restaurar a O.S. %s", atualizacao["original"].get("numero"))
        raise

    return {
        "fontes": len(fontes),
        "atualizadas": len(atualizacoes),
        "preservadas": len(os_atuais) - len(atualizacoes),
        "duplicadas_fonte": duplicadas_fonte,
        "itens": sum(len(atualizacao["corrigida"].get("itens") or []) for atualizacao in atualizacoes),
        "componentes": sum(
            len(atualizacao["corrigida"].get("composicao") or []) for atualizacao in atualizacoes
        ),
    }


def adicionar_os_ausentes(arquivos):
    if not supabase_data.enabled():
        raise ValueError("A inclusao seletiva de O.S. exige a base Supabase.")
    fontes = _fontes_os_reconciliacao(arquivos)
    fontes_por_numero, duplicadas_fonte = _parsear_fontes_os_reconciliacao(fontes)
    historico_atual = supabase_data.carregar_documentos(force=True)
    os_atuais = [documento for documento in historico_atual if documento.get("tipo") == "os"]
    numeros_existentes = {
        str(documento.get("numero") or "").strip()
        for documento in os_atuais
    }
    fontes_ausentes = {
        numero: fonte
        for numero, fonte in fontes_por_numero.items()
        if numero not in numeros_existentes
    }
    ignoradas_existentes = len(fontes_por_numero) - len(fontes_ausentes)
    if not fontes_ausentes:
        return {
            "fontes": len(fontes),
            "inseridas": 0,
            "preservadas": len(os_atuais),
            "ignoradas_existentes": ignoradas_existentes,
            "duplicadas_fonte": duplicadas_fonte,
            "numeros": [],
        }

    referencias = [*historico_atual, *_carregar_historico_local()]
    contexto = _contexto_reconciliacao_os(historico_atual)
    novos = []
    for numero, fonte in sorted(fontes_ausentes.items(), key=lambda item: item[0]):
        referencia = _referencia_os_recente(
            referencias,
            numero,
            fonte["dados"].get("chassis", ""),
        )
        novos.append(_montar_os_reconciliada(fonte, referencia, contexto))

    ids_novos = []
    try:
        for inicio in range(0, len(novos), 10):
            lote = novos[inicio:inicio + 10]
            salvos = supabase_data.salvar_documentos(lote)
            ids_novos.extend(
                str(documento["id"])
                for documento in salvos
                if documento.get("id") is not None
            )
            if len(salvos) != len(lote) or any(documento.get("id") is None for documento in salvos):
                raise RuntimeError("O Supabase nao confirmou todas as O.S. ausentes.")

        final = supabase_data.carregar_documentos(force=True)
        final_por_id = {str(documento.get("id")): documento for documento in final}
        esperadas_por_numero = {
            str(documento.get("numero") or "").strip(): documento
            for documento in novos
        }
        divergentes = []
        for documento_id in ids_novos:
            salvo = final_por_id.get(documento_id)
            esperado = esperadas_por_numero.get(str((salvo or {}).get("numero") or "").strip())
            if (
                not salvo
                or not esperado
                or _assinatura_os_recalculada(salvo) != _assinatura_os_recalculada(esperado)
            ):
                divergentes.append(str((salvo or {}).get("numero") or documento_id))
        if len(ids_novos) != len(novos) or divergentes:
            raise RuntimeError(
                "A verificacao final divergiu nas O.S. adicionadas: "
                + ", ".join(divergentes[:10])
            )
    except Exception:
        if ids_novos:
            try:
                supabase_data.excluir_documentos(ids_novos)
            except Exception:
                app.logger.exception("Falha ao reverter O.S. apos erro na inclusao seletiva")
        raise

    return {
        "fontes": len(fontes),
        "inseridas": len(novos),
        "preservadas": len(os_atuais),
        "ignoradas_existentes": ignoradas_existentes,
        "duplicadas_fonte": duplicadas_fonte,
        "numeros": [documento["numero"] for documento in novos],
    }


def reconciliar_os_marco_zero(arquivos):
    if not supabase_data.enabled():
        raise ValueError("A reconciliacao do marco zero exige a base Supabase.")
    fontes = _fontes_os_reconciliacao(arquivos)
    fontes_por_numero, duplicadas_fonte = _parsear_fontes_os_reconciliacao(fontes)
    historico_atual = supabase_data.carregar_documentos(force=True)
    referencias = [*historico_atual, *_carregar_historico_local()]

    contexto = _contexto_reconciliacao_os(historico_atual)

    novos = []
    for numero, fonte in sorted(fontes_por_numero.items(), key=lambda item: item[0]):
        referencia = _referencia_os_recente(referencias, numero, fonte["dados"].get("chassis", ""))
        novos.append(_montar_os_reconciliada(fonte, referencia, contexto))

    ids_antigos = [
        str(documento.get("id")) for documento in historico_atual
        if documento.get("tipo") == "os" and documento.get("id") is not None
    ]
    ids_novos = []
    try:
        for inicio in range(0, len(novos), 10):
            lote = novos[inicio:inicio + 10]
            salvos = supabase_data.salvar_documentos(lote)
            if len(salvos) != len(lote) or any(documento.get("id") is None for documento in salvos):
                raise RuntimeError("O Supabase nao confirmou todas as O.S. do novo marco zero.")
            ids_novos.extend(str(documento["id"]) for documento in salvos)
        supabase_data.excluir_documentos(ids_antigos)
    except Exception:
        if ids_novos:
            try:
                supabase_data.excluir_documentos(ids_novos)
            except Exception:
                app.logger.exception("Falha ao reverter O.S. novas apos erro na reconciliacao")
        raise

    final = supabase_data.carregar_documentos(force=True)
    os_finais = [documento for documento in final if documento.get("tipo") == "os"]
    numeros_finais = {str(documento.get("numero") or "").strip() for documento in os_finais}
    numeros_esperados = set(fontes_por_numero)
    if len(os_finais) != len(numeros_esperados) or numeros_finais != numeros_esperados:
        raise RuntimeError("A verificacao final encontrou divergencia na base de O.S.")
    residuos = len({
        str(documento.get("numero") or "").strip()
        for documento in historico_atual
        if documento.get("tipo") == "os"
    } - numeros_esperados)
    return {
        "fontes": len(fontes),
        "ordens": len(novos),
        "antigas_removidas": len(ids_antigos),
        "duplicadas_fonte": duplicadas_fonte,
        "residuos": residuos,
    }


@app.route("/reconciliar_os_marco_zero", methods=["POST"])
@permission_required("suprimentos.work_order.import")
def reconciliar_os_marco_zero_route():
    try:
        if request.form.get("confirmar_substituicao_os") != "sim":
            raise ValueError("Confirme a substituicao das O.S. antes de sincronizar.")
        arquivos = request.files.getlist("arquivos_os_marco_zero")
        resultado = reconciliar_os_marco_zero(arquivos)
        status = (
            f"Marco zero concluido: {resultado['ordens']} O.S. recalculada(s), "
            f"{resultado['antigas_removidas']} registro(s) antigo(s) removido(s), "
            f"{resultado['residuos']} numero(s) residual(is) eliminado(s)."
        )
        if resultado["duplicadas_fonte"]:
            status += f" {resultado['duplicadas_fonte']} arquivo(s) duplicado(s) descartado(s)."
    except ValueError as exc:
        status = str(exc)
    except Exception:
        app.logger.exception("Falha ao reconciliar marco zero de O.S.")
        status = "Falha ao reconciliar o marco zero de O.S. Consulte o log e o relatorio de backup."
    return redirect(url_for("index", tab="gestao-os", documento_status=status))


@app.route("/recalcular_os_importadas", methods=["POST"])
@permission_required("suprimentos.work_order.import")
def recalcular_os_importadas_route():
    try:
        if request.form.get("confirmar_recalculo_os") != "sim":
            raise ValueError("Confirme o recalculo seletivo das O.S. importadas.")
        arquivos = request.files.getlist("arquivos_os_recalculo")
        resultado = recalcular_os_importadas(arquivos)
        status = (
            f"Recalculo seletivo concluido: {resultado['atualizadas']} O.S. atualizada(s), "
            f"{resultado['preservadas']} O.S. nova(s) preservada(s), "
            f"{resultado['itens']} item(ns) e {resultado['componentes']} componente(s) recalculado(s)."
        )
        if resultado["duplicadas_fonte"]:
            status += f" {resultado['duplicadas_fonte']} arquivo(s) duplicado(s) descartado(s)."
    except ValueError as exc:
        status = str(exc)
    except Exception:
        app.logger.exception("Falha ao recalcular O.S. importadas")
        status = "Falha ao recalcular as O.S. importadas. Nenhuma O.S. nova foi excluida."
    return redirect(url_for("index", tab="gestao-os", documento_status=status))


@app.route("/adicionar_os_ausentes", methods=["POST"])
@permission_required("suprimentos.work_order.import")
def adicionar_os_ausentes_route():
    try:
        if request.form.get("confirmar_inclusao_os") != "sim":
            raise ValueError("Confirme a inclusao seletiva das O.S. ausentes.")
        arquivos = request.files.getlist("arquivos_os_ausentes")
        resultado = adicionar_os_ausentes(arquivos)
        numeros = ", ".join(resultado["numeros"]) or "nenhuma"
        status = (
            f"Inclusao seletiva concluida: {resultado['inseridas']} O.S. adicionada(s) "
            f"({numeros}), {resultado['preservadas']} O.S. existente(s) preservada(s)."
        )
        if resultado["ignoradas_existentes"]:
            status += (
                f" {resultado['ignoradas_existentes']} O.S. ja existente(s) "
                "foi(ram) ignorada(s)."
            )
        if resultado["duplicadas_fonte"]:
            status += f" {resultado['duplicadas_fonte']} arquivo(s) duplicado(s) descartado(s)."
    except ValueError as exc:
        status = str(exc)
    except Exception:
        app.logger.exception("Falha ao adicionar O.S. ausentes")
        status = "Falha ao adicionar as O.S. ausentes. As insercoes foram revertidas."
    return redirect(url_for("index", tab="gestao-os", documento_status=status))


def _bulk_document_required_permission():
    tipo = request.form.get("tipo_baixa_documentos", "")
    return (
        "suprimentos.purchase.bulk_manage"
        if str(tipo).strip().lower() == "oc"
        else "suprimentos.work_order.import"
    )


@app.route("/importar_baixa_documentos", methods=["POST"])
@permission_required(_bulk_document_required_permission)
def importar_baixa_documentos_route():
    arquivo = request.files.get("arquivo_baixa_documentos")
    tipo = request.form.get("tipo_baixa_documentos", "")
    next_tab = str(request.form.get("next_tab") or "dashboard").strip().lower()
    if next_tab not in {"dashboard", "gestao-oc", "gestao-os"}:
        next_tab = "dashboard"
    if not arquivo or not arquivo.filename:
        return redirect(url_for("index", tab=next_tab, documento_status="Selecione uma planilha XLSX ou CSV para baixa em massa."))
    try:
        resultado = importar_baixas_documentos(arquivo, tipo)
        status = (
            f"Baixa em massa: {resultado['atualizados']} documento(s) atualizado(s), "
            f"{resultado['excluidos']} excluido(s), "
            f"{resultado.get('linhas_atualizadas', 0)} linha(s) atualizada(s), "
            f"{resultado.get('linhas_excluidas', 0)} linha(s) excluida(s), "
            f"{resultado['ignorados']} linha(s) ignorada(s)."
        )
        if resultado["erros"]:
            status += " " + "; ".join(resultado["erros"][:3])
            if len(resultado["erros"]) > 3:
                status += f"; mais {len(resultado['erros']) - 3} erro(s)."
    except ValueError as exc:
        status = str(exc)
    except Exception:
        app.logger.exception("Falha ao importar baixa em massa de documentos")
        status = "Falha inesperada ao importar a baixa em massa. Consulte o log."
    return redirect(url_for("index", tab=next_tab, documento_status=status))


@app.route("/healthz")
def healthz():
    produtos_count = None
    if supabase_catalog.enabled():
        try:
            produtos_count = len(supabase_catalog.carregar_produtos())
        except Exception as exc:
            app.logger.exception("Falha no healthz Supabase")
            return {
                "ok": False,
                "catalog": supabase_catalog.status(),
                "error": str(exc),
            }, 500
    rbac_status = supabase_data.shared_rbac_schema_status()
    if rbac_status["enabled"] and not rbac_status["ready"]:
        return {
            "ok": False,
            "catalog": supabase_catalog.status(),
            "data": supabase_data.status(),
            "shared_rbac": {
                "enabled": True,
                "ready": False,
            },
        }, 503
    return {
        "ok": True,
        "catalog": supabase_catalog.status(),
        "data": supabase_data.status(),
        "shared_rbac": {
            "enabled": rbac_status["enabled"],
            "ready": rbac_status["ready"],
        },
        "produtos_count": produtos_count,
    }


def _purchase_save_required_permission():
    historico_id = (request.form.get("oc_historico_id", "") or "").strip()
    if historico_id:
        return "suprimentos.purchase.edit"
    submit_token = (request.form.get("oc_submit_token", "") or "").strip()
    if submit_token and obter_historico_por_submit_token(submit_token):
        return "suprimentos.purchase.edit"
    return "suprimentos.purchase.create"


@app.route("/gerar_oc", methods=["POST"])
@permission_required(_purchase_save_required_permission)
def gerar_oc():

    acao = (request.form.get("acao", "imprimir") or "imprimir").strip().lower()
    historico_id = (request.form.get("oc_historico_id", "") or "").strip()
    submit_token = (request.form.get("oc_submit_token", "") or "").strip()
    historico_existente = obter_historico_documento(historico_id) if historico_id else None
    if historico_id and (not historico_existente or historico_existente.get("tipo") != "oc"):
        return "O.C. indicada para edicao nao foi encontrada.", 404
    atualizar_skus_automatico()
    fornecedor = request.form.get("fornecedor", "")
    fornecedores = carregar_fornecedores()
    fornecedor_info = fornecedores.get(fornecedor, {})

    itens = []

    codigos = request.form.getlist("codigo[]")
    descricoes = request.form.getlist("descricao[]")
    unidades = request.form.getlist("unidade[]")
    qtds = request.form.getlist("qtd[]")
    valores = request.form.getlist("valor[]")
    descontos = request.form.getlist("desconto[]")
    ipis = request.form.getlist("ipi[]")
    icmss = request.form.getlist("icms[]")
    cofins_list = request.form.getlist("cofins[]")
    line_ids = request.form.getlist("oc_line_id[]")

    produtos = carregar_produtos()
    for i in range(len(codigos)):
        codigo_item = normalizar_codigo(codigos[i])
        if not codigo_item:
            continue
        ipi = ipis[i] if i < len(ipis) else ""
        icms = icmss[i] if i < len(icmss) else ""
        cofins = cofins_list[i] if i < len(cofins_list) else ""

        qtd = _parse_numero_form(qtds[i] if i < len(qtds) else "", 0.0)
        valor = _parse_numero_form(valores[i] if i < len(valores) else "", 0.0)
        desconto = _parse_numero_form(descontos[i] if i < len(descontos) else "", 0.0)

        produto_info = produtos.get(codigo_item, {})
        campos_extras = produto_info.get("campos_extras") if isinstance(produto_info.get("campos_extras"), dict) else {}
        desc_form = descricoes[i] if i < len(descricoes) else ""
        unidade_form = unidades[i] if i < len(unidades) else ""
        descricao_final = produto_info.get("descricao") or desc_form
        descricao_secundaria = campos_extras.get("descricao_secundaria") or produto_info.get("descricao_secundaria") or ""
        unidade_final = produto_info.get("unidade") or unidade_form
        ipi_val = ipi if ipi != "" else produto_info.get("ipi")
        icms_val = icms if icms != "" else produto_info.get("icms")
        cofins_val = cofins if cofins != "" else produto_info.get("cofins")

        total = calcular_total_item(qtd, valor, desconto, ipi_val, icms_val, cofins_val)
        itens.append({
            "line_id": line_ids[i].strip() if i < len(line_ids) else "",
            "codigo": codigo_item,
            "descricao": descricao_final,
            "descricao_primaria": descricao_final,
            "descricao_secundaria": descricao_secundaria,
            "unidade": unidade_final,
            "qtd": qtd,
            "valor": valor,
            "desconto": desconto,
            "ipi": ipi_val,
            "icms": icms_val,
            "cofins": cofins_val,
            "total": total
        })

    total_itens = 0
    for item in itens:
        total_itens += item["total"]

    frete_raw = request.form.get("frete", "")
    frete_val = 0
    if frete_raw != "":
        frete_val = _parse_numero_form(frete_raw, 0.0)

    prazo_raw = request.form.get("prazo", "")
    prazo_int = None
    if prazo_raw != "":
        try:
            prazo_int = int(prazo_raw)
        except ValueError:
            prazo_int = None

    vencimento = ""
    if prazo_int is not None:
        vencimento = (date.today() + timedelta(days=prazo_int)).strftime("%d/%m/%Y")

    dados_anteriores = ((historico_existente or {}).get("dados") or {})
    categoria_oc = str(
        request.form.get("oc_categoria")
        or dados_anteriores.get("oc_categoria")
        or "GERAL"
    ).strip().upper()
    destino_form = request.form.get("destino")
    destino_oc = str(
        (destino_form if destino_form is not None else dados_anteriores.get("destino", ""))
        or ""
    ).strip()
    dados_pedido = {
        "cnpj": fornecedor_info.get("cnpj", request.form.get("cnpj", "")),
        "bairro": fornecedor_info.get("bairro", request.form.get("bairro", "")),
        "cidade": fornecedor_info.get("cidade", request.form.get("cidade", "")),
        "uf": fornecedor_info.get("uf", request.form.get("uf", "")),
        "email": fornecedor_info.get("email", request.form.get("email", "")),
        "razao_social": fornecedor_info.get("razao_social", request.form.get("razao_social", "")),
        "endereco": fornecedor_info.get("endereco", request.form.get("endereco", "")),
        "cep": fornecedor_info.get("cep", request.form.get("cep", "")),
        "telefone": fornecedor_info.get("telefone", request.form.get("telefone", "")),
        "previsao": request.form.get("previsao", ""),
        "tipo_frete": request.form.get("tipo_frete", ""),
        "frete": frete_val,
        "total_itens": total_itens,
        "total_pedido": total_itens + frete_val,
        "forma_pagamento": request.form.get("forma_pagamento", ""),
        "prazo": prazo_raw,
        "vencimento": vencimento,
        "obs": request.form.get("obs", ""),
        "oc_categoria": categoria_oc,
        "destino": destino_oc,
    }

    numero_oc = (
        (request.form.get("oc_numero", "") or "").strip()
        or str((historico_existente or {}).get("numero") or "").strip()
        or proximo_numero_oc()
    )

    fornecedor_nome = fornecedor_info.get("fornecedor") or fornecedor_info.get("razao_social") or fornecedor
    oc_mode = request.form.get("oc_mode", "completo")
    incluir_composicao = oc_mode != "resumido"
    componentes = carregar_os_componentes()
    dados_hist = dict(dados_pedido)
    dados_hist["fornecedor"] = fornecedor_nome
    status_documento = (
        ((historico_existente or {}).get("status") or "emitido")
        if historico_existente
        else "rascunho"
    )
    sync_result = None
    if historico_existente and status_documento != "rascunho":
        try:
            sync_result = _sync_emitted_legacy_oc_to_erp(
                historico_existente,
                dados_pedido,
                itens,
                numero_oc,
                fornecedor_nome,
            )
        except Exception as exc:
            app.logger.exception(
                "Falha ao atualizar a O.C. %s no ERP; documento legado nao foi sobrescrito.",
                numero_oc,
            )
            return redirect(url_for(
                "index",
                tab="oc",
                documento_status=(
                    "A compra não foi salva porque a atualização integrada falhou: "
                    f"{exc}"
                ),
            ))
        if sync_result and sync_result.get("locked"):
            return redirect(url_for(
                "index",
                tab="gestao-oc",
                documento_status=(
                    "A compra não foi alterada: já existe recebimento confirmado no Estoque. "
                    "Faça a correção por estorno/ajuste rastreável antes de mudar suas linhas."
                ),
            ))
        if sync_result and sync_result.get("id"):
            dados_hist["erp_purchase_order_id"] = str(sync_result["id"])
    if acao == "salvar":
        historico_salvo = registrar_historico(
            "oc",
            numero_oc,
            dados_hist,
            itens=itens,
            documento_id=historico_id or None,
            status=status_documento,
            submit_token=submit_token,
        )
        if sync_result and sync_result.get("id"):
            vincular_documento_erp(
                historico_salvo,
                "erp_purchase_order_id",
                sync_result["id"],
            )
        limpar_importacao(_user_scoped_file(OC_IMPORT_FILE))
        return redirect(url_for("index", tab="dashboard", documento_status="Compra salva sem impressao."))
    arquivo = gerar_word(
        numero_oc,
        fornecedor_nome,
        dados_pedido,
        itens,
        incluir_composicao=incluir_composicao,
        componentes=componentes,
    )
    nome_docx = construir_nome_oc(numero_oc, fornecedor_nome, dados_pedido)
    historico_emitido = registrar_historico(
        "oc",
        numero_oc,
        dados_hist,
        itens=itens,
        documento_id=historico_id or None,
        status="emitido",
        submit_token=submit_token,
    )
    try:
        if sync_result is None:
            sync_result = _sync_emitted_legacy_oc_to_erp(
                historico_emitido, dados_pedido, itens, numero_oc, fornecedor_nome
            )
        if sync_result and sync_result.get("id"):
            vincular_documento_erp(
                historico_emitido,
                "erp_purchase_order_id",
                sync_result["id"],
            )
        if sync_result and sync_result.get("locked"):
            app.logger.warning("O.C. %s emitida, mas ja possui recebimento e nao foi alterada no ERP.", numero_oc)
    except Exception:
        # Preserve the existing document workflow if the optional integration is
        # unavailable.  The failure is explicit in the application log and no
        # receipt/movement is ever fabricated in Estoque.
        app.logger.exception("Falha ao publicar O.C. %s no ERP", numero_oc)
    limpar_importacao(_user_scoped_file(OC_IMPORT_FILE))
    @after_this_request
    def _cleanup_oc(response):
        try:
            shutil.rmtree(os.path.dirname(arquivo), ignore_errors=True)
        except Exception:
            pass
        return response

    resp = send_file(arquivo, as_attachment=True, download_name=nome_docx)
    return resp



def _parse_os_composition_form(form):
    raw_json = (form.get("os_composicao_json", "") or "").strip()
    if raw_json:
        try:
            source_rows = json.loads(raw_json)
        except json.JSONDecodeError as exc:
            raise ValueError("A composicao da O.S. foi enviada em formato invalido.") from exc
        if not isinstance(source_rows, list):
            raise ValueError("A composicao da O.S. deve ser uma lista de itens.")
    else:
        comp_itens = form.getlist("os_comp_item[]")
        comp_codigos = form.getlist("os_comp_codigo[]")
        comp_descricoes = form.getlist("os_comp_descricao[]")
        comp_unidades = form.getlist("os_comp_unidade[]")
        comp_qtds = form.getlist("os_comp_qtd[]")
        comp_levels = form.getlist("os_comp_level[]")
        comp_setores = form.getlist("os_comp_setor[]")
        comp_setores_manuais = form.getlist("os_comp_setor_manual[]")
        comp_line_ids = form.getlist("os_comp_line_id[]")
        source_rows = []
        for idx in range(len(comp_codigos)):
            source_rows.append(
                {
                    "line_id": comp_line_ids[idx] if idx < len(comp_line_ids) else "",
                    "item": comp_itens[idx] if idx < len(comp_itens) else "",
                    "codigo": comp_codigos[idx],
                    "descricao": comp_descricoes[idx] if idx < len(comp_descricoes) else "",
                    "unidade": comp_unidades[idx] if idx < len(comp_unidades) else "",
                    "qtd": comp_qtds[idx] if idx < len(comp_qtds) else "",
                    "level": comp_levels[idx] if idx < len(comp_levels) else 0,
                    "setor": comp_setores[idx] if idx < len(comp_setores) else "",
                    "setor_manual": comp_setores_manuais[idx] if idx < len(comp_setores_manuais) else False,
                }
            )

    composicao = []
    for raw_row in source_rows:
        if not isinstance(raw_row, dict):
            continue
        item_pai = normalizar_codigo(raw_row.get("item", ""))
        codigo = normalizar_codigo(raw_row.get("codigo", ""))
        descricao = str(raw_row.get("descricao", "") or "").strip()
        unidade = str(raw_row.get("unidade", "") or "").strip()
        qtd = str(raw_row.get("qtd", "") or "").strip()
        try:
            level = int(raw_row.get("level", 0) or 0)
        except (TypeError, ValueError):
            level = 0
        if not (codigo or descricao or qtd or unidade):
            continue
        linha = {
            "line_id": _normalizar_line_id(raw_row.get("line_id")),
            "item": item_pai,
            "codigo": codigo,
            "descricao": descricao,
            "unidade": unidade,
            "qtd": qtd,
            "level": level,
        }
        setor = str(raw_row.get("setor", "") or "").strip().upper()
        setor_manual_raw = raw_row.get("setor_manual", False)
        setor_manual = setor_manual_raw is True or str(setor_manual_raw).strip().lower() in {"1", "true", "sim"}
        if setor in {SETOR_EXPEDICAO, SETOR_PREPARACAO, SETOR_FATURAMENTO_DIRETO}:
            linha["setor"] = setor
            linha["setor_manual"] = setor_manual
        composicao.append(linha)
    return composicao


@app.route("/gerar_os", methods=["POST"])
@permission_required("suprimentos.work_order.manage")
def gerar_os():
    atualizar_skus_automatico()
    acao = (request.form.get("acao", "imprimir") or "imprimir").strip().lower()
    submit_token = (request.form.get("os_submit_token", "") or "").strip()
    composicao_source = (request.form.get("os_composicao_source", "bom") or "bom").strip().lower()
    historico_form_id = (request.form.get("os_historico_id", "") or "").strip()
    historico_existente = obter_historico_documento(historico_form_id) if historico_form_id else None
    if historico_form_id and (not historico_existente or historico_existente.get("tipo") != "os"):
        historico_form_id = ""
        historico_existente = None
    usando_composicao_historica = composicao_source == "custom" and bool(historico_form_id)
    cliente_selecionado = _limpar_valor_busca(
        request.form.get("os_cliente", "") or request.form.get("os_cliente_busca", "")
    )
    cliente = _resolver_nome_cliente_os(cliente_selecionado)
    os_produtos = carregar_os_produtos()
    produtos_catalogo = carregar_produtos()
    bom_dir = get_bom_dir()
    if not supabase_data.enabled() and bom_dir and os.path.isdir(bom_dir):
        resultado_bom = importar_bom_diretorio(bom_dir, somente_se_mais_novo=True)
        if resultado_bom.get("falhas"):
            detalhes = []
            if resultado_bom.get("em_uso"):
                detalhes.append(f"arquivos em uso: {', '.join(resultado_bom['em_uso'][:3])}")
            if resultado_bom.get("com_erro"):
                detalhes.append(f"arquivos com erro: {', '.join(resultado_bom['com_erro'][:3])}")
            resumo = "; ".join(detalhes) or "falha ao ler a base"
            return f"Nao foi possivel atualizar a B.O.M antes de gerar a O.S. ({resumo}).", 400
    itens = []

    codigos = request.form.getlist("os_codigo[]")
    qtds = request.form.getlist("os_qtd[]")
    series = request.form.getlist("os_serie[]")
    unidades = request.form.getlist("os_unidade[]")
    descricoes = request.form.getlist("os_descricao[]")
    fornecedores_linha = request.form.getlist("os_fornecedor_item[]")
    luminarias_linha = request.form.getlist("os_luminaria[]")
    luminarias_qtd_linha = request.form.getlist("os_luminaria_qtd[]")
    popup_itens_linha = request.form.getlist("os_popup_itens[]")
    line_ids = request.form.getlist("os_line_id[]")
    luminarias_extra = []
    popup_itens_extra = []
    regras_popup_por_gatilho = {}
    for regra in carregar_regras_popup_item():
        regras_popup_por_gatilho.setdefault(regra.get("gatilho", ""), []).append(regra)
    componentes = carregar_os_componentes()
    for idx in range(len(codigos)):
        codigo_item = normalizar_codigo(codigos[idx])
        if not codigo_item:
            continue
        qtd_raw = str(qtds[idx]).strip() if idx < len(qtds) else ""
        qtd = _parse_numero_form(qtd_raw, 1.0)
        if qtd <= 0:
            qtd = 1.0

        item_info = os_produtos.get(codigo_item, {})
        descricao_final = item_info.get("descricao") or (descricoes[idx] if idx < len(descricoes) else "")
        unidade_final = unidades[idx] if idx < len(unidades) else item_info.get("unidade", "")
        categoria_final = item_info.get("categoria", "") or ""
        fornecedor_linha = fornecedores_linha[idx].strip() if idx < len(fornecedores_linha) else ""
        fornecedor_final = item_info.get("fornecedor", "") or ""
        if _eh_faturamento_direto(descricao_final):
            fornecedor_final = _fornecedor_faturamento_direto(
                descricao_final,
                categoria_final,
                fornecedor_linha,
                item_info.get("fornecedor", ""),
            )
        itens.append(
            {
                "line_id": line_ids[idx].strip() if idx < len(line_ids) else "",
                "codigo": codigo_item,
                "descricao": descricao_final,
                "qtd": qtd,
                "serie": series[idx] if idx < len(series) else "",
                "unidade": unidade_final or item_info.get("unidade", ""),
                "grupo": item_info.get("grupo", "") or "",
                "categoria": categoria_final,
                "fornecedor": fornecedor_final,
                "valor": 0,
                "total": calcular_total_item(qtd, 0, 0),
            }
        )

        luminaria_codigo = normalizar_codigo(luminarias_linha[idx]) if idx < len(luminarias_linha) else ""
        if luminaria_codigo and luminaria_codigo != POPUP_ITEM_NAO_APLICAVEL:
            luminaria_info = produtos_catalogo.get(luminaria_codigo, {}) or os_produtos.get(luminaria_codigo, {})
            luminaria_qtd_raw = str(luminarias_qtd_linha[idx]).strip() if idx < len(luminarias_qtd_linha) else ""
            luminaria_qtd = _parse_numero_form(luminaria_qtd_raw, 1.0)
            if luminaria_qtd <= 0:
                luminaria_qtd = 1.0
            luminarias_extra.append(
                {
                    "item": codigo_item,
                    "codigo": luminaria_codigo,
                    "descricao": luminaria_info.get("descricao", "") or luminaria_info.get("nome", "") or "",
                    "qtd": luminaria_qtd,
                    "serie": "",
                    "unidade": luminaria_info.get("unidade", "") or "",
                    "grupo": luminaria_info.get("grupo", "") or "",
                    "categoria": luminaria_info.get("categoria", "") or "",
                    "fornecedor": luminaria_info.get("fornecedor", "") or "",
                    "valor": 0,
                    "total": calcular_total_item(qtd, 0, 0),
                    "level": 0,
                }
            )

        popup_json = popup_itens_linha[idx] if idx < len(popup_itens_linha) else "[]"
        try:
            popup_selecoes = json.loads(popup_json or "[]")
        except Exception:
            popup_selecoes = []
        if usando_composicao_historica:
            popup_selecoes_validas = popup_selecoes if isinstance(popup_selecoes, list) else []
        else:
            popup_selecoes_validas, popup_erro = _resolver_selecoes_popup_item(
                codigo_item,
                popup_selecoes,
                regras_popup_por_gatilho,
                componentes,
            )
            if popup_erro:
                return popup_erro, 400
        for selecao in popup_selecoes_validas:
            if not isinstance(selecao, dict):
                continue
            relacionado_codigo = normalizar_codigo(selecao.get("codigo", ""))
            if not relacionado_codigo or relacionado_codigo == POPUP_ITEM_NAO_APLICAVEL:
                continue
            relacionado_qtd = _parse_numero_form(selecao.get("qtd", 1), 1.0)
            if relacionado_qtd <= 0:
                relacionado_qtd = 1
            relacionado_info = produtos_catalogo.get(relacionado_codigo, {}) or os_produtos.get(relacionado_codigo, {})
            popup_itens_extra.append(
                {
                    "item": codigo_item,
                    "codigo": relacionado_codigo,
                    "descricao": relacionado_info.get("descricao", "") or relacionado_info.get("nome", "") or "",
                    "qtd": relacionado_qtd,
                    "serie": "",
                    "unidade": relacionado_info.get("unidade", "") or "",
                    "grupo": relacionado_info.get("grupo", "") or "",
                    "categoria": relacionado_info.get("categoria", "") or "",
                    "fornecedor": relacionado_info.get("fornecedor", "") or "",
                    "valor": 0,
                    "total": calcular_total_item(relacionado_qtd, 0, 0),
                    "level": 0,
                }
            )

    total_itens = sum(item["total"] for item in itens)

    processos = carregar_os_processos()
    relacoes_processo_item = carregar_relacoes_processo_item()
    processo_por_item = construir_processo_por_item(
        os_produtos,
        processos.keys(),
        relacoes_processo_item,
    )
    conjunto_processo_form = (request.form.get("os_processo_conjunto", "") or "").strip()
    codigos_para_processo = [
        item.get("codigo", "")
        for item in [*itens, *luminarias_extra, *popup_itens_extra]
        if item.get("codigo", "")
    ]
    conjuntos_processo_disponiveis = resolver_processos_transformacao(
        codigos_para_processo,
        processo_por_item,
    )
    if conjunto_processo_form and conjunto_processo_form in processos:
        conjuntos_processo_disponiveis.append(conjunto_processo_form)
    conjuntos_processo_disponiveis = list(dict.fromkeys(conjuntos_processo_disponiveis))
    conjunto_processo = " + ".join(conjuntos_processo_disponiveis)
    processos_modelo = mesclar_processos_modelo(processos, conjuntos_processo_disponiveis)

    dados = {
        "cliente": cliente,
        "previsao_inicio": request.form.get("os_previsao_inicio", ""),
        "previsao_termino": request.form.get("os_previsao_termino", ""),
        "chassis": request.form.get("os_chassis", ""),
        "municipio": request.form.get("os_municipio", ""),
        "mmv": request.form.get("os_mmv", ""),
        "descricao_servico": request.form.get("os_descricao_servico", ""),
        "total_itens": total_itens,
        "total_pedido": total_itens,
        "obs_materiais": request.form.get("os_obs_materiais", ""),
        "obs": request.form.get("os_obs", ""),
        "processo_conjunto": conjunto_processo,
    }

    processos_final = {nome: [] for nome in PROCESSOS_ORDEM}
    algum_processo_informado = False
    for processo in PROCESSOS_OS:
        nome = processo["nome"]
        key = processo["key"]
        atividades = request.form.getlist(f"proc_{key}_atividade[]")
        responsaveis = request.form.getlist(f"proc_{key}_responsavel[]")
        line_ids_processo = request.form.getlist(f"proc_{key}_line_id[]")
        linhas = []
        for idx in range(len(atividades)):
            atividade = atividades[idx].strip()
            responsavel = responsaveis[idx].strip() if idx < len(responsaveis) else ""
            if atividade:
                linhas.append({
                    "line_id": line_ids_processo[idx].strip() if idx < len(line_ids_processo) else "",
                    "atividade": atividade,
                    "responsavel": responsavel,
                })
        if linhas:
            algum_processo_informado = True
        processos_final[nome] = linhas

    if not algum_processo_informado and any(processos_modelo.values()):
        processos_final = processos_modelo

    numero_manual = request.form.get("os_numero", "").strip()
    numero_os = numero_manual or str((historico_existente or {}).get("numero") or "").strip() or proximo_numero_os()

    layout_pdf = request.files.get("os_layout_pdf")
    composicao_importada = _parse_os_composition_form(request.form)

    if composicao_source == "custom":
        composicao_final = composicao_importada
    else:
        composicao_final = resolver_composicao_final(itens, componentes, composicao_importada or None)
    extras_composicao = (
        []
        if composicao_source == "custom"
        else expandir_composicao_referenciada(
            [*luminarias_extra, *popup_itens_extra],
            componentes,
        )
    )
    if extras_composicao:
        existentes = {
            (normalizar_codigo(linha.get("item", "")), normalizar_codigo(linha.get("codigo", "")))
            for linha in composicao_final
        }
        for extra in extras_composicao:
            chave = (normalizar_codigo(extra.get("item", "")), normalizar_codigo(extra.get("codigo", "")))
            if chave in existentes:
                continue
            composicao_final.append(extra)
            existentes.add(chave)
    composicao_enriquecida = propagar_setor_preparacao(
        enriquecer_composicao(composicao_final, os_produtos),
        os_produtos,
        componentes,
    )
    pendencias_faturamento_direto = filtrar_linhas_faturamento_direto(composicao_enriquecida)
    pendencias_expedicao = filtrar_linhas_setor(composicao_enriquecida, SETOR_EXPEDICAO)
    pendencias_preparacao = filtrar_linhas_preparacao(composicao_enriquecida)
    chaves_preparacao = {
        (
            normalizar_codigo(linha.get("codigo", "")),
            normalizar_codigo(linha.get("item", "")),
            str(linha.get("qtd", "")),
        )
        for linha in pendencias_preparacao
    }
    for linha in linhas_layout_preparacao(itens, os_produtos):
        chave = (
            normalizar_codigo(linha.get("codigo", "")),
            normalizar_codigo(linha.get("item", "")),
            str(linha.get("qtd", "")),
        )
        if chave not in chaves_preparacao:
            pendencias_preparacao.append(linha)
            chaves_preparacao.add(chave)
    requisicao_materiais = [*pendencias_expedicao, *pendencias_preparacao, *pendencias_faturamento_direto]

    itens_expedicao = construir_itens_os_expedicao(pendencias_expedicao)
    itens_preparacao = construir_itens_os_preparacao(pendencias_preparacao)
    for item in itens_expedicao + itens_preparacao:
        item["qtd"] = _formatar_qtd_saida(item.get("qtd", ""))

    itens_por_modo = {
        "originais": itens,
        "expedicao": itens_expedicao,
        "preparacao": itens_preparacao,
    }
    modos_pacote = ["completa", "expedicao", "preparacao"]
    layout_bytes = b""
    if layout_pdf and layout_pdf.filename:
        try:
            layout_pdf.stream.seek(0)
        except Exception:
            pass
        layout_bytes = layout_pdf.read() or b""

    def _criar_layout_clone():
        if not layout_bytes:
            return None
        return FileStorage(
            stream=io.BytesIO(layout_bytes),
            filename=layout_pdf.filename,
            content_type=layout_pdf.content_type,
        )

    dados_historico = dict(dados)
    dados_historico["modo_os"] = "pacote_os"
    status_documento = (historico_existente or {}).get("status") or "emitido"
    if acao == "salvar" and not historico_existente:
        status_documento = "rascunho"
    if acao == "salvar":
        registrar_historico(
            "os",
            numero_os,
            dados_historico,
            itens=itens,
            processos=processos_final,
            componentes=componentes,
            composicao=composicao_enriquecida,
            documento_id=historico_form_id or None,
            status=status_documento,
            submit_token=submit_token,
        )
        limpar_importacao(_user_scoped_file(OS_IMPORT_FILE))
        return redirect(url_for("index", tab="dashboard", documento_status="O.S. salva sem impressao."))

    arquivos_docx = []
    for modo_key in modos_pacote:
        config_modo = OS_MODE_CONFIGS[modo_key]
        itens_saida = itens_por_modo.get(config_modo["itens_source"], itens)
        layout_clone = _criar_layout_clone()
        arquivos_docx.append(
            gerar_os_docx(
                numero_os,
                dados,
                itens_saida,
                componentes,
                processos_final,
                layout_clone,
                composicao_final or None,
                modo=config_modo["doc_mode"],
                titulo_arquivo=config_modo["titulo"],
                incluir_cliente_nome=config_modo.get("incluir_cliente_nome", True),
                cliente_nome_limite=config_modo.get("cliente_nome_limite"),
            )
        )

    for fornecedor, linhas_fornecedor in agrupar_linhas_por_fornecedor(pendencias_faturamento_direto):
        itens_fornecedor = construir_itens_os_setor(agrupar_linhas_setor(linhas_fornecedor))
        for item in itens_fornecedor:
            item["qtd"] = _formatar_qtd_saida(item.get("qtd", ""))
        fornecedor_titulo = _sanitize_output_name(fornecedor) or "SEM FORNECEDOR"
        dados_requisicao = dict(dados)
        dados_requisicao["fornecedor_requisicao"] = fornecedor
        arquivos_docx.append(
            gerar_os_docx(
                numero_os,
                dados_requisicao,
                itens_fornecedor,
                componentes,
                processos_final,
                _criar_layout_clone(),
                composicao_resolvida=[],
                modo="faturamento_direto",
                titulo_arquivo=f"FATURAMENTO DIRETO - {fornecedor_titulo}",
                incluir_cliente_nome=False,
            )
        )
    arquivo_requisicao_materiais = _criar_planilha_requisicao_materiais(
        numero_os,
        dados,
        requisicao_materiais,
        "03 - Requisicao de Materiais",
    )

    chassi_nome = (dados.get("chassis", "") or "").strip()
    registrar_historico(
        "os",
        numero_os,
        dados_historico,
        itens=itens,
        processos=processos_final,
        componentes=componentes,
        composicao=composicao_enriquecida,
        documento_id=historico_form_id or None,
        status="emitido",
        submit_token=submit_token,
    )
    limpar_importacao(_user_scoped_file(OS_IMPORT_FILE))

    arquivos_saida = [
        *arquivos_docx,
        arquivo_requisicao_materiais,
    ]
    nome_zip = f"02 - Pacote O.S - {chassi_nome}.zip".strip(" -")
    zip_path, download_name = _criar_zip_temporario(arquivos_saida, nome_zip)

    @after_this_request
    def _cleanup_os_zip(response):
        try:
            os.remove(zip_path)
        except Exception:
            pass
        for path in arquivos_saida:
            try:
                shutil.rmtree(os.path.dirname(path), ignore_errors=True)
            except Exception:
                pass
        return response

    resp = send_file(zip_path, as_attachment=True, download_name=download_name)
    return resp


@app.route("/salvar_caminhos", methods=["POST"])
@permission_required("suprimentos.system.admin")
def salvar_caminhos():
    pedidos_dir = request.form.get("pedidos_dir", "")
    os_dir = request.form.get("os_dir", "")
    bom_dir = request.form.get("bom_dir", "")
    skus_file = request.form.get("skus_file", "")
    processos_dir = request.form.get("processos_dir", "")
    set_save_paths(pedidos_dir, os_dir)
    set_bom_dir(bom_dir)
    set_skus_file(skus_file)
    set_processos_dir(processos_dir)
    return redirect(url_for("index", tab="cadastro"))


@app.route("/cadastrar_fornecedor", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def cadastrar_fornecedor():

    fornecedores = carregar_fornecedores()

    fornecedor = request.form.get("fornecedor", "").strip()
    cnpj = request.form.get("cnpj", "").strip()

    chave = cnpj if cnpj else fornecedor
    if chave:
        atual = fornecedores.get(chave, {})
        def _pick(campo, valor):
            return valor if valor != "" else atual.get(campo, "")
        fornecedores[chave] = {
            "fornecedor": _pick("fornecedor", fornecedor),
            "razao_social": _pick("razao_social", request.form.get("razao_social", "").strip()),
            "cnpj": _pick("cnpj", cnpj),
            "email": _pick("email", request.form.get("email", "").strip()),
            "telefone": _pick("telefone", request.form.get("telefone", "").strip()),
            "endereco": _pick("endereco", request.form.get("endereco", "").strip()),
            "bairro": _pick("bairro", request.form.get("bairro", "").strip()),
            "cidade": _pick("cidade", request.form.get("cidade", "").strip()),
            "uf": _pick("uf", request.form.get("uf", "").strip()),
            "cep": _pick("cep", request.form.get("cep", "").strip()),
        }

        salvar_fornecedores(fornecedores)

    return index()


@app.route("/cadastrar_pessoa", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def cadastrar_pessoa():
    pessoa = {
        "identificador": request.form.get("identificador", "").strip(),
        "pessoa_fisica": request.form.get("pessoa_fisica") == "1",
        "nome_fantasia": request.form.get("nome_fantasia", "").strip(),
        "razao_social": request.form.get("razao_social", "").strip(),
        "cnpj_cpf": request.form.get("cnpj_cpf", "").strip(),
        "rg": request.form.get("rg", "").strip(),
        "ie": request.form.get("ie", "").strip(),
        "logradouro": request.form.get("logradouro", "").strip(),
        "logradouro_numero": request.form.get("logradouro_numero", "").strip(),
        "complemento": request.form.get("complemento", "").strip(),
        "bairro": request.form.get("bairro", "").strip(),
        "cidade": request.form.get("cidade", "").strip(),
        "uf": request.form.get("uf", "").strip(),
        "cep": request.form.get("cep", "").strip(),
        "telefone": request.form.get("telefone", "").strip(),
        "whatsapp": request.form.get("whatsapp", "").strip(),
        "celular": request.form.get("celular", "").strip(),
        "email": request.form.get("email", "").strip(),
        "site": request.form.get("site", "").strip(),
        "cliente": request.form.get("cliente") == "1",
        "fornecedor": request.form.get("fornecedor") == "1",
        "colaborador": request.form.get("colaborador") == "1",
        "transportadora": request.form.get("transportadora") == "1",
        "pessoa_grupo": request.form.get("pessoa_grupo", "").strip(),
        "categoria": request.form.get("categoria", "").strip(),
        "observacoes": request.form.get("observacoes", "").strip(),
    }
    if any(pessoa.get(campo) for campo in ("nome_fantasia", "razao_social", "cnpj_cpf", "identificador")):
        if supabase_data.enabled():
            supabase_data.salvar_pessoas([pessoa])
        else:
            nome = pessoa.get("nome_fantasia") or pessoa.get("razao_social") or pessoa.get("cnpj_cpf")
            if pessoa.get("fornecedor"):
                fornecedores = carregar_fornecedores()
                fornecedores[pessoa.get("cnpj_cpf") or nome] = {
                    "fornecedor": nome,
                    "razao_social": pessoa.get("razao_social", ""),
                    "cnpj": pessoa.get("cnpj_cpf", ""),
                    "email": pessoa.get("email", ""),
                    "telefone": pessoa.get("telefone", ""),
                    "endereco": pessoa.get("logradouro", ""),
                    "bairro": pessoa.get("bairro", ""),
                    "cidade": pessoa.get("cidade", ""),
                    "uf": pessoa.get("uf", ""),
                    "cep": pessoa.get("cep", ""),
                }
                salvar_fornecedores(fornecedores)
            if pessoa.get("cliente"):
                clientes = carregar_os_fornecedores()
                clientes[nome] = {"cliente": nome}
                salvar_os_fornecedores(clientes)
    return redirect(url_for("index", tab="cadastro"))


@app.route("/cadastrar_item", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def cadastrar_item():

    produtos = carregar_produtos()

    codigo = request.form.get("codigo", "").strip()
    if codigo:
        atual = produtos.get(codigo, {})
        novos = {
            "descricao": request.form.get("descricao", "").strip(),
            "unidade": request.form.get("unidade", "").strip(),
            "grupo": request.form.get("grupo", "").strip(),
            "categoria": request.form.get("categoria", "").strip(),
            "processo_conjunto": request.form.get("processo_conjunto", "").strip(),
        }
        produtos[codigo] = _mesclar_dados_item(atual, novos)

        salvar_json(PRODUTOS_FILE, produtos)

    return index()


@app.route("/cadastrar_os_fornecedor", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def cadastrar_os_fornecedor():

    fornecedores = carregar_os_fornecedores()

    cliente = request.form.get("os_fornecedor", "").strip()
    if cliente:
        atual = fornecedores.get(cliente, {})
        fornecedores[cliente] = {"cliente": cliente if cliente != "" else atual.get("cliente", "")}

        salvar_os_fornecedores(fornecedores)

    return index()


@app.route("/cadastrar_os_item", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def cadastrar_os_item():

    produtos = carregar_os_produtos()
    componentes = carregar_os_componentes()

    codigo = normalizar_codigo(request.form.get("os_item_codigo", ""))
    descricao = request.form.get("os_item_descricao", "").strip()
    unidade = request.form.get("os_item_unidade", "").strip()
    grupo = request.form.get("os_item_grupo", "").strip()
    categoria = request.form.get("os_item_categoria", "").strip()
    fornecedor = request.form.get("os_item_fornecedor", "").strip()
    processo_conjunto = request.form.get("os_item_processo_conjunto", "").strip()

    comp_codigos = request.form.getlist("os_comp_codigo[]")
    comp_descricoes = request.form.getlist("os_comp_descricao[]")
    comp_unidades = request.form.getlist("os_comp_unidade[]")
    comp_qtds = request.form.getlist("os_comp_qtd[]")

    comps = []
    for i in range(len(comp_codigos)):
        if not comp_codigos[i] and not comp_descricoes[i]:
            continue
        comps.append({
            "codigo": normalizar_codigo(comp_codigos[i]),
            "descricao": comp_descricoes[i],
            "unidade": comp_unidades[i],
            "quantidade": comp_qtds[i],
        })

    if codigo and descricao:
        atual_prod = produtos.get(codigo, {})
        produtos[codigo] = _mesclar_dados_item(
            atual_prod,
            {
                "descricao": descricao,
                "unidade": (unidade or "").strip() or atual_prod.get("unidade", "UN") or "UN",
                "unidade_comercial": (unidade or "").strip() or atual_prod.get("unidade_comercial", "") or atual_prod.get("unidade", "UN") or "UN",
                "unidade_interna": atual_prod.get("unidade_interna", ""),
                "grupo": grupo,
                "categoria": categoria,
                "tipo": atual_prod.get("tipo", ""),
                "fornecedor": fornecedor,
                "ncm": "",
                "origem": "",
                "valor": "",
                "ipi": "",
                "icms": "",
                "cofins": "",
                "observacao": "",
                "processo_conjunto": processo_conjunto,
            },
        )
        if comps:
            componentes[codigo] = comps
        elif codigo in componentes:
            componentes[codigo] = componentes.get(codigo, [])
        else:
            componentes[codigo] = []

        salvar_json(OS_PRODUTOS_FILE, produtos)
        salvar_json(OS_COMPONENTES_FILE, componentes)

    return index()


@app.route("/cadastrar_relacao_processo_item", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def cadastrar_relacao_processo_item():
    relacoes = carregar_relacoes_processo_item()
    codigo = normalizar_codigo(request.form.get("relacao_item", ""))
    processos = [
        processo.strip()
        for processo in request.form.getlist("relacao_processo_conjunto[]")
        if processo.strip()
    ]
    if not processos:
        processo_legado = request.form.get("relacao_processo_conjunto", "").strip()
        processos = [processo_legado] if processo_legado else []
    if codigo and processos:
        relacoes[codigo] = list(dict.fromkeys(processos))
        salvar_relacoes_processo_item(relacoes)
    return redirect(url_for("index", tab="cadastro"))


@app.route("/excluir_relacao_processo_item", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def excluir_relacao_processo_item():
    relacoes = carregar_relacoes_processo_item()
    codigo = normalizar_codigo(request.form.get("relacao_item", ""))
    if codigo and codigo in relacoes:
        relacoes.pop(codigo, None)
        salvar_relacoes_processo_item(relacoes)
    return redirect(url_for("index", tab="cadastro"))


@app.route("/cadastrar_regra_popup_item", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def cadastrar_regra_popup_item():
    regras = carregar_regras_popup_item()
    gatilho = normalizar_codigo(request.form.get("popup_regra_gatilho", ""))
    opcoes = [
        normalizar_codigo(codigo)
        for codigo in request.form.getlist("popup_regra_opcao[]")
        if normalizar_codigo(codigo)
    ]
    try:
        quantidade = float(request.form.get("popup_regra_quantidade", "1") or 1)
    except Exception:
        quantidade = 1
    if quantidade <= 0:
        quantidade = 1
    if gatilho and opcoes:
        proximo_id = max(
            [
                int(re.sub(r"\D", "", str(regra.get("id", ""))) or 0)
                for regra in regras
            ]
            or [0]
        ) + 1
        regras.append(
            {
                "id": f"regra-{proximo_id}",
                "gatilho": gatilho,
                "opcoes": list(dict.fromkeys(opcoes)),
                "quantidade": quantidade,
                "quantidade_editavel": request.form.get("popup_regra_quantidade_editavel") == "1",
            }
        )
        salvar_regras_popup_item(regras)
    return redirect(url_for("index", tab="cadastro"))


@app.route("/excluir_regra_popup_item", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def excluir_regra_popup_item():
    regra_id = str(request.form.get("popup_regra_id", "") or "").strip()
    if regra_id:
        regras = [
            regra
            for regra in carregar_regras_popup_item()
            if str(regra.get("id", "") or "").strip() != regra_id
        ]
        salvar_regras_popup_item(regras)
    return redirect(url_for("index", tab="cadastro"))


def _status_importacao_planilha(resultado, nome):
    status = (
        f"{nome}: {resultado['novas']} nova(s), "
        f"{resultado['atualizadas']} atualizada(s) e "
        f"{resultado['ignoradas']} ignorada(s)."
    )
    erros = resultado.get("erros") or []
    if erros:
        status += " " + "; ".join(erros[:3])
        if len(erros) > 3:
            status += f"; mais {len(erros) - 3} erro(s)."
    return status


@app.route("/importar_regras_popup_item", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def importar_regras_popup_item_route():
    arquivo = request.files.get("arquivo_regras_popup_item")
    if not arquivo or not arquivo.filename:
        status = "Selecione uma planilha XLSX de parametros de item relacionado."
    else:
        try:
            resultado = importar_regras_popup_item_planilha(arquivo)
            status = _status_importacao_planilha(resultado, "Parametros importados")
        except ValueError as exc:
            status = f"Falha na importacao: {exc}"
        except Exception:
            app.logger.exception("Falha ao importar parametros de item relacionado")
            status = "Falha inesperada ao importar os parametros. Consulte o log."
    destino = url_for("index", tab="cadastro", regras_popup_status=status)
    return redirect(f"{destino}#parametros-item-relacionado")


@app.route("/importar_relacoes_processo_item", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def importar_relacoes_processo_item_route():
    arquivo = request.files.get("arquivo_relacoes_processo_item")
    if not arquivo or not arquivo.filename:
        status = "Selecione uma planilha XLSX de relacao processo x item."
    else:
        try:
            resultado = importar_relacoes_processo_item_planilha(arquivo)
            status = _status_importacao_planilha(resultado, "Relacoes importadas")
        except ValueError as exc:
            status = f"Falha na importacao: {exc}"
        except Exception:
            app.logger.exception("Falha ao importar relacoes processo x item")
            status = "Falha inesperada ao importar as relacoes. Consulte o log."
    destino = url_for("index", tab="cadastro", relacoes_processo_status=status)
    return redirect(f"{destino}#relacao-processo-item")


@app.route("/importar_produtos", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def importar_produtos_route():

    arquivo = request.files.get("arquivo_produtos")
    if arquivo and arquivo.filename:
        importar_produtos(arquivo)

    return redirect(url_for("index", tab="oc"))


@app.route("/importar_fornecedores", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def importar_fornecedores_route():

    arquivo = request.files.get("arquivo_fornecedores")
    if arquivo and arquivo.filename:
        importar_fornecedores(arquivo)

    return redirect(url_for("index", tab="oc"))


@app.route("/importar_pessoas", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def importar_pessoas_route():
    arquivo = request.files.get("arquivo_pessoas")
    status = ""
    if arquivo and arquivo.filename:
        try:
            count = importar_pessoas(arquivo)
            status = f"Pessoas importadas: {count} registro(s)."
        except Exception:
            app.logger.exception("Falha ao importar cadastro geral de pessoas")
            status = "Falha ao importar cadastro geral de pessoas. Consulte o log."
    return redirect(url_for("index", tab="cadastro", pessoas_status=status))


@app.route("/importar_os_produtos", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def importar_os_produtos_route():

    arquivo = request.files.get("arquivo_os_produtos")
    if arquivo and arquivo.filename:
        importar_os_produtos(arquivo)

    return redirect(url_for("index", tab="os"))


@app.route("/importar_os_fornecedores", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def importar_os_fornecedores_route():

    arquivo = request.files.get("arquivo_os_fornecedores")
    if arquivo and arquivo.filename:
        importar_os_fornecedores(arquivo)

    return redirect(url_for("index", tab="os"))


@app.route("/importar_os_componentes", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def importar_os_componentes_route():
    if supabase_data.enabled():
        status = "B.O.M. deve ser importada no ModuloCadastro; o suprimentos apenas consome a base Supabase."
        return redirect(url_for("index", tab="cadastro", bom_status=status))

    arquivo = request.files.get("arquivo_os_componentes")
    if arquivo and arquivo.filename:
        importar_os_componentes(arquivo)

    return redirect(url_for("index", tab="os"))


@app.route("/atualizar_bom", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def atualizar_bom():
    if supabase_data.enabled():
        status = "B.O.M. lida diretamente do Supabase pelo ModuloCadastro."
        return redirect(url_for("index", tab="cadastro", bom_status=status))
    bom_dir = get_bom_dir()
    if not bom_dir:
        status = "Caminho da B.O.M. não foi configurado."
        return redirect(url_for("index", tab="cadastro", bom_status=status))
    if not os.path.isdir(bom_dir):
        status = f"Caminho da B.O.M. inválido: {bom_dir}"
        return redirect(url_for("index", tab="cadastro", bom_status=status))

    arquivos = listar_arquivos_excel_base(bom_dir)

    if not arquivos:
        status = f"Nenhuma planilha Excel (.xls/.xlsx/.xlsm/.xltx/.xltm) encontrada em {bom_dir}"
        return redirect(url_for("index", tab="cadastro", bom_status=status))

    resultado = importar_bom_diretorio(bom_dir)
    arquivos_processados = resultado["arquivos"]
    linhas_importadas = resultado["linhas"]
    falhas = resultado["falhas"]
    arquivos_em_uso = resultado["em_uso"]
    arquivos_com_erro = resultado["com_erro"]

    if arquivos_processados == 0:
        status = f"Nenhum arquivo válido encontrado em {bom_dir}"
    elif falhas:
        detalhes = []
        if arquivos_em_uso:
            detalhes.append(f"{len(arquivos_em_uso)} em uso no Excel")
        if arquivos_com_erro:
            detalhes.append(f"{len(arquivos_com_erro)} com erro de leitura")
        resumo_falhas = ", ".join(detalhes) if detalhes else f"{falhas} falhas"
        status = (
            f"{arquivos_processados} arquivos processados, "
            f"{linhas_importadas} linhas importadas e {resumo_falhas}. Veja o log."
        )
    else:
        status = f"{arquivos_processados} arquivos importados ({linhas_importadas} linhas)."

    return redirect(url_for("index", tab="cadastro", bom_status=status))


@app.route("/atualizar_skus", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def atualizar_skus():
    if supabase_catalog.enabled():
        try:
            supabase_catalog.clear_cache()
            produtos = supabase_catalog.carregar_produtos(force=True)
            status = f"Catalogo Supabase recarregado: {len(produtos)} SKU(s)."
        except Exception as exc:
            app.logger.exception("Falha ao recarregar catalogo Supabase")
            status = f"Falha ao recarregar catalogo Supabase: {exc}"
        return redirect(url_for("index", tab="cadastro", skus_status=status))

    skus_file = get_skus_file()
    resultado = atualizar_skus_arquivo(skus_file, somente_se_mais_novo=False)
    if resultado.get("erro"):
        status = resultado["erro"]
    elif resultado.get("linhas", 0) == 0:
        status = f"Nenhum SKU importado de {skus_file}."
    else:
        status = f"SKUs atualizados: {resultado['linhas']} linha(s) importada(s) de {skus_file}."
    return redirect(url_for("index", tab="cadastro", skus_status=status))


@app.route("/atualizar_processos", methods=["POST"])
@permission_required("suprimentos.master_data.manage")
def atualizar_processos():
    tab_destino = (request.form.get("next_tab", "") or "").strip() or "cadastro"
    if supabase_data.enabled():
        status = "Processos persistidos no Supabase. Use importacao/cadastro de processos para alterar a base."
        return redirect(url_for("index", tab=tab_destino, os_processos_status=status))
    processos_dir = get_processos_dir()
    if not processos_dir:
        status = "Caminho da base de processos não foi configurado."
        return redirect(url_for("index", tab=tab_destino, os_processos_status=status))
    if not os.path.isdir(processos_dir):
        status = f"Caminho da base de processos inválido: {processos_dir}"
        return redirect(url_for("index", tab=tab_destino, os_processos_status=status))

    excel_exts = {".xls", ".xlsx", ".xlsm", ".xltx", ".xltm"}
    arquivos = []
    for root, _, nomes in os.walk(processos_dir):
        for nome in nomes:
            if os.path.splitext(nome)[1].lower() in excel_exts:
                arquivos.append(os.path.join(root, nome))

    if not arquivos:
        status = f"Nenhuma planilha Excel (.xls/.xlsx/.xlsm/.xltx/.xltm) encontrada em {processos_dir}"
        return redirect(url_for("index", tab=tab_destino, os_processos_status=status))

    arquivos_processados = 0
    linhas_importadas = 0
    falhas = 0
    arquivos_em_uso = []
    arquivos_com_erro = []

    for caminho in sorted(arquivos):
        try:
            with _open_for_read(caminho) as f:
                storage = FileStorage(stream=f, filename=os.path.basename(caminho))
                linhas_importadas += importar_os_processos_atualizado(storage)
            arquivos_processados += 1
            app.logger.info("Importada base de processos %s", caminho)
        except PermissionError:
            arquivos_em_uso.append(os.path.basename(caminho))
            falhas += 1
        except Exception:
            app.logger.exception("Falha ao importar base de processos %s", caminho)
            arquivos_com_erro.append(os.path.basename(caminho))
            falhas += 1

    if arquivos_processados == 0:
        status = f"Nenhum arquivo válido encontrado em {processos_dir}"
    elif falhas:
        detalhes = []
        if arquivos_em_uso:
            detalhes.append(f"{len(arquivos_em_uso)} em uso no Excel")
        if arquivos_com_erro:
            detalhes.append(f"{len(arquivos_com_erro)} com erro de leitura")
        resumo_falhas = ", ".join(detalhes) if detalhes else f"{falhas} falhas"
        status = (
            f"{arquivos_processados} arquivos processados, "
            f"{linhas_importadas} linhas importadas e {resumo_falhas}. Veja o log."
        )
    else:
        status = f"{arquivos_processados} arquivos de processos importados ({linhas_importadas} linhas)."

    return redirect(url_for("index", tab=tab_destino, os_processos_status=status))


@app.route("/importar_os_processos", methods=["POST"])
@permission_required("suprimentos.work_order.import")
def importar_os_processos_route():

    tab_destino = (request.form.get("next_tab", "") or "").strip() or "os"
    arquivos = [arquivo for arquivo in request.files.getlist("arquivo_os_processos") if arquivo and arquivo.filename]
    if not arquivos:
        status = "Nenhum arquivo de processo foi selecionado."
        return redirect(url_for("index", tab=tab_destino, os_processos_status=status))

    linhas_importadas = 0
    arquivos_importados = 0
    for arquivo in arquivos:
        linhas = importar_os_processos_atualizado(arquivo)
        linhas_importadas += linhas
        arquivos_importados += 1

    status = f"{arquivos_importados} arquivo(s) de processo importado(s) ({linhas_importadas} linhas)."
    return redirect(url_for("index", tab=tab_destino, os_processos_status=status))


@app.route("/importar_oc_documento", methods=["POST"])
@permission_required("suprimentos.purchase.create")
def importar_oc_documento():
    arquivo = request.files.get("arquivo_oc_template")
    if arquivo and arquivo.filename:
        ext = os.path.splitext(secure_filename(arquivo.filename))[1].lower()
        if ext == ".docx":
            data = parse_oc_docx(arquivo)
        elif ext == ".pdf":
            data = parse_oc_pdf(arquivo)
        else:
            data = {}
        if data:
            data["itens"] = _atribuir_line_ids(
                data.get("itens") or [],
                "oc-item",
                campos_chave=("codigo", "descricao", "qtd", "unidade"),
            )
            salvar_json(_user_scoped_file(OC_IMPORT_FILE), data)
    return redirect(url_for("index", tab="oc"))


@app.route("/importar_os_documento", methods=["POST"])
@permission_required("suprimentos.work_order.manage")
def importar_os_documento():
    arquivo = request.files.get("arquivo_os_template")
    if arquivo and arquivo.filename:
        ext = os.path.splitext(secure_filename(arquivo.filename))[1].lower()
        if ext == ".docx":
            data = parse_os_docx_atualizado(arquivo)
        elif ext == ".pdf":
            data = parse_os_pdf(arquivo)
        else:
            data = {}
        if data:
            # A importacao de O.S serve como base de preenchimento. A composicao
            # deve ser recalculada pela B.O.M atual no Supabase para evitar
            # duplicidade, dados antigos e payloads grandes demais.
            data["itens"] = _atribuir_line_ids(
                data.get("itens") or [],
                "os-item",
                campos_chave=("codigo", "descricao", "qtd", "unidade", "serie"),
            )
            data["processos"] = _atribuir_line_ids_processos(data.get("processos") or {})
            data["composicao"] = []
            salvar_json(_user_scoped_file(OS_IMPORT_FILE), data)
    return redirect(url_for("index", tab="os"))


@app.route("/exportar_modelo_produtos")
@permission_required("suprimentos.master_data.manage")
def exportar_modelo_produtos():
    path, nome = _criar_modelo_xlsx(
        MODELO_ITENS_HEADERS,
        "modelo_produtos.xlsx",
        header_row=2,
    )
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_modelo_fornecedores")
@permission_required("suprimentos.master_data.manage")
def exportar_modelo_fornecedores():
    path, nome = _criar_modelo_xlsx(
        ["fornecedor", "razao_social", "cnpj", "email", "telefone", "endereco", "bairro", "cidade", "uf", "cep"],
        "modelo_fornecedores.xlsx",
    )
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_modelo_os_clientes")
@permission_required("suprimentos.master_data.manage")
def exportar_modelo_os_clientes():
    path, nome = _criar_modelo_xlsx(["cliente"], "modelo_os_clientes.xlsx")
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_modelo_os_itens")
@permission_required("suprimentos.master_data.manage")
def exportar_modelo_os_itens():
    path, nome = _criar_modelo_xlsx(
        MODELO_ITENS_HEADERS,
        "modelo_os_itens.xlsx",
        header_row=2,
    )
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_modelo_os_componentes")
@permission_required("suprimentos.master_data.manage")
def exportar_modelo_os_componentes():
    path, nome = _criar_modelo_xlsx(
        ["item_codigo", "componente_codigo", "descricao", "unidade", "quantidade"],
        "modelo_os_componentes.xlsx",
    )
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_modelo_os_processos")
@permission_required("suprimentos.work_order.import")
def exportar_modelo_os_processos():
    path, nome = _criar_modelo_os_processos_xlsx()
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_modelo_regras_popup_item")
@permission_required("suprimentos.master_data.manage")
def exportar_modelo_regras_popup_item():
    path, nome = _criar_planilha_regras_popup_item()
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_regras_popup_item")
@permission_required("suprimentos.master_data.manage")
def exportar_regras_popup_item():
    path, nome = _criar_planilha_regras_popup_item(
        carregar_regras_popup_item(),
        nome_arquivo="parametros_item_relacionado_atuais.xlsx",
    )
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_modelo_relacoes_processo_item")
@permission_required("suprimentos.master_data.manage")
def exportar_modelo_relacoes_processo_item():
    path, nome = _criar_planilha_relacoes_processo_item()
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/exportar_relacoes_processo_item")
@permission_required("suprimentos.master_data.manage")
def exportar_relacoes_processo_item():
    path, nome = _criar_planilha_relacoes_processo_item(
        carregar_relacoes_processo_item(),
        nome_arquivo="relacao_processo_item_atual.xlsx",
    )
    return send_file(path, as_attachment=True, download_name=nome)


@app.route("/gerar_zip_release", methods=["POST"])
@permission_required("suprimentos.system.admin")
def gerar_zip_release():
    ok, err = _run_release_build()
    if not ok:
        return err, 400
    result = _build_release_zip()
    if not result:
        return "Nao encontrei build em dist. Gere o EXE antes e tente novamente.", 400
    zip_path, temp_root = result

    @after_this_request
    def _cleanup(response):
        try:
            shutil.rmtree(temp_root, ignore_errors=True)
        except Exception:
            pass
        return response

    return send_file(zip_path, as_attachment=True, download_name=RELEASE_ZIP_NAME)


@app.route("/resetar_base", methods=["POST"])
@permission_required("suprimentos.system.admin")
def resetar_base():
    return "Reset da base de dados desativado.", 403


@app.route("/exportar_dashboard", methods=["GET"])
def exportar_dashboard():
    tipo_filtro = (request.args.get("tipo", "") or "").strip().lower()
    if tipo_filtro not in {"", "oc", "os"}:
        return "Tipo de relatorio invalido.", 400
    required_permissions = []
    if tipo_filtro in {"", "oc"}:
        required_permissions.append("suprimentos.purchase.export")
    if tipo_filtro in {"", "os"}:
        required_permissions.append("suprimentos.work_order.view")
    for permission in required_permissions:
        if not can(permission):
            return _authorization_denied(permission)
    historico = [
        entry for entry in carregar_historico()
        if not tipo_filtro or entry.get("tipo") == tipo_filtro
    ]
    wb = Workbook()
    wb.remove(wb.active)
    base_headers = ["ID", "Status", "Data Criacao", "Numero", "Criado Por", "Atualizado Por"]
    ws_oc = ws_oc_itens = ws_os = ws_os_itens = ws_os_proc = ws_os_comp = None
    if tipo_filtro in {"", "oc"}:
        ws_oc = wb.create_sheet("Compras")
        ws_oc.append(base_headers + [
            "Fornecedor", "Razao Social", "CNPJ", "Email", "Telefone", "Endereco", "Bairro",
            "Cidade", "UF", "CEP", "Previsao", "Tipo Frete", "Frete", "Total Itens",
            "Total Pedido", "Forma Pagamento", "Prazo", "Vencimento", "Observacoes", "ACAO"
        ])
        ws_oc_itens = wb.create_sheet("Compras Itens")
        ws_oc_itens.append(base_headers + [
            "ID Linha", "Indice", "Codigo", "Descricao", "Unidade", "Qtd", "Valor", "Desconto",
            "IPI", "ICMS", "COFINS", "Total", "Status Linha", "ACAO"
        ])
    if tipo_filtro in {"", "os"}:
        ws_os = wb.create_sheet("Ordens de Servico")
        ws_os.append(base_headers + [
            "Cliente", "Chassis", "Municipio", "MMV", "Previsao Inicio", "Previsao Termino",
            "Descricao Servico", "Processo Vinculado", "Observacoes Materiais", "Observacoes", "ACAO"
        ])
        ws_os_itens = wb.create_sheet("OS Itens")
        ws_os_itens.append(base_headers + [
            "ID Linha", "Indice", "Codigo", "Descricao", "Quantidade", "Serie", "Unidade", "Grupo",
            "Categoria", "Fornecedor", "Status Linha", "ACAO"
        ])
        ws_os_proc = wb.create_sheet("OS Processos")
        ws_os_proc.append(base_headers + ["ID Linha", "Grupo", "Indice", "Atividade", "Responsavel", "Status Linha", "ACAO"])
        ws_os_comp = wb.create_sheet("OS Componentes")
        ws_os_comp.append(base_headers + [
            "ID Linha", "Indice", "Item Pai", "Codigo", "Descricao", "Unidade", "Quantidade", "Nivel",
            "Destino", "Destino Manual", "Status Linha", "ACAO"
        ])

    for entry in historico:
        tipo = entry.get("tipo")
        data = entry.get("data_criacao", "")
        numero = entry.get("numero", "")
        dados = entry.get("dados", {}) or {}
        itens = entry.get("itens", []) or []
        base = [
            entry.get("id", ""),
            entry.get("status", "emitido"),
            data,
            numero,
            entry.get("criado_por", ""),
            entry.get("atualizado_por", ""),
        ]
        if tipo == "oc":
            ws_oc.append(base + [
                dados.get("fornecedor", ""),
                dados.get("razao_social", ""),
                dados.get("cnpj", ""),
                dados.get("email", ""),
                dados.get("telefone", ""),
                dados.get("endereco", ""),
                dados.get("bairro", ""),
                dados.get("cidade", ""),
                dados.get("uf", ""),
                dados.get("cep", ""),
                dados.get("previsao", ""),
                dados.get("tipo_frete", ""),
                dados.get("frete", ""),
                dados.get("total_itens", ""),
                dados.get("total_pedido", ""),
                dados.get("forma_pagamento", ""),
                dados.get("prazo", ""),
                dados.get("vencimento", ""),
                dados.get("obs", ""),
                "",
            ])
            for idx, item in enumerate(itens, start=1):
                ws_oc_itens.append(base + [
                    item.get("line_id", ""),
                    idx,
                    item.get("codigo", ""),
                    item.get("descricao", ""),
                    item.get("unidade", ""),
                    item.get("qtd", ""),
                    item.get("valor", ""),
                    item.get("desconto", ""),
                    item.get("ipi", ""),
                    item.get("icms", ""),
                    item.get("cofins", ""),
                    item.get("total", ""),
                    _linha_status(item),
                    "",
                ])
        elif tipo == "os":
            ws_os.append(base + [
                dados.get("cliente", ""),
                dados.get("chassis", ""),
                dados.get("municipio", ""),
                dados.get("mmv", ""),
                dados.get("previsao_inicio", ""),
                dados.get("previsao_termino", ""),
                dados.get("descricao_servico", ""),
                dados.get("processo_conjunto", ""),
                dados.get("obs_materiais", ""),
                dados.get("obs", ""),
                "",
            ])
            for idx, item in enumerate(itens, start=1):
                ws_os_itens.append(base + [
                    item.get("line_id", ""),
                    idx,
                    item.get("codigo", ""),
                    item.get("descricao", ""),
                    item.get("qtd", ""),
                    item.get("serie", ""),
                    item.get("unidade", ""),
                    item.get("grupo", ""),
                    item.get("categoria", ""),
                    item.get("fornecedor", ""),
                    _linha_status(item),
                    "",
                ])
            processos = entry.get("processos", {}) or {}
            for grupo, linhas in processos.items():
                for idx, linha in enumerate(linhas, start=1):
                    ws_os_proc.append(base + [
                        linha.get("line_id", ""),
                        grupo,
                        idx,
                        linha.get("atividade", ""),
                        linha.get("responsavel", ""),
                        _linha_status(linha),
                        "",
                    ])
            composicao = entry.get("composicao", []) or []
            for idx, comp in enumerate(composicao, start=1):
                ws_os_comp.append(base + [
                    comp.get("line_id", ""),
                    idx,
                    comp.get("item", ""),
                    comp.get("codigo", ""),
                    comp.get("descricao", ""),
                    comp.get("unidade", ""),
                    comp.get("qtd", ""),
                    comp.get("level", ""),
                    comp.get("setor", ""),
                    "Sim" if comp.get("setor_manual") else "Nao",
                    _linha_status(comp),
                    "",
                ])

    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="0B1C3A")
            cell.alignment = Alignment(horizontal="center")
        for col_idx, column_cells in enumerate(ws.columns, start=1):
            width = max((len(str(cell.value or "")) for cell in column_cells), default=0) + 2
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max(width, 10), 45)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    tmp.close()
    wb.save(tmp.name)
    nome = {"oc": "relatorio_compras", "os": "relatorio_ordens_servico"}.get(tipo_filtro, "relatorio_suprimentos")
    @after_this_request
    def _cleanup_relatorio(response):
        try:
            os.remove(tmp.name)
        except OSError:
            pass
        return response
    return send_file(tmp.name, as_attachment=True, download_name=f"{nome}_{date.today().isoformat()}.xlsx")


def _erp_stock_request(path, method="GET", payload=None):
    base = os.environ.get("ERP_STOCK_API_URL", "").rstrip("/")
    token = os.environ.get("ERP_BACKEND_TOKEN", "")
    if not base or not token:
        raise ValueError("Integracao ERP nao configurada. Defina ERP_STOCK_API_URL e ERP_BACKEND_TOKEN no backend.")
    body = json.dumps(payload).encode("utf-8") if payload is not None else None
    req = urllib.request.Request(
        f"{base}/api/erp/internal/{path.lstrip('/')}", data=body, method=method,
        headers={
            "Content-Type": "application/json",
            "X-ERP-Backend-Token": token,
            "X-ERP-Actor": current_username(),
            "X-ERP-Actor-ID": current_user_id(),
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=25) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        data = json.loads(exc.read().decode("utf-8") or "{}")
        raise ValueError(data.get("error") or "Falha no Estoque.")


def _erp_stock_binary_request(path):
    base = os.environ.get("ERP_STOCK_API_URL", "").rstrip("/")
    token = os.environ.get("ERP_BACKEND_TOKEN", "")
    if not base or not token:
        raise ValueError("Integracao ERP nao configurada para exportacao.")
    req = urllib.request.Request(
        f"{base}/api/erp/internal/{path.lstrip('/')}",
        method="GET",
        headers={
            "X-ERP-Backend-Token": token,
            "X-ERP-Actor": current_username(),
            "X-ERP-Actor-ID": current_user_id(),
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as response:
            return response.read()
    except urllib.error.HTTPError as exc:
        try:
            data = json.loads(exc.read().decode("utf-8") or "{}")
        except (UnicodeDecodeError, json.JSONDecodeError):
            data = {}
        raise ValueError(data.get("error") or "Falha ao gerar relatório no Estoque.") from exc
    except urllib.error.URLError as exc:
        raise ValueError(f"Estoque indisponível: {exc.reason}") from exc


def _erp_mes_request(path, method="GET", payload=None):
    base = os.environ.get("ERP_MES_API_URL", "").rstrip("/")
    token = os.environ.get("ERP_BACKEND_TOKEN", "")
    if not base or not token:
        raise ValueError("Integração MES não configurada. Defina ERP_MES_API_URL e ERP_BACKEND_TOKEN no backend.")
    timeout_seconds = _positive_env_int("ERP_MES_API_TIMEOUT_SECONDS", 75)
    body = json.dumps(payload).encode("utf-8") if payload is not None else None
    req = urllib.request.Request(
        f"{base}/api/erp/internal/{path.lstrip('/')}",
        data=body,
        method=method,
        headers={
            "Content-Type": "application/json",
            "X-ERP-Backend-Token": token,
            "X-ERP-Actor": current_username(),
            "X-ERP-Actor-ID": current_user_id(),
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout_seconds) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        data = json.loads(exc.read().decode("utf-8") or "{}")
        raise ValueError(data.get("error") or "Falha no MES.")
    except TimeoutError as exc:
        raise ValueError(
            f"O MES demorou mais de {timeout_seconds} segundos para responder. "
            "Aguarde a inicialização e tente novamente."
        ) from exc
    except urllib.error.URLError as exc:
        if isinstance(exc.reason, TimeoutError) or "timed out" in str(exc.reason).lower():
            raise ValueError(
                f"O MES demorou mais de {timeout_seconds} segundos para responder. "
                "Aguarde a inicialização e tente novamente."
            ) from exc
        raise ValueError(f"MES indisponível: {exc.reason}") from exc


def _resolve_linked_legacy_os_work_id(documento):
    """Resolve and persist the structural MES link without guessing duplicates."""
    direct_id = str(
        (documento or {}).get("erp_work_order_id")
        or ((documento or {}).get("dados") or {}).get("erp_work_order_id")
        or ""
    ).strip()
    if direct_id:
        return direct_id

    numero = str((documento or {}).get("numero") or "").strip()
    if not numero:
        return None
    orders = _erp_mes_request("work-orders").get("orders", [])
    matches = [
        item for item in orders
        if str(item.get("numero_os") or "").strip() == numero and item.get("work_order_id")
    ]
    if len(matches) > 1:
        raise ValueError(
            f"A O.S. {numero} possui mais de um vínculo possível no MES. "
            "Associe o UUID antes da conclusão técnica."
        )
    if not matches:
        return None
    work_id = str(matches[0]["work_order_id"])
    vincular_documento_erp(documento, "erp_work_order_id", work_id)
    return work_id


def _close_linked_legacy_os_in_mes(documento, motivo):
    work_id = _resolve_linked_legacy_os_work_id(documento)
    if not work_id:
        # Documentos antigos sem uma O.S. MES correspondente continuam
        # consultáveis, mas não podem arquivar um registro inexistente.
        return None
    return _erp_mes_request(
        f"work-orders/{work_id}/technical-close",
        "POST",
        {"motivo": str(motivo or "")},
    )


def _reopen_linked_legacy_os_in_mes(documento, motivo):
    work_id = _resolve_linked_legacy_os_work_id(documento)
    if not work_id:
        return None
    return _erp_mes_request(
        f"work-orders/{work_id}/technical-reopen",
        "POST",
        {"motivo": str(motivo or "")},
    )


def _erp_iso_date(value):
    """Normalize legacy form dates before sending them to the ERP service."""
    if isinstance(value, (date, datetime)):
        return value.date().isoformat() if isinstance(value, datetime) else value.isoformat()
    raw = str(value or "").strip()
    if not raw:
        return None
    for pattern in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(raw, pattern).date().isoformat()
        except ValueError:
            pass
    return None


def _sync_emitted_legacy_oc_to_erp(historico, dados_pedido, itens, numero_oc, fornecedor_nome):
    """Publish the existing full Suprimentos O.C. as the single receipt source.

    The legacy history id is deliberately used as a stable idempotency key.  A
    reprint or edit before receiving updates the same ERP O.C.; it never creates
    a duplicate pending receipt in Estoque.
    """
    if not erp_feature_enabled():
        return None
    history_id = str((historico or {}).get("id") or "").strip()
    if not history_id:
        raise ValueError("A O.C. nao recebeu identificador de historico para sincronizacao.")
    payload = {
        "numero_oc": numero_oc,
        "categoria": (dados_pedido.get("oc_categoria") or "GERAL").strip().upper(),
        "fornecedor_nome": fornecedor_nome,
        "data_emissao": _erp_iso_date((historico or {}).get("data_criacao")) or date.today().isoformat(),
        "data_necessidade": _erp_iso_date(dados_pedido.get("previsao")),
        "destino": (dados_pedido.get("destino") or "").strip(),
        "frete": dados_pedido.get("frete", 0),
        "observacoes": dados_pedido.get("obs", ""),
        "idempotency_key": f"suprimentos-oc:{history_id}",
        "lines": [
            {
                "sku_codigo": item.get("codigo"),
                "descricao_original": item.get("descricao"),
                "unidade": item.get("unidade") or "UN",
                "quantidade_pedida": item.get("qtd"),
                "valor_unitario_pedido": item.get("valor"),
                "destino": (dados_pedido.get("destino") or "").strip(),
                "data_necessidade": _erp_iso_date(dados_pedido.get("previsao")),
            }
            for item in itens
        ],
    }
    return _erp_stock_request("purchase-orders/legacy-sync", "POST", payload)


def _cancel_emitted_legacy_oc_in_erp(historico, motivo):
    if str((historico or {}).get("status") or "").lower() == "rascunho":
        return None
    history_id = str((historico or {}).get("id") or "").strip()
    if not history_id:
        raise ValueError("A O.C. nao possui identificador para cancelamento integrado.")
    return _erp_stock_request("purchase-orders/legacy-cancel", "POST", {
        "idempotency_key": f"suprimentos-oc:{history_id}", "motivo": motivo or "",
    })


def _close_emitted_legacy_oc_in_erp(historico, motivo):
    if str((historico or {}).get("status") or "").lower() == "rascunho":
        return None
    history_id = str((historico or {}).get("id") or "").strip()
    if not history_id:
        raise ValueError("A O.C. nao possui identificador para conclusao integrada.")
    return _erp_stock_request("purchase-orders/legacy-close", "POST", {
        "idempotency_key": f"suprimentos-oc:{history_id}", "motivo": motivo or "",
    })


@app.route("/erp/ordens-compra")
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.view")
def erp_purchase_orders_screen():
    """Operational monitoring; O.C. creation remains in the complete legacy form."""
    stock_public_url = (
        os.environ.get("ERP_STOCK_PUBLIC_URL")
        or os.environ.get("ERP_STOCK_API_URL")
        or "http://127.0.0.1:5000"
    ).rstrip("/")
    return render_template(
        "erp_ordens_compra.html",
        current_user=current_user(),
        stock_public_url=stock_public_url,
    )


@app.route("/erp/relatorios/compras-inspecao.xlsx")
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.export")
def erp_purchase_inspection_report():
    try:
        content = _erp_stock_binary_request("reports/purchases-inspections.xlsx")
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400
    return send_file(
        io.BytesIO(content),
        as_attachment=True,
        download_name=f"Compras_Bancos_e_Inspecao_{date.today().isoformat()}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/api/erp/purchase-orders", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.create")
def erp_purchase_order_proxy():
    try:
        result = _erp_stock_request("purchase-orders", "POST", request.get_json(silent=True) or {})
        return jsonify(result), 201 if not result.get("replayed") else 200
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/receipts/pending")
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.view")
def erp_pending_receipts_proxy():
    try: return jsonify(_erp_stock_request("receipts/pending"))
    except ValueError as exc: return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/dashboard")
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.view")
def erp_dashboard_proxy():
    try: return jsonify(_erp_stock_request("dashboard"))
    except ValueError as exc: return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/purchase-orders/<order_id>/technical-close", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.technical_close")
def erp_purchase_order_technical_close_proxy(order_id):
    try:
        return jsonify(_erp_stock_request(
            f"purchase-orders/{order_id}/technical-close",
            "POST",
            request.get_json(silent=True) or {},
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/purchase-orders/<order_id>/correct-number", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.edit")
def erp_purchase_order_correct_number_proxy(order_id):
    """Controlled correction of the visible O.C. number.

    The stock service keeps the UUID, receipts and movements intact.  It also
    updates the linked buyer document only when its immutable ERP UUID matches.
    """
    try:
        return jsonify(_erp_stock_request(
            f"purchase-orders/{order_id}/correct-number",
            "POST",
            request.get_json(silent=True) or {},
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/purchase-orders/<order_id>/financial-close", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.financial_close")
def erp_purchase_order_financial_close_proxy(order_id):
    try:
        return jsonify(_erp_stock_request(
            f"purchase-orders/{order_id}/financial-close",
            "POST",
            request.get_json(silent=True) or {},
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/purchase-orders/<order_id>/financial-detail")
@login_required
@erp_feature_required
@permission_required("suprimentos.purchase.view")
def erp_purchase_order_financial_detail_proxy(order_id):
    try:
        return jsonify(_erp_stock_request(
            f"purchase-orders/{order_id}/financial-detail",
            "GET",
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/erp/gestao-os")
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.view")
def erp_work_order_management_screen():
    return render_template(
        "erp_gestao_os.html",
        mes_url=os.environ.get("ERP_MES_PUBLIC_URL") or os.environ.get("ERP_MES_API_URL", "http://127.0.0.1:8010"),
        current_user=current_user(),
    )


@app.route("/erp/ordens-producao")
@login_required
@erp_feature_required
@permission_required("suprimentos.production_order.view")
def erp_production_orders_screen():
    return render_template("erp_ordens_producao.html", current_user=current_user())


@app.route("/api/erp/production-orders")
@login_required
@erp_feature_required
@permission_required("suprimentos.production_order.view")
def erp_production_orders_proxy():
    try:
        return jsonify(_erp_stock_request("production-orders", "GET"))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/production-orders/catalog/<sku>")
@login_required
@erp_feature_required
@permission_required("suprimentos.production_order.view")
def erp_production_order_catalog_proxy(sku):
    try:
        return jsonify({"ok": True, "catalog": supabase_catalog.registration_by_sku(sku)})
    except supabase_catalog.SupabaseCatalogError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 404


@app.route("/api/erp/production-orders", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.production_order.manage")
def erp_production_order_create_proxy():
    try:
        result = _erp_stock_request("production-orders", "POST", request.get_json(silent=True) or {})
        return jsonify(result), 200 if result.get("replayed") else 201
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/production-orders/<order_id>")
@login_required
@erp_feature_required
@permission_required("suprimentos.production_order.view")
def erp_production_order_detail_proxy(order_id):
    try:
        return jsonify(_erp_stock_request(f"production-orders/{order_id}", "GET"))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/production-orders/<order_id>/<action>", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.production_order.execute")
def erp_production_order_action_proxy(order_id, action):
    if action not in {"commit", "complete", "cancel"}:
        return jsonify({"ok": False, "error": "Ação de O.P. inválida."}), 404
    try:
        return jsonify(_erp_stock_request(
            f"production-orders/{order_id}/{action}",
            "POST",
            request.get_json(silent=True) or {},
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/erp/ordens-producao/<order_id>/documento.docx")
@login_required
@erp_feature_required
@permission_required("suprimentos.production_order.view")
def erp_production_order_document(order_id):
    try:
        order = _erp_stock_request(f"production-orders/{order_id}", "GET").get("order")
        if not order:
            raise ValueError("O.P. não encontrada.")
        return send_file(
            build_production_order_docx(order),
            as_attachment=True,
            download_name=f"{order.get('numero_op') or 'ordem_producao'}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.route("/api/erp/os-management")
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.view")
def erp_work_order_management_list():
    try:
        return jsonify(_erp_mes_request("work-orders"))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.route("/api/erp/os-management/catalogs")
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.view")
def erp_work_order_catalogs():
    try:
        return jsonify(_erp_mes_request("catalogs"))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.route("/api/erp/os-management/<work_id>")
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.view")
def erp_work_order_management_detail(work_id):
    try:
        return jsonify(_erp_mes_request(f"work-orders/{work_id}"))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/os-management/work-orders/<work_id>/materials")
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.view")
def erp_work_order_materials_proxy(work_id):
    try:
        return jsonify(_erp_stock_request(f"work-orders/{work_id}/materials"))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.route("/api/erp/os-management/entries", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.manage")
def erp_vehicle_entry_proxy():
    try:
        return jsonify(_erp_mes_request("vehicle-entries", "POST", request.get_json(silent=True) or {})), 201
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.route("/api/erp/os-management/entries/<entry_id>/work-orders", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.manage")
def erp_work_order_create_proxy(entry_id):
    try:
        return jsonify(_erp_mes_request(
            f"vehicle-entries/{entry_id}/work-orders", "POST", request.get_json(silent=True) or {}
        )), 201
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.route("/api/erp/os-management/work-orders/<work_id>", methods=["PUT"])
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.manage")
def erp_work_order_update_proxy(work_id):
    try:
        return jsonify(_erp_mes_request(
            f"work-orders/{work_id}", "PUT", request.get_json(silent=True) or {}
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.route("/api/erp/os-management/work-orders/<work_id>/activate", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.manage")
def erp_work_order_activate_proxy(work_id):
    try:
        return jsonify(_erp_mes_request(f"work-orders/{work_id}/activate", "POST", {}))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.route("/api/erp/os-management/work-orders/<work_id>/technical-close", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.technical_close")
def erp_work_order_technical_close_proxy(work_id):
    try:
        return jsonify(_erp_mes_request(
            f"work-orders/{work_id}/technical-close",
            "POST",
            request.get_json(silent=True) or {},
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/os-management/work-orders/<work_id>/technical-reopen", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.technical_close")
def erp_work_order_technical_reopen_proxy(work_id):
    try:
        return jsonify(_erp_mes_request(
            f"work-orders/{work_id}/technical-reopen",
            "POST",
            request.get_json(silent=True) or {},
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@app.route("/api/erp/os-management/work-orders/<work_id>/schedules", methods=["POST"])
@login_required
@erp_feature_required
@permission_required("suprimentos.work_order.schedule")
def erp_work_order_schedule_proxy(work_id):
    try:
        return jsonify(_erp_mes_request(
            f"work-orders/{work_id}/schedules", "POST", request.get_json(silent=True) or {}
        ))
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


if __name__ == "__main__":
    port = int(os.environ.get("PORT", _get_free_port()))
    _abrir_navegador(port)
    app.run(debug=False, use_reloader=False, port=port)
