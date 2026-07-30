"""DOCX generator for the Serralheria production order.

The document is generated from the immutable snapshots persisted in the O.P.,
not from a later live Cadastro lookup. This keeps the printed instruction and
the stock conversion auditable even if a SKU parameter is edited afterwards.
"""

from io import BytesIO
from decimal import Decimal, InvalidOperation

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
LIGHT = "D9EAF7"


def _quantity(value):
    """Formats a stored 3-decimal stock quantity for people, not for math."""
    try:
        normalized = Decimal(str(value or "0").replace(",", "."))
    except (InvalidOperation, ValueError):
        return str(value or "")
    rendered = format(normalized, "f").rstrip("0").rstrip(".")
    return rendered or "0"


def _cell(cell, text, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)


def _table_header(table, labels):
    for cell, label in zip(table.rows[0].cells, labels):
        _cell(cell, label, bold=True, fill=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)


def _title(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def _section(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def _format_parameters(order):
    rows = order.get("selected_parameters") or []
    if not rows:
        return []
    result = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        label = str(row.get("label") or row.get("key") or "Parâmetro")
        value = str(row.get("value") or "")
        if value:
            result.append((label, value))
    return result


def build_production_order_docx(order):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    _title(document, "ORDEM DE FABRICAÇÃO - SERRALHERIA")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Nº {order.get('numero_op', '')}  |  Status: {order.get('status', '')}")
    run.bold = True
    run.font.size = Pt(9)

    _section(document, "DADOS DA ORDEM DE FABRICAÇÃO")
    info = document.add_table(rows=3, cols=4)
    info.style = "Table Grid"
    values = [
        ("Tipo", "Produção para estoque" if order.get("producao_tipo") == "ESTOQUE" else "Produção destinada"),
        ("Setor interno", order.get("setor") or "SERRALHERIA"),
        ("Destino", order.get("destino_descricao") or "-"),
        ("Quantidade", f"{_quantity(order.get('quantidade_planejada'))} {order.get('unidade', '')}"),
        ("Chassis / lote", order.get("chassi_lote") or "-"),
        ("Cliente", order.get("cliente_nome") or "-"),
        ("Município", order.get("municipio") or "-"),
        ("MMV", order.get("mmv") or "-"),
        ("Data de criação", (order.get("created_at") or "").replace("T", " ")[:16]),
        ("Conclusão", (order.get("completed_at") or "-").replace("T", " ")[:16]),
        ("Responsável", "____________________________"),
        ("Qualidade", "____________________________"),
    ]
    for index, (label, value) in enumerate(values):
        row = info.rows[index // 4]
        cell = row.cells[index % 4]
        _cell(cell, f"{label}:\n{value}", fill=LIGHT if index % 2 == 0 else None)

    _section(document, "PRODUTO FINAL (SKU DESTINO)")
    output = document.add_table(rows=2, cols=4)
    output.style = "Table Grid"
    _table_header(output, ["CÓDIGO", "PRODUTO FINAL", "QTDE", "UN."])
    _cell(output.rows[1].cells[0], order.get("target_sku"))
    _cell(output.rows[1].cells[1], order.get("target_descricao"))
    _cell(output.rows[1].cells[2], _quantity(order.get("quantidade_planejada")), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(output.rows[1].cells[3], order.get("unidade"), align=WD_ALIGN_PARAGRAPH.CENTER)

    parameters = _format_parameters(order)
    if parameters:
        _section(document, "PARÂMETROS EXPLÍCITOS DO PRODUTO FINAL")
        parameters_table = document.add_table(rows=1, cols=2)
        parameters_table.style = "Table Grid"
        _table_header(parameters_table, ["CAMPO ALTERADO", "ESPECIFICAÇÃO DO SKU FINAL"])
        for label, value in parameters:
            row = parameters_table.add_row().cells
            _cell(row[0], label, bold=True)
            _cell(row[1], value)

    _section(document, "COMPOSIÇÃO / MATERIAL DE ORIGEM A CONSUMIR")
    inputs = document.add_table(rows=1, cols=4)
    inputs.style = "Table Grid"
    _table_header(inputs, ["CÓDIGO", "COMPOSIÇÃO", "QTDE", "UN."])
    for item in order.get("inputs") or []:
        row = inputs.add_row().cells
        _cell(row[0], item.get("sku"))
        _cell(row[1], item.get("descricao"))
        _cell(row[2], _quantity(item.get("quantidade_planejada")), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row[3], item.get("unidade"), align=WD_ALIGN_PARAGRAPH.CENTER)

    _section(document, "PROCESSOS DE PRODUÇÃO - SERRALHERIA")
    processes = order.get("process_snapshot") or []
    if not processes:
        processes = ["Conferir material de origem e identificação da O.P."] + [
            f"Adaptar {label} para {value}." for label, value in parameters
        ] + ["Conferir produto final, acabamento e liberar para estoque."]
    process_table = document.add_table(rows=1, cols=6)
    process_table.style = "Table Grid"
    _table_header(process_table, ["#", "ATIVIDADE", "OK/NOK", "RESPONSÁVEL", "INÍCIO", "FIM"])
    for index, process in enumerate(processes, start=1):
        row = process_table.add_row().cells
        _cell(row[0], index, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row[1], process)
        _cell(row[2], "", align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row[3], "")
        _cell(row[4], "")
        _cell(row[5], "")

    _section(document, "OBSERVAÇÕES")
    notes = document.add_table(rows=1, cols=1)
    notes.style = "Table Grid"
    _cell(notes.rows[0].cells[0], order.get("observacoes") or "")

    _section(document, "VISTOS DE QUALIDADE E GESTÃO")
    signatures = document.add_table(rows=2, cols=2)
    signatures.style = "Table Grid"
    _cell(signatures.rows[0].cells[0], "ASSINATURA DA QUALIDADE", bold=True, fill=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(signatures.rows[0].cells[1], "ASSINATURA DE GESTÃO", bold=True, fill=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(signatures.rows[1].cells[0], "\n\n_______________________________", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(signatures.rows[1].cells[1], "\n\n_______________________________", align=WD_ALIGN_PARAGRAPH.CENTER)

    _section(document, "DESENHO TÉCNICO (SE APLICÁVEL)")
    drawing = document.add_table(rows=1, cols=1)
    drawing.style = "Table Grid"
    _cell(drawing.rows[0].cells[0], "\n\n\n\n")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
