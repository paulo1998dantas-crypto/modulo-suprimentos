import logging
import os
import struct
import tempfile
from copy import deepcopy
from datetime import datetime, timedelta

import pypdfium2 as pdfium
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from composicao import resolver_composicao_final
from config import TEMPLATE_OS, TEMPLATE_REQUISICAO_EXPEDICAO
from os_template import encontrar_linha_cabecalho, mapear_tabelas_os

logger = logging.getLogger(__name__)


def _resolve_unique_path(path):
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    for idx in range(1, 100):
        candidato = f"{base} - R{idx:02d}{ext}"
        if not os.path.exists(candidato):
            return candidato
    return path


def _sanitize_nome_arquivo(texto):
    texto = "" if texto is None else str(texto).strip()
    invalid = '<>:"/\\\\|?*'
    for ch in invalid:
        texto = texto.replace(ch, " ")
    return " ".join(texto.split())


def _limitar_nome_arquivo(pasta, nome, limite_path=245):
    nome = _sanitize_nome_arquivo(nome) or "documento.docx"
    base, ext = os.path.splitext(nome)
    disponivel = limite_path - len(os.path.abspath(pasta)) - 1
    if disponivel <= len(ext) + 8:
        return f"02 - OS{ext or '.docx'}"
    if len(nome) <= disponivel:
        return nome
    base_limite = max(8, disponivel - len(ext))
    partes = base.split(" - ")
    if len(partes) >= 3:
        sufixo = partes[-1].strip()
        prefixo = " - ".join(partes[:-1]).strip()
        prefixo_limite = base_limite - len(sufixo) - 3
        if sufixo and prefixo_limite >= 8:
            return f"{prefixo[:prefixo_limite].strip(' -')} - {sufixo}{ext}"
    return f"{base[:base_limite].strip(' -')}{ext}"


def _cliente_resumido_nome(cliente, limite=6):
    cliente = _sanitize_nome_arquivo(cliente)
    if not cliente:
        return ""
    primeiro = cliente.split()[0] if cliente.split() else cliente
    return primeiro[:limite].strip(" -")


def _montar_nome_documento_os(titulo_arquivo, cliente, chassi, incluir_cliente=False, cliente_limite=None):
    partes = ["02", titulo_arquivo]
    if incluir_cliente:
        cliente_nome = _cliente_resumido_nome(cliente, cliente_limite) if cliente_limite else _sanitize_nome_arquivo(cliente)
        if cliente_nome:
            partes.append(cliente_nome)
    chassi_nome = _sanitize_nome_arquivo(chassi)
    if chassi_nome:
        partes.append(chassi_nome)
    return " - ".join(parte for parte in partes if parte).strip(" -") + ".docx"


def _safe_save_doc(doc, path):
    try:
        doc.save(path)
        return path
    except Exception as exc:
        try:
            fallback_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            os.makedirs(fallback_dir, exist_ok=True)
            fallback_path = _resolve_unique_path(os.path.join(fallback_dir, os.path.basename(path)))
            doc.save(fallback_path)
            logger.warning("Falha ao salvar em %s. Salvo em %s. Erro: %s", path, fallback_path, exc)
            return fallback_path
        except Exception:
            raise


def _set_cell_text(cell, texto):
    cell.text = "" if texto is None else str(texto)


def _set_cell_align(cell, alignment):
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment


def _formatar_cell_atividade(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.08)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0


def _celulas_unicas_linha(row):
    vistas = set()
    celulas = []
    for cell in row.cells:
        chave = id(cell._tc)
        if chave in vistas:
            continue
        vistas.add(chave)
        celulas.append(cell)
    return celulas


def _set_vertical_merge(cell, mode=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    vmerge = tc_pr.find(qn("w:vMerge"))
    if vmerge is None:
        vmerge = OxmlElement("w:vMerge")
        tc_pr.append(vmerge)
    if mode:
        vmerge.set(qn("w:val"), mode)
    else:
        if qn("w:val") in vmerge.attrib:
            del vmerge.attrib[qn("w:val")]


def _set_os_numero_legacy_unused(cell, numero_os):
    while len(cell.paragraphs) < 2:
        cell.add_paragraph()
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.text = ""
    p0 = cell.paragraphs[0]
    p1 = cell.paragraphs[1]
    p0.add_run("Nº")
    run_num = p1.add_run(str(numero_os))
    run_num.bold = True


def _set_os_numero(cell, numero_os):
    while len(cell.paragraphs) < 2:
        cell.add_paragraph()
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.text = ""
    p0 = cell.paragraphs[0]
    p1 = cell.paragraphs[1]
    p0.add_run("N\u00ba Ordem de Servi\u00e7o")
    run_num = p1.add_run(str(numero_os))
    run_num.bold = True


def _limpar_linhas_apos(tabela, header_index):
    for idx in range(len(tabela.rows) - 1, header_index, -1):
        tabela._tbl.remove(tabela.rows[idx]._tr)


def _formatar_datetime_local(texto):
    if not texto:
        return ""
    texto = str(texto).strip()
    if "T" in texto:
        data, hora = texto.split("T", 1)
        partes = data.split("-")
        if len(partes) == 3:
            ano, mes, dia = partes
            return f"{dia}/{mes}/{ano} - {hora}"
    return texto


def _parse_datetime_local(texto):
    if not texto:
        return None
    texto = str(texto).strip()
    formatos = (
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%dT%H:%M",
        "%d/%m/%Y - %H:%M:%S",
        "%d/%m/%Y - %H:%M",
        "%d/%m/%Y %H:%M:%S",
        "%d/%m/%Y %H:%M",
        "%d/%m/%Y",
    )
    for formato in formatos:
        try:
            return datetime.strptime(texto, formato)
        except Exception:
            continue
    try:
        return datetime.fromisoformat(texto.replace("Z", "+00:00")).replace(tzinfo=None)
    except Exception:
        return None


def _formatar_data_necessidade(texto):
    dt = _parse_datetime_local(texto)
    if dt is not None:
        return (dt - timedelta(days=1)).strftime("%d/%m/%Y")
    texto_formatado = _formatar_datetime_local(texto)
    if " - " in texto_formatado:
        return texto_formatado.split(" - ", 1)[0]
    return texto_formatado


def _configurar_cabecalho_requisicao(doc, refs, dados):
    if refs.get("cabecalho") is not None and refs["cabecalho"] < len(doc.tables):
        tabela_cabecalho = doc.tables[refs["cabecalho"]]
        if tabela_cabecalho.rows and len(tabela_cabecalho.rows[0].cells) > 1:
            titulo_cabecalho = dados.get("cabecalho_requisicao", "ORDEM DE REQUISI\u00c7\u00c3O")
            _set_cell_text(tabela_cabecalho.cell(0, 1), titulo_cabecalho)

    if refs.get("dados") is None or refs["dados"] >= len(doc.tables):
        return

    tabela_dados = doc.tables[refs["dados"]]
    if tabela_dados.rows:
        titulo = dados.get("titulo_requisicao", "DADOS DA ORDEM DE REQUISI\u00c7\u00c3O")
        for cell in tabela_dados.rows[0].cells:
            _set_cell_text(cell, titulo)
    if len(tabela_dados.rows) > 3 and len(tabela_dados.rows[3].cells) >= 4:
        _set_cell_text(tabela_dados.cell(3, 0), "DATA DE NECESSIDADE:")
        _set_cell_text(
            tabela_dados.cell(3, 1),
            _formatar_data_necessidade(dados.get("previsao_inicio", "")),
        )
        _set_cell_text(tabela_dados.cell(3, 2), "")
        _set_cell_text(tabela_dados.cell(3, 3), "")


def _preencher_tabela_produtos(tabela, itens):
    _limpar_linhas_apos(tabela, 0)
    for item in itens:
        row = tabela.add_row().cells
        _set_cell_text(row[0], item.get("codigo", ""))
        _set_cell_text(row[1], item.get("descricao", ""))
        _set_cell_text(row[2], item.get("qtd", ""))
        _set_cell_text(row[3], item.get("serie", ""))
        if len(row) >= 6:
            _set_cell_text(row[4], item.get("visto", ""))
            _set_cell_text(row[5], item.get("unidade", ""))
        else:
            _set_cell_text(row[4], item.get("unidade", ""))
        for cell in row:
            _set_cell_align(cell, WD_ALIGN_PARAGRAPH.CENTER)


def _preencher_tabela_componentes(tabela, composicao):
    _limpar_linhas_apos(tabela, 1)
    for comp in composicao or []:
        row = tabela.add_row().cells
        try:
            level = int(comp.get("level", 0) or 0)
        except Exception:
            level = 0
        prefixo = " >" * level
        descricao = f"{prefixo} {comp.get('descricao', '')}".strip() if level else comp.get("descricao", "")
        _set_cell_text(row[0], comp.get("codigo", ""))
        _set_cell_text(row[1], descricao)
        _set_cell_text(row[2], comp.get("qtd", ""))
        _set_cell_text(row[3], comp.get("unidade", ""))
        for cell in row:
            _set_cell_align(cell, WD_ALIGN_PARAGRAPH.CENTER)


def _preencher_tabela_processo(tabela, linhas):
    linhas = list(linhas or [])
    header_idx = encontrar_linha_cabecalho(tabela, "ATIVIDADE")
    if header_idx is None:
        header_idx = 3 if len(tabela.rows) >= 4 else max(len(tabela.rows) - 1, 0)

    modelo_idx = header_idx + 1 if header_idx + 1 < len(tabela.rows) else header_idx
    modelo_idx = max(0, min(modelo_idx, len(tabela.rows) - 1))
    modelo_tr = deepcopy(tabela.rows[modelo_idx]._tr)

    def _clonar_linha_formato():
        novo_tr = deepcopy(modelo_tr)
        tabela._tbl.append(novo_tr)

    atividade_inicio = header_idx + 1
    while len(tabela.rows) > atividade_inicio:
        tabela._tbl.remove(tabela.rows[-1]._tr)

    while len(tabela.rows) < atividade_inicio + max(len(linhas), 1):
        _clonar_linha_formato()

    total_linhas = max(len(linhas), 1)
    for idx in range(atividade_inicio, atividade_inicio + total_linhas):
        row = tabela.rows[idx]
        for cell in _celulas_unicas_linha(row):
            _set_cell_text(cell, "")
            _set_cell_align(cell, WD_ALIGN_PARAGRAPH.CENTER)

    for offset, linha in enumerate(linhas):
        row = tabela.rows[atividade_inicio + offset]
        cells_reais = _celulas_unicas_linha(row)
        if len(cells_reais) > 0:
            _set_cell_text(cells_reais[0], str(offset + 1))
            _set_cell_align(cells_reais[0], WD_ALIGN_PARAGRAPH.CENTER)
        if len(cells_reais) > 1:
            _set_cell_text(cells_reais[1], linha.get("atividade", ""))
            _formatar_cell_atividade(cells_reais[1])


def _limpar_layout(tabela):
    if len(tabela.rows) <= 1 or len(tabela.rows[1].cells) == 0:
        return
    cell = tabela.cell(1, 0)
    cell.text = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""


def _salvar_bitmap_bmp(bitmap, path):
    import numpy as np

    pixels = bitmap.to_numpy()
    if pixels.ndim == 2:
        pixels = np.repeat(pixels[:, :, None], 3, axis=2)
    if pixels.shape[2] < 3:
        pixels = np.repeat(pixels[:, :, :1], 3, axis=2)
    pixels = pixels[:, :, :3]
    if str(bitmap.mode).upper().startswith("RGB"):
        pixels = pixels[:, :, ::-1]
    pixels = np.ascontiguousarray(pixels[::-1])
    altura, largura = pixels.shape[:2]
    padding = (4 - (largura * 3) % 4) % 4
    row_padding = b"\x00" * padding
    tamanho_pixels = (largura * 3 + padding) * altura
    offset = 14 + 40
    with open(path, "wb") as arquivo:
        arquivo.write(struct.pack("<2sIHHI", b"BM", offset + tamanho_pixels, 0, 0, offset))
        arquivo.write(struct.pack("<IIIHHIIIIII", 40, largura, altura, 1, 24, 0, tamanho_pixels, 2835, 2835, 0, 0))
        for linha in pixels:
            arquivo.write(linha.tobytes())
            arquivo.write(row_padding)


def _inserir_layout_pdf(tabela, file_storage):
    if not tabela or not file_storage or not file_storage.filename:
        return

    tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp_pdf.close()
    file_storage.save(tmp_pdf.name)

    pdf = pdfium.PdfDocument(tmp_pdf.name)
    if len(pdf) == 0:
        return

    page = pdf[0]
    try:
        bitmap = page.render(scale=2)
    except Exception as exc:
        logger.warning("Falha ao renderizar PDF para imagem: %s", exc)
        return

    try:
        pil_image = bitmap.to_pil()
        tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp_img.close()
        pil_image.save(tmp_img.name)
        largura_px, altura_px = pil_image.size
    except Exception:
        tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".bmp")
        tmp_img.close()
        _salvar_bitmap_bmp(bitmap, tmp_img.name)
        largura_px, altura_px = bitmap.width, bitmap.height

    cell = tabela.cell(1, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    largura_max = 6.8
    altura_max = 8.4
    proporcao = min(largura_max / largura_px, altura_max / altura_px)
    paragraph.add_run().add_picture(
        tmp_img.name,
        width=Inches(largura_px * proporcao),
        height=Inches(altura_px * proporcao),
    )


def _isolar_layout_na_ultima_pagina(doc, tabela):
    if not tabela:
        return
    body = doc._body._element
    tabela_xml = tabela._tbl
    parent = tabela_xml.getparent()
    if parent is None:
        return
    parent.remove(tabela_xml)
    for quebra_antiga in tabela_xml.xpath('.//w:br[@w:type="page"]'):
        quebra_antiga.getparent().remove(quebra_antiga)
    sect_pr = body.sectPr
    while len(body):
        ultimo = body[-2] if sect_pr is not None and body[-1] is sect_pr else body[-1]
        if ultimo.tag != qn("w:p"):
            break
        if "".join(ultimo.itertext()).strip() or ultimo.xpath(".//w:drawing"):
            break
        body.remove(ultimo)
    quebra = doc.add_paragraph()
    quebra._p.get_or_add_pPr().append(OxmlElement("w:pageBreakBefore"))
    if sect_pr is not None:
        body.insert(body.index(sect_pr), tabela_xml)
    else:
        body.append(tabela_xml)


def _remover_tabela(doc, tabela):
    try:
        tabela._tbl.getparent().remove(tabela._tbl)
    except Exception:
        pass


def _alinhar_descricao_esquerda(tabela, coluna, inicio_linha=0):
    if not tabela:
        return
    for row in tabela.rows[inicio_linha:]:
        if len(row.cells) <= coluna:
            continue
        for paragraph in row.cells[coluna].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _alinhar_tudo_centro(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _alinhar_tabelas_processo(refs, doc):
    for idx in refs.get("processos", {}).values():
        if idx >= len(doc.tables):
            continue
        tabela = doc.tables[idx]
        for row in tabela.rows[2:]:
            for cell_idx, cell in enumerate(_celulas_unicas_linha(row)):
                if cell_idx == 1:
                    _formatar_cell_atividade(cell)
                else:
                    _set_cell_align(cell, WD_ALIGN_PARAGRAPH.CENTER)


def _salvar_documento_os(
    doc,
    numero_os,
    dados,
    titulo_arquivo="O.S",
    incluir_cliente_nome=True,
    cliente_nome_limite=None,
):
    pasta = tempfile.mkdtemp(prefix="modulo-suprimentos-os-")
    cliente = (dados.get("cliente", "") or "").strip()
    chassi = (dados.get("chassis", "") or "").strip()
    nome = _montar_nome_documento_os(
        titulo_arquivo,
        cliente,
        chassi,
        incluir_cliente=incluir_cliente_nome,
        cliente_limite=cliente_nome_limite,
    )
    nome = _limitar_nome_arquivo(pasta, nome)
    path = os.path.join(pasta, nome)
    path = _resolve_unique_path(path)
    return _safe_save_doc(doc, path)


def gerar_os_docx(
    numero_os,
    dados,
    itens,
    componentes,
    processos,
    layout_pdf=None,
    composicao_resolvida=None,
    modo="completa",
    titulo_arquivo="O.S",
    incluir_cliente_nome=True,
    cliente_nome_limite=None,
):
    template_path = (
        TEMPLATE_REQUISICAO_EXPEDICAO
        if modo in {"expedicao", "faturamento_direto"} and os.path.exists(TEMPLATE_REQUISICAO_EXPEDICAO)
        else TEMPLATE_OS
    )
    doc = Document(template_path)
    refs = mapear_tabelas_os(doc)
    processo_preparacao = None
    for nome in processos.keys():
        if "PREPARA" in str(nome).upper():
            processo_preparacao = nome
            break

    if refs.get("cabecalho") is not None:
        _set_os_numero(doc.tables[refs["cabecalho"]].cell(0, 2), numero_os)

    if refs.get("dados") is not None:
        tabela_dados = doc.tables[refs["dados"]]
        _set_cell_text(tabela_dados.cell(1, 1), dados.get("chassis", ""))
        _set_cell_text(tabela_dados.cell(1, 3), dados.get("municipio", ""))
        _set_cell_text(tabela_dados.cell(2, 1), dados.get("cliente", ""))
        _set_cell_text(tabela_dados.cell(2, 3), dados.get("mmv", ""))
        _set_cell_text(tabela_dados.cell(3, 1), _formatar_datetime_local(dados.get("previsao_inicio", "")))
        _set_cell_text(tabela_dados.cell(3, 3), _formatar_datetime_local(dados.get("previsao_termino", "")))

    if modo in {"expedicao", "preparacao", "faturamento_direto"}:
        dados_requisicao = dict(dados)
        if modo == "expedicao":
            dados_requisicao["titulo_requisicao"] = "DADOS DA ORDEM DE REQUISI\u00c7\u00c3O EXPEDI\u00c7\u00c3O"
        elif modo == "preparacao":
            dados_requisicao["titulo_requisicao"] = "DADOS DA ORDEM DE REQUISI\u00c7\u00c3O PREPARA\u00c7\u00c3O"
        elif modo == "faturamento_direto":
            dados_requisicao["cabecalho_requisicao"] = "FATURAMENTO DIRETO"
            dados_requisicao["titulo_requisicao"] = "DADOS DO FATURAMENTO DIRETO"
        _configurar_cabecalho_requisicao(doc, refs, dados_requisicao)

    if refs.get("itens") is not None:
        _preencher_tabela_produtos(doc.tables[refs["itens"]], itens)

    composicao_final = composicao_resolvida or resolver_composicao_final(itens, componentes)
    ocultar_composicao = modo in {"resumida", "producao", "expedicao", "preparacao", "faturamento_direto"}
    ocultar_observacoes = modo in {"mascara", "producao", "expedicao", "preparacao", "faturamento_direto"}

    if modo == "completa":
        ocultar_processos = False
        processos_exibicao = processos
    elif modo == "producao":
        ocultar_processos = False
        processos_exibicao = {
            nome: linhas
            for nome, linhas in processos.items()
            if nome != processo_preparacao
        }
    elif modo == "preparacao":
        ocultar_processos = False
        processos_exibicao = (
            {
                nome: processos.get(nome, [])
                for nome in processos.keys()
                if nome == processo_preparacao
            }
            if processo_preparacao
            else {}
        )
    elif modo in {"expedicao", "faturamento_direto"}:
        ocultar_processos = False
        processos_exibicao = {}
    else:
        ocultar_processos = True
        processos_exibicao = {nome: [] for nome in processos.keys()}

    if not ocultar_composicao and refs.get("composicao") is not None:
        _preencher_tabela_componentes(doc.tables[refs["composicao"]], composicao_final)

    if ocultar_composicao and refs.get("composicao") is not None:
        _remover_tabela(doc, doc.tables[refs["composicao"]])
        refs = mapear_tabelas_os(doc)

    if modo == "mascara":
        indices_remover = []
        if refs.get("observacoes") is not None:
            indices_remover.append(refs["observacoes"])
        indices_remover.extend(refs.get("processos", {}).values())
        for idx in sorted(set(indices_remover), reverse=True):
            if idx < len(doc.tables):
                _remover_tabela(doc, doc.tables[idx])
        refs = mapear_tabelas_os(doc)
        if refs.get("layout") is not None:
            tabela_layout = doc.tables[refs["layout"]]
            _limpar_layout(tabela_layout)
            _inserir_layout_pdf(tabela_layout, layout_pdf)
            _isolar_layout_na_ultima_pagina(doc, tabela_layout)
        _alinhar_tudo_centro(doc)
        if refs.get("itens") is not None and refs["itens"] < len(doc.tables):
            _alinhar_descricao_esquerda(doc.tables[refs["itens"]], 1)
        if refs.get("composicao") is not None and refs["composicao"] < len(doc.tables):
            _alinhar_descricao_esquerda(doc.tables[refs["composicao"]], 1, inicio_linha=1)
        return _salvar_documento_os(
            doc,
            numero_os,
            dados,
            titulo_arquivo=titulo_arquivo,
            incluir_cliente_nome=incluir_cliente_nome,
            cliente_nome_limite=cliente_nome_limite,
        )

    if ocultar_observacoes and refs.get("observacoes") is not None:
        _remover_tabela(doc, doc.tables[refs["observacoes"]])
        refs = mapear_tabelas_os(doc)
    elif refs.get("observacoes") is not None:
        tabela_obs = doc.tables[refs["observacoes"]]
        if len(tabela_obs.rows) > 1:
            _set_cell_text(tabela_obs.cell(1, 0), dados.get("obs_materiais", ""))

    if ocultar_processos:
        indices_processos_vazios = list(refs.get("processos", {}).values())
    else:
        indices_processos_vazios = [
            idx
            for nome, idx in refs.get("processos", {}).items()
            if not (processos_exibicao.get(nome) or [])
        ]
    for idx in sorted(set(indices_processos_vazios), reverse=True):
        if idx < len(doc.tables):
            _remover_tabela(doc, doc.tables[idx])

    refs = mapear_tabelas_os(doc)
    for nome, idx in refs.get("processos", {}).items():
        linhas = processos_exibicao.get(nome, [])
        if idx < len(doc.tables) and linhas:
            _preencher_tabela_processo(doc.tables[idx], linhas)

    if dados.get("obs"):
        doc.add_paragraph("")
        doc.add_paragraph(f"OBS FINAL: {dados.get('obs')}")

    if refs.get("layout") is not None and refs["layout"] < len(doc.tables):
        tabela_layout = doc.tables[refs["layout"]]
        _limpar_layout(tabela_layout)
        _inserir_layout_pdf(tabela_layout, layout_pdf)
        _isolar_layout_na_ultima_pagina(doc, tabela_layout)

    _alinhar_tudo_centro(doc)
    refs = mapear_tabelas_os(doc)

    if refs.get("itens") is not None and refs["itens"] < len(doc.tables):
        _alinhar_descricao_esquerda(doc.tables[refs["itens"]], 1)

    if refs.get("composicao") is not None and refs["composicao"] < len(doc.tables):
        _alinhar_descricao_esquerda(doc.tables[refs["composicao"]], 1, inicio_linha=1)

    if refs.get("dados") is not None and refs["dados"] < len(doc.tables):
        tabela_dados = doc.tables[refs["dados"]]
        for row_idx, col_idx in [(1, 1), (2, 1), (3, 1)]:
            if row_idx < len(tabela_dados.rows) and col_idx < len(tabela_dados.rows[row_idx].cells):
                _set_cell_align(tabela_dados.rows[row_idx].cells[col_idx], WD_ALIGN_PARAGRAPH.LEFT)

    _alinhar_tabelas_processo(refs, doc)
    return _salvar_documento_os(
        doc,
        numero_os,
        dados,
        titulo_arquivo=titulo_arquivo,
        incluir_cliente_nome=incluir_cliente_nome,
        cliente_nome_limite=cliente_nome_limite,
    )
