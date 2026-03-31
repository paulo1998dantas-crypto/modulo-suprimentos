import os
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from composicao import expandir_composicao_itens, normalizar_componentes
from config import TEMPLATE_WORD, TEMPLATE_OS, pasta_ano, OS_COMPONENTES_FILE
import logging


def _substituir_paragrafo(paragrafo, chave, valor):
    if chave not in paragrafo.text:
        return

    texto_completo = "".join(run.text for run in paragrafo.runs)

    # Caso o placeholder esteja inteiro em um run, substitui preservando formatacao
    substituiu_em_run = False
    for run in paragrafo.runs:
        if chave in run.text:
            run.text = run.text.replace(chave, str(valor))
            substituiu_em_run = True

    if substituiu_em_run:
        return

    # Fallback: placeholder pode estar dividido em varios runs (Word costuma fazer isso).
    # Reconstroi o paragrafo com o texto ja substituido. Isso pode perder formatacao
    # apenas nesse paragrafo, mas garante a substituicao correta.
    texto_substituido = texto_completo.replace(chave, str(valor))
    if not paragrafo.runs:
        paragrafo.text = texto_substituido
        return

    paragrafo.runs[0].text = texto_substituido
    for run in paragrafo.runs[1:]:
        run.text = ""


def _format_currency(valor):
    if isinstance(valor, (int, float)):
        return _format_brl(float(valor))

    if valor is None:
        return ""

    texto = str(valor).strip()
    if texto == "":
        return ""

    if texto.startswith("R$"):
        return texto

    try:
        numero = float(texto.replace(".", "").replace(",", "."))
    except ValueError:
        return texto

    return _format_brl(numero)


def _format_brl(numero):
    inteiro = int(round(numero * 100)) // 100
    centavos = int(round(numero * 100)) % 100
    inteiro_str = f"{inteiro:,}".replace(",", ".")
    return f"R$ {inteiro_str},{centavos:02d}"


def _format_date(texto):
    if not texto:
        return ""

    texto = str(texto).strip()
    if texto == "":
        return ""

    try:
        if "-" in texto:
            data = datetime.strptime(texto, "%Y-%m-%d")
            return data.strftime("%d/%m/%Y")
    except ValueError:
        pass

    return texto


def construir_nome_oc(numero_oc, fornecedor, dados_pedido):
    ano_atual = datetime.now().strftime("%y")
    valor_final = _format_currency(dados_pedido.get("total_pedido", ""))
    if valor_final and not valor_final.startswith("R$"):
        valor_final = f"R$ {valor_final}"
    return f"{numero_oc}-{ano_atual} - {fornecedor} - {valor_final}.docx"


def substituir(doc, chave, valor):

    def _processar_container(container):
        for p in container.paragraphs:
            _substituir_paragrafo(p, chave, valor)

        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _substituir_paragrafo(p, chave, valor)

    _processar_container(doc)

    # inclui cabecalhos e rodapes
    for section in doc.sections:
        _processar_container(section.header)
        _processar_container(section.footer)
        _processar_container(section.first_page_header)
        _processar_container(section.first_page_footer)
        _processar_container(section.even_page_header)
        _processar_container(section.even_page_footer)


def gerar_word(numero_oc, fornecedor, dados_pedido, itens, incluir_composicao=True):

    doc = Document(TEMPLATE_WORD)

    substituir(doc, "{OC_NUMERO}", numero_oc)
    substituir(doc, "{DATA}", datetime.now().strftime("%d/%m/%Y"))
    substituir(doc, "{FORNECEDOR}", fornecedor)
    substituir(doc, "{CNPJ}", dados_pedido.get("cnpj", ""))
    substituir(doc, "{BAIRRO}", dados_pedido.get("bairro", ""))
    substituir(doc, "{CIDADE}", dados_pedido.get("cidade", ""))
    substituir(doc, "{UF}", dados_pedido.get("uf", ""))
    substituir(doc, "{EMAIL}", dados_pedido.get("email", ""))
    substituir(doc, "{RAZAO_SOCIAL}", dados_pedido.get("razao_social", ""))
    substituir(doc, "{ENDERECO}", dados_pedido.get("endereco", ""))
    substituir(doc, "{CEP}", dados_pedido.get("cep", ""))
    substituir(doc, "{TELEFONE}", dados_pedido.get("telefone", ""))
    substituir(doc, "{PREVISAO}", _format_date(dados_pedido.get("previsao", "")))
    substituir(doc, "{TIPO_FRETE}", dados_pedido.get("tipo_frete", ""))
    substituir(doc, "{FRETE}", _format_currency(dados_pedido.get("frete", "")))
    substituir(doc, "{TOTAL_ITENS}", _format_currency(dados_pedido.get("total_itens", "")))
    substituir(doc, "{TOTAL_PEDIDO}", _format_currency(dados_pedido.get("total_pedido", "")))
    substituir(doc, "{FORMA DE PAGAMENTO}", dados_pedido.get("forma_pagamento", ""))
    substituir(doc, "{FORMA_PAGAMENTO}", dados_pedido.get("forma_pagamento", ""))
    substituir(doc, "{PRAZO}", dados_pedido.get("prazo", ""))
    substituir(doc, "{VENCIMENTO}", _format_date(dados_pedido.get("vencimento", "")))
    substituir(doc, "{OBS}", dados_pedido.get("obs", ""))

    tabela = None

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if ("{CODIGO}" in cell.text or "{DESCRICAO}" in cell.text or
                        "{UNIDADE}" in cell.text or "{QTD}" in cell.text or
                        "{VALOR_UNIT}" in cell.text or "{DESCONTO}" in cell.text or
                        "{TOTAL_ITEM}" in cell.text):
                    tabela = table
                    break
            if tabela:
                break
        if tabela:
            break

    if not tabela:
        raise Exception("Tabela de produtos nao encontrada no template")

    # Centraliza cabecalho da coluna de codigo
    for p in tabela.rows[0].cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    template_row = tabela.rows[1]

    def _apply_cell_style(src_cell, dst_cell):
        dst_cell.vertical_alignment = src_cell.vertical_alignment or WD_ALIGN_VERTICAL.CENTER
        src_p = src_cell.paragraphs[0] if src_cell.paragraphs else None
        for p in dst_cell.paragraphs:
            if src_p is not None:
                p.alignment = src_p.alignment
                try:
                    p.style = src_p.style
                except Exception:
                    pass
                if src_p.runs:
                    src_run = src_p.runs[0]
                    for run in p.runs:
                        if src_run.font.name:
                            run.font.name = src_run.font.name
                        if src_run.font.size:
                            run.font.size = src_run.font.size
                        if src_run.bold is not None:
                            run.bold = src_run.bold
                        if src_run.italic is not None:
                            run.italic = src_run.italic

    def _set_cell_text(cell, texto, template_cell):
        cell.text = "" if texto is None else str(texto)
        _apply_cell_style(template_cell, cell)

    for idx, item in enumerate(itens):
        row_cells = template_row.cells if idx == 0 else tabela.add_row().cells
        _set_cell_text(row_cells[0], item["codigo"], template_row.cells[0])
        _set_cell_text(row_cells[1], item["descricao"], template_row.cells[1])
        _set_cell_text(row_cells[2], item["unidade"], template_row.cells[2])
        _set_cell_text(row_cells[3], item["qtd"], template_row.cells[3])
        _set_cell_text(row_cells[4], _format_brl(item["valor"]), template_row.cells[4])
        _set_cell_text(row_cells[5], _format_brl(item["desconto"]), template_row.cells[5])
        _set_cell_text(row_cells[6], _format_brl(item["total"]), template_row.cells[6])

    # Composicao de materiais (B.O.M) usando a tabela do template, se existir
    componentes = {}
    if os.path.exists(OS_COMPONENTES_FILE):
        try:
            with open(OS_COMPONENTES_FILE, "r", encoding="utf-8") as f:
                componentes = json.load(f) or {}
        except Exception:
            componentes = {}
    componentes = normalizar_componentes(componentes)

    composicao_rows = []
    if incluir_composicao:
        composicao_rows = expandir_composicao_itens(itens, componentes)

    tabela_comp = None
    template_row_comp = None
    template_row_idx = None
    for table in doc.tables:
        for idx_row, row in enumerate(table.rows):
            for cell in row.cells:
                # Identifica a tabela de composição pelo placeholder exclusivo {ITEM_M}
                if "{ITEM_M}" in cell.text:
                    tabela_comp = table
                    template_row_comp = row
                    template_row_idx = idx_row
                    break
            if tabela_comp:
                break
        if tabela_comp:
            break

    if composicao_rows and tabela_comp and template_row_comp and template_row_idx is not None:
        while len(tabela_comp.rows) > (template_row_idx + 1):
            tabela_comp._tbl.remove(tabela_comp.rows[-1]._tr)

        template_texts = [c.text for c in template_row_comp.cells]

        def _copiar_estilo_celula(src_cell, dst_cell):
            try:
                dst_cell.vertical_alignment = src_cell.vertical_alignment
            except Exception:
                pass
            if not src_cell.paragraphs or not dst_cell.paragraphs:
                return
            src_p = src_cell.paragraphs[0]
            for p in dst_cell.paragraphs:
                p.alignment = src_p.alignment
                try:
                    p.style = src_p.style
                except Exception:
                    pass
                if src_p.runs:
                    src_run = src_p.runs[0]
                    for run in p.runs:
                        if src_run.font.name:
                            run.font.name = src_run.font.name
                        if src_run.font.size:
                            run.font.size = src_run.font.size
                        if src_run.bold is not None:
                            run.bold = src_run.bold
                        if src_run.italic is not None:
                            run.italic = src_run.italic

        for idx, comp in enumerate(composicao_rows):
            row_cells = template_row_comp.cells if idx == 0 else tabela_comp.add_row().cells
            for i, c in enumerate(row_cells):
                if i < len(template_texts):
                    c.text = template_texts[i]
                texto = c.text
                texto = texto.replace("{ITEM_M}", str(comp.get("item", "")))
                texto = texto.replace("{CODIGO}", str(comp.get("codigo", "")))
                texto = texto.replace("{DESCRICAO}", str(comp.get("descricao", "")))
                texto = texto.replace("{UNIDADE}", str(comp.get("unidade", "")))
                texto = texto.replace("{QTD}", str(comp.get("qtd", "")))
                c.text = texto
            for c_idx in range(min(len(template_row_comp.cells), len(row_cells))):
                _copiar_estilo_celula(template_row_comp.cells[c_idx], row_cells[c_idx])

    # Se nao houver composicao (ou se estiver desativada), remove a tabela do template
    if (not composicao_rows or not incluir_composicao) and tabela_comp is not None:
        try:
            tabela_comp._tbl.getparent().remove(tabela_comp._tbl)
        except Exception:
            pass

    pasta = pasta_ano()

    nome = construir_nome_oc(numero_oc, fornecedor, dados_pedido)
    path = os.path.join(pasta, nome)

    # Centraliza texto dentro de todas as celulas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Mantem colunas de descricao alinhadas à esquerda
    for table in doc.tables:
        desc_idx = None
        if table.rows:
            header_cells = table.rows[0].cells
            for idx, cell in enumerate(header_cells):
                texto = (cell.text or "").upper()
                if "DESCRI" in texto:
                    desc_idx = idx
                    break
        if desc_idx is not None:
            for row in table.rows:
                if desc_idx < len(row.cells):
                    for p in row.cells[desc_idx].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Forca descricao da composicao (dados) alinhada a esquerda
    for table in doc.tables:
        is_comp = False
        for row in table.rows:
            for cell in row.cells:
                texto = (cell.text or "").upper()
                if "COMPOSI" in texto and "MATERIA" in texto:
                    is_comp = True
                    break
            if is_comp:
                break
        if not is_comp or not table.rows:
            continue
        header_row_idx = None
        desc_idx = None
        for r_idx, row in enumerate(table.rows):
            row_text = " ".join([(c.text or "") for c in row.cells]).upper()
            if "DESCRI" in row_text:
                header_row_idx = r_idx
                for c_idx, cell in enumerate(row.cells):
                    if "DESCRI" in (cell.text or "").upper():
                        desc_idx = c_idx
                        break
                break
        if header_row_idx is None or desc_idx is None:
            continue
        for row in table.rows[header_row_idx + 1:]:
            if desc_idx < len(row.cells):
                for p in row.cells[desc_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
# Campos de endereco/contato e dados do fornecedor devem ficar à esquerda
    palavras_esq = [
        "ENDERE", "BAIRRO", "CIDADE", "CEP", "TELEFONE", "E-MAIL", "EMAIL",
        "DADOS DO FORNECEDOR", "NOME FANTASIA", "RAZAO SOCIAL", "RAZÃO SOCIAL",
        "CNPJ", "CPF", "CNPJ/CPF", "CIDADE/UF",
    ]
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join([c.text or "" for c in row.cells]).upper()
            if any(p in row_text for p in palavras_esq):
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    path = _safe_save_doc(doc, path)

    return path




logger = logging.getLogger(__name__)

def _safe_save_doc(doc, path):
    try:
        doc.save(path)
        return path
    except Exception as exc:
        try:
            fallback_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            os.makedirs(fallback_dir, exist_ok=True)
            fallback_path = os.path.join(fallback_dir, os.path.basename(path))
            doc.save(fallback_path)
            logger.warning("Falha ao salvar em %s. Salvo em %s. Erro: %s", path, fallback_path, exc)
            return fallback_path
        except Exception:
            raise
